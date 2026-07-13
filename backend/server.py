from fastapi import FastAPI, APIRouter, HTTPException, Depends, Request, Response, Body, Query, WebSocket, WebSocketDisconnect, UploadFile, File
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import httpx
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, Union
import uuid
import random
import secrets
from datetime import datetime, timezone, timedelta, date
import re
import base64
import asyncio
import hmac
import hashlib
import resend
import json
from dateutil.relativedelta import relativedelta
try:
    from apscheduler.schedulers.asyncio import AsyncIOScheduler
    _APSCHEDULER_AVAILABLE = True
except ImportError:
    _APSCHEDULER_AVAILABLE = False
from google import genai as google_genai
from google.genai import types as genai_types
import hafizz
try:
    from slowapi import Limiter, _rate_limit_exceeded_handler
    from slowapi.util import get_remote_address
    from slowapi.errors import RateLimitExceeded
    _RATE_LIMIT_AVAILABLE = True
except ImportError:
    _RATE_LIMIT_AVAILABLE = False

try:
    from better_profanity import profanity as _profanity_lib
    _profanity_lib.load_censor_words()
    _PROFANITY_AVAILABLE = True
except ImportError:
    _PROFANITY_AVAILABLE = False

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url, minPoolSize=1, maxPoolSize=10, serverSelectionTimeoutMS=10000)
db = client[os.environ['DB_NAME']]
users_collection = db.users
docs_collection = db.documents

# ============ CLOUDFLARE R2 ============
import boto3
from botocore.config import Config as BotoConfig

R2_ACCOUNT_ID = os.environ.get('R2_ACCOUNT_ID', '')
R2_ACCESS_KEY = os.environ.get('R2_ACCESS_KEY', '')
R2_SECRET_KEY = os.environ.get('R2_SECRET_KEY', '')
R2_BUCKET = os.environ.get('R2_BUCKET', 'zet-studio-international')
R2_PUBLIC_URL = os.environ.get('R2_PUBLIC_URL', 'https://pub-c3e7d4d3e4bd4819871a4bf4809a2a9b.r2.dev')

def _get_r2():
    return boto3.client(
        's3',
        endpoint_url=f'https://{R2_ACCOUNT_ID}.r2.cloudflarestorage.com',
        aws_access_key_id=R2_ACCESS_KEY,
        aws_secret_access_key=R2_SECRET_KEY,
        config=BotoConfig(signature_version='s3v4'),
        region_name='auto',
    )

app = FastAPI()
CEO_EMAIL = "muhammadbahaddinyilmaz@gmail.com"
ADMIN_EMAILS = {"info@zetstudiointl.com", "support@zetstudiointl.com", "ideas@zetstudiointl.com"}
api_router = APIRouter(prefix="/api")

# ============ RATE LIMITING ============
if _RATE_LIMIT_AVAILABLE:
    limiter = Limiter(key_func=get_remote_address, default_limits=["200/minute"])
    app.state.limiter = limiter
    app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ============ PROFANITY FILTER ============
EXTRA_BANNED_TR = [
    "orospu", "amk", "bok", "sik", "pic", "göt", "piç", "oç", "kahpe",
    "ibne", "bok", "lan", "yarrak", "sikerim", "amına", "orospu",
]

def contains_profanity(text: str) -> bool:
    if not text:
        return False
    tl = text.lower()
    if any(w in tl for w in EXTRA_BANNED_TR):
        return True
    if _PROFANITY_AVAILABLE:
        return _profanity_lib.contains_profanity(text)
    return False

def sanitize_text(text: str, max_len: int = 2000) -> str:
    """Temel XSS / injection temizliği + uzunluk sınırı."""
    import re, html
    if not text:
        return ""
    text = html.escape(text[:max_len])
    text = re.sub(r'javascript:', '', text, flags=re.IGNORECASE)
    text = re.sub(r'on\w+\s*=', '', text, flags=re.IGNORECASE)
    return text.strip()

# ── Basit in-memory rate limiter ──────────────────────────────────────────────
_rl_store: dict = {}  # { "{user_id}:{action}": [timestamp, ...] }

def check_rate_limit(user_id: str, action: str, limit: int, window: int):
    """window saniye içinde limit kadar istek geçince 429 fırlatır."""
    import time
    key = f"{user_id}:{action}"
    now = time.time()
    timestamps = [t for t in _rl_store.get(key, []) if now - t < window]
    if len(timestamps) >= limit:
        raise HTTPException(status_code=429, detail=f"Çok fazla istek. Lütfen {window} saniye bekleyin.")
    timestamps.append(now)
    _rl_store[key] = timestamps

# ============ GEMINI RETRY HELPER ============

async def gemini_generate(client, model: str, contents, config, max_retries: int = 3):
    """generate_content wrapper with retry on 503 (overloaded)."""
    import time
    last_exc = None
    for attempt in range(1, max_retries + 1):
        try:
            return await asyncio.to_thread(
                client.models.generate_content,
                model=model,
                contents=contents,
                config=config,
            )
        except Exception as e:
            last_exc = e
            err_str = str(e).lower()
            if "503" in err_str or "overloaded" in err_str or "service unavailable" in err_str:
                if attempt < max_retries:
                    logging.warning(f"Gemini 503, deneme {attempt}/{max_retries} — 3s bekleniyor")
                    await asyncio.sleep(3)
                    continue
            raise  # non-503 hataları direkt fırlat
    raise last_exc

_FAKE_DOMAINS = {'example.com', 'example.org', 'example.net', 'localhost', 'test.com', 'test.org', 'placeholder.com'}

def extract_sources(resp) -> list:
    """Gemini grounding_metadata'dan web kaynaklarını çıkarır."""
    from urllib.parse import urlparse
    sources = []
    try:
        candidates = getattr(resp, 'candidates', None) or []
        if not candidates:
            return sources
        gm = getattr(candidates[0], 'grounding_metadata', None)
        if not gm:
            return sources
        chunks = getattr(gm, 'grounding_chunks', None) or []
        seen = set()
        for chunk in chunks:
            web = getattr(chunk, 'web', None)
            if not web:
                continue
            url   = getattr(web, 'uri',   '') or ''
            title = getattr(web, 'title', '') or url
            if not url or not url.startswith(('http://', 'https://')):
                continue
            try:
                netloc = urlparse(url).netloc.lower().removeprefix('www.')
                if netloc in _FAKE_DOMAINS:
                    continue
            except Exception:
                continue
            if url not in seen:
                seen.add(url)
                sources.append({"title": title, "url": url})
    except Exception:
        pass
    return sources

# ============ MODELS ============

def normalize_subscription(sub: Any) -> Dict[str, Any]:
    """Subscription'ı her zaman dict formatına normalize et."""
    if sub is None:
        return {"plan": "free", "status": "active"}
    if isinstance(sub, str):
        return {"plan": sub if sub else "free", "status": "active"}
    if isinstance(sub, dict):
        return sub
    return {"plan": "free", "status": "active"}

def sanitize_user_doc(user: dict) -> dict:
    """MongoDB user belgesini Pydantic User modeline geçmeden önce güvenli hale getirir."""
    u = dict(user)
    u.pop("_id", None)

    # Subscription alanlarını normalize et
    for field in ("subscription", "mindshare_subscription", "judge_subscription"):
        if field in u:
            u[field] = normalize_subscription(u[field])

    # Geriye dönük uyumluluk: eski alanlar → yeni alanlar
    u.setdefault("mindshare_credits", u.get("credits", 0))
    u.setdefault("judge_credits", 0)
    u.setdefault("mindshare_rank", u.get("rank", "iron"))
    u.setdefault("judge_rank", "iron")
    u.setdefault("mindshare_xp", u.get("quest_xp", u.get("xp", 0)))
    u.setdefault("judge_xp", 0)
    u.setdefault("zc_balance", u.get("sp_balance", 0))
    u.setdefault("name", "")
    u.setdefault("bio", "")

    # Sayısal alanları int'e zorla
    for int_field in ("credits", "bonus_credits", "sp_balance", "zc_balance",
                      "xp", "quest_xp", "mindshare_credits", "judge_credits",
                      "mindshare_xp", "judge_xp", "followers_count", "following_count",
                      "active_time_seconds"):
        if int_field in u:
            try:
                u[int_field] = int(u[int_field])
            except (TypeError, ValueError):
                u[int_field] = 0

    # Bool alanlarını normalize et
    for bool_field in ("picture_custom", "name_custom", "creative_station",
                       "identity_verified", "needs_onboarding", "cancel_pending"):
        if bool_field in u:
            u[bool_field] = bool(u[bool_field])

    return u


class User(BaseModel):
    model_config = {"extra": "ignore", "arbitrary_types_allowed": True}

    # Sadece bunlar zorunlu
    user_id: str
    email: str

    # Profil
    name: Optional[str] = ""
    picture: Optional[str] = None
    username: Optional[str] = None
    display_name: Optional[str] = None
    bio: Optional[str] = ""
    picture_custom: Optional[bool] = False
    name_custom: Optional[bool] = False

    # Subscription — sanitize_user_doc() tarafından normalize edilmiş dict gelir
    subscription: Optional[Dict[str, Any]] = None
    subscription_date: Optional[Any] = None
    cancel_pending: Optional[bool] = False
    cancel_requested_at: Optional[str] = None
    cancel_token: Optional[str] = None

    # Mindshare specific
    mindshare_credits: Optional[int] = 0
    mindshare_subscription: Optional[Dict[str, Any]] = None
    mindshare_settings: Optional[Any] = None
    mindshare_rank: Optional[str] = "iron"
    mindshare_xp: Optional[int] = 0

    # Judge specific
    judge_credits: Optional[int] = 0
    judge_subscription: Optional[Dict[str, Any]] = None
    judge_settings: Optional[Any] = None
    judge_rank: Optional[str] = "iron"
    judge_xp: Optional[int] = 0

    # Eski alanlar — geriye dönük uyumluluk
    credits: Optional[int] = 0
    bonus_credits: Optional[int] = 0
    sp_balance: Optional[int] = 0
    zc_balance: Optional[int] = 0
    rank: Optional[str] = "iron"
    xp: Optional[int] = 0
    quest_xp: Optional[int] = 0
    completed_quests: Optional[List[Any]] = []

    # Creative Station
    creative_station: Optional[bool] = False
    cs_expires_at: Optional[str] = None

    # Social
    verified_type: Optional[str] = None
    following: Optional[List[Any]] = []
    followers: Optional[List[Any]] = []
    following_count: Optional[int] = 0
    followers_count: Optional[int] = 0
    identity_verified: Optional[bool] = False

    # Activity
    active_time_seconds: Optional[int] = 0
    last_heartbeat: Optional[str] = None
    created_at: Optional[Any] = None
    needs_onboarding: Optional[bool] = False
    inventory: Optional[List[Dict[str, Any]]] = []
    last_daily_case: Optional[str] = None

class Document(BaseModel):
    doc_id: str = Field(default_factory=lambda: f"doc_{uuid.uuid4().hex[:12]}")
    user_id: str
    title: str
    doc_type: str = "document"
    content: dict = Field(default_factory=dict)
    pages: List[dict] = Field(default_factory=lambda: [{"page_id": "page_1", "content": {}}])
    pinned: bool = False
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class DocumentCreate(BaseModel):
    title: str
    doc_type: str = "document"
    file_data: Optional[str] = None  # Base64 PDF data for editing

class NoteCreate(BaseModel):
    content: str
    reminder_time: Optional[str] = None  # ISO datetime string for alarm

class SubscriptionUpdate(BaseModel):
    plan: str  # 'free', 'plus', 'pro', 'creative_station'
    action: str  # 'subscribe' or 'cancel'
    billing_cycle: Optional[str] = "monthly"  # 'monthly' or 'yearly'

class SubscriptionCreate(BaseModel):
    plan: str                       # pro, ultra, creative_station, entertainment_pocket
    billing_cycle: str = "monthly"  # monthly, yearly
    payment_provider: str           # lemonsqueezy, stripe
    external_subscription_id: str

class CheckoutRequest(BaseModel):
    plan: str
    billing_cycle: str = "monthly"

class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    subtitle: Optional[str] = None
    content: Optional[dict] = None
    pages: Optional[List[dict]] = None
    pinned: Optional[bool] = None
    settings: Optional[dict] = None  # per-document editor settings (margins, font, bg...)

class QuickNote(BaseModel):
    note_id: str = Field(default_factory=lambda: f"note_{uuid.uuid4().hex[:12]}")
    user_id: str
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class QuickNoteCreate(BaseModel):
    content: str
    reminder_time: Optional[str] = None
    notebook_id: Optional[str] = None

class NoteContentUpdate(BaseModel):
    content: str

class NotebookCreate(BaseModel):
    name: str
    color: Optional[str] = "#292F91"

class NotebookUpdate(BaseModel):
    name: Optional[str] = None
    color: Optional[str] = None
    pinned: Optional[bool] = None
    order: Optional[int] = None

class NotebookSetPassword(BaseModel):
    password: str

class NotebookVerifyPassword(BaseModel):
    password: str

class ChatMessage(BaseModel):
    role: str
    content: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

ZETA_MODEL_MAP = {
    "spark": ("gemini-2.5-flash", 1024),    # hızlı + ucuz
    "prime": ("gemini-2.5-flash", 4096),    # dengeli
    "aziz": ("gemini-2.5-pro", 8192),       # en güçlü
}

class ZetaChatRequest(BaseModel):
    message: str
    doc_id: Optional[str] = None
    session_id: Optional[str] = None
    document_content: Optional[str] = None
    image: Optional[str] = None
    mode: Optional[str] = "fast"
    mood: Optional[str] = "professional"
    emoji_level: Optional[str] = "medium"
    custom_prompt: Optional[str] = ""
    personality: Optional[str] = "normal"
    is_ceo: Optional[bool] = False
    model: Optional[str] = "prime"  # spark | prime | aziz
    canvas_context: Optional[str] = None  # live editor state snapshot from frontend
    shared_memory: Optional[bool] = False  # ortak hafıza: diğer oturumlardan bağlam

class ParseSourceRequest(BaseModel):
    filename: str
    content_base64: str
    mime_type: str = "application/octet-stream"

class ZetaAutoWriteRequest(BaseModel):
    prompt: str
    page_count: int = 1
    writing_style: str = "profesyonel"  # akademik, yaratici, resmi, gunluk, hikaye, profesyonel

class ZetaDeepAnalysisRequest(BaseModel):
    topic: str
    document_content: Optional[str] = None

class ZetaImageRequest(BaseModel):
    prompt: str
    reference_image: Optional[str] = None
    pro: Optional[bool] = False
    aspect_ratio: Optional[str] = "16:9"

class ZetaDocEditRequest(BaseModel):
    user_request: str
    page_elements: List[Dict[str, Any]] = []
    page_size: Dict[str, int] = {"width": 794, "height": 1123}
    page_index: int = 0
    doc_id: Optional[str] = None
    is_ceo: bool = False
    all_pages: Optional[List[Dict[str, Any]]] = None  # [{page_index, elements}]
    attached_image_b64: Optional[str] = None   # base64 encoded image
    attached_image_mime: Optional[str] = None  # e.g. "image/jpeg"
    doc_settings: Optional[Dict[str, Any]] = None  # current editor settings

class TranslateRequest(BaseModel):
    text: str
    target_language: str

class ShareLinkCreate(BaseModel):
    permission: str = "view"  # "view" or "edit"
    expires_hours: Optional[int] = None

class CommentCreate(BaseModel):
    content: str
    element_id: Optional[str] = None
    x: Optional[float] = None
    y: Optional[float] = None
    page_index: Optional[int] = 0

class CommentReply(BaseModel):
    content: str

class JudgeFeedbackCreate(BaseModel):
    session_id: Optional[str] = None
    message_index: int = 0
    feedback_type: str  # 'positive' or 'negative'
    message_content: str = ''

class EmailAuthRequest(BaseModel):
    email: str
    password: str
    name: Optional[str] = None

class ZetIdLoginRequest(BaseModel):
    access_token: str
    refresh_token: Optional[str] = None

class ZetIdConnectTokenRequest(BaseModel):
    id_token: str

class ZetIdConnectEmailRequest(BaseModel):
    email: str
    password: str

class ZetIdTokenRequest(BaseModel):
    grant_type: str
    code: Optional[str] = None
    client_id: str
    client_secret: str
    redirect_uri: Optional[str] = None

# ============ COLLABORATION WEBSOCKET MANAGER ============

class CollaborationManager:
    def __init__(self):
        self.active_rooms: Dict[str, Dict[str, dict]] = {}  # doc_id -> {user_id: {ws, user_info}}

    async def connect(self, doc_id: str, user_id: str, user_name: str, ws: WebSocket):
        await ws.accept()
        if doc_id not in self.active_rooms:
            self.active_rooms[doc_id] = {}
        color = f"hsl({hash(user_id) % 360}, 70%, 60%)"
        self.active_rooms[doc_id][user_id] = {
            "ws": ws, "name": user_name, "color": color,
            "cursor": None, "connected_at": datetime.now(timezone.utc).isoformat()
        }
        # Notify others
        await self.broadcast(doc_id, user_id, {
            "type": "user_joined",
            "user_id": user_id, "name": user_name, "color": color,
            "users": self._get_users(doc_id)
        })

    def disconnect(self, doc_id: str, user_id: str):
        if doc_id in self.active_rooms:
            self.active_rooms[doc_id].pop(user_id, None)
            if not self.active_rooms[doc_id]:
                del self.active_rooms[doc_id]

    async def broadcast(self, doc_id: str, sender_id: str, message: dict):
        if doc_id not in self.active_rooms:
            return
        dead = []
        for uid, info in self.active_rooms[doc_id].items():
            if uid == sender_id:
                continue
            try:
                await info["ws"].send_json(message)
            except Exception:
                dead.append(uid)
        for uid in dead:
            self.active_rooms[doc_id].pop(uid, None)

    def _get_users(self, doc_id: str):
        if doc_id not in self.active_rooms:
            return []
        return [{"user_id": uid, "name": info["name"], "color": info["color"]}
                for uid, info in self.active_rooms[doc_id].items()]

    def get_online_count(self, doc_id: str):
        return len(self.active_rooms.get(doc_id, {}))

collab_manager = CollaborationManager()

# ============ EMAIL HELPER ============

resend.api_key = os.environ.get("RESEND_API_KEY")
SENDER_EMAIL = os.environ.get("SENDER_EMAIL", "ZET Mindshare <info@zetstudiointl.com>")

# ── Lemon Squeezy ─────────────────────────────────────────────────────────────
LS_API_KEY = os.getenv("LEMONSQUEEZY_API_KEY", "")
LS_WEBHOOK_SECRET = os.getenv("LEMONSQUEEZY_WEBHOOK_SECRET", "")
LS_STORE_ID = os.getenv("LEMONSQUEEZY_STORE_ID", "342968")
LS_VARIANTS: Dict[str, Dict[str, str]] = {
    "plus":             {"monthly": os.getenv("LS_VARIANT_PLUS_MONTHLY", "1633716"), "yearly": os.getenv("LS_VARIANT_PLUS_YEARLY", "1633755")},
    "pro":              {"monthly": os.getenv("LS_VARIANT_PRO_MONTHLY", "1633736"),  "yearly": os.getenv("LS_VARIANT_PRO_YEARLY", "1633759")},
    "creative_station": {"monthly": os.getenv("LS_VARIANT_CS_MONTHLY", "1633753"),  "yearly": os.getenv("LS_VARIANT_CS_YEARLY", "1633854")},
}

async def send_email(to_email: str, subject: str, html_content: str) -> dict:
    """Send email using Resend API"""
    try:
        params = {
            "from": SENDER_EMAIL,
            "to": [to_email],
            "subject": subject,
            "html": html_content
        }
        result = await asyncio.to_thread(resend.Emails.send, params)
        logging.info(f"Email sent to {to_email}: {result}")
        return {"success": True, "email_id": result.get("id")}
    except Exception as e:
        logging.error(f"Failed to send email to {to_email}: {str(e)}")
        return {"success": False, "error": str(e)}

# ============ PLAN PRICING ============

PLAN_PRICES_USD = {
    "plus":             {"monthly": 9.99,  "yearly": 99.99},
    "pro":              {"monthly": 19.99, "yearly": 199.99},
    "creative_station": {"monthly": 49.00, "yearly": 490.00},
}

def compute_next_renewal(from_date: date, billing_cycle: str) -> date:
    """Return the next renewal date using calendar month/year arithmetic.
    relativedelta handles month-end edge cases (e.g. Jan 31 → Feb 28)."""
    delta = relativedelta(months=1) if billing_cycle == "monthly" else relativedelta(years=1)
    return from_date + delta

async def send_renewal_email(to_email: str, plan: str, renewal_date: date, amount: float, next_date: date):
    currency = "USD"
    html = f"""
    <div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;border-radius:12px;">
        <h2 style="color:#4ca8ad;margin-bottom:20px;">🔄 ZET Mindshare Aboneliğiniz Yenilendi</h2>
        <p style="font-size:16px;line-height:1.6;">Aboneliğiniz başarıyla yenilendi. Aşağıda fatura detaylarını bulabilirsiniz.</p>
        <div style="background:rgba(255,255,255,0.08);padding:16px;border-radius:8px;margin:20px 0;">
            <p style="margin:6px 0;color:#4ca8ad;"><strong>Plan:</strong> {plan.upper()}</p>
            <p style="margin:6px 0;color:#4ca8ad;"><strong>Yenileme Tarihi:</strong> {renewal_date.strftime('%d.%m.%Y')}</p>
            <p style="margin:6px 0;color:#4ca8ad;"><strong>Tahsil Edilen Tutar:</strong> {amount:.2f} {currency}</p>
            <p style="margin:6px 0;color:#4ca8ad;"><strong>Sonraki Yenileme:</strong> {next_date.strftime('%d.%m.%Y')}</p>
        </div>
        <a href="https://zetmindshare.com" style="display:inline-block;background:#4ca8ad;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;">ZET Mindshare'e Git</a>
        <p style="color:#888;font-size:12px;margin-top:20px;">Sorularınız için info@zetstudiointl.com adresine yazabilirsiniz.</p>
    </div>
    """
    await send_email(to_email, f"🔄 ZET Mindshare {plan.upper()} Planı Yenilendi", html)

async def send_renewal_reminder_email(to_email: str, plan: str, renewal_date: date, amount: float):
    currency = "USD"
    html = f"""
    <div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;border-radius:12px;">
        <h2 style="color:#f59e0b;margin-bottom:20px;">⏰ Abonelik Yenileme Hatırlatması</h2>
        <p style="font-size:16px;line-height:1.6;">ZET Mindshare aboneliğiniz <strong>3 gün içinde</strong> yenilenecektir.</p>
        <div style="background:rgba(255,255,255,0.08);padding:16px;border-radius:8px;margin:20px 0;">
            <p style="margin:6px 0;color:#f59e0b;"><strong>Plan:</strong> {plan.upper()}</p>
            <p style="margin:6px 0;color:#f59e0b;"><strong>Yenileme Tarihi:</strong> {renewal_date.strftime('%d.%m.%Y')}</p>
            <p style="margin:6px 0;color:#f59e0b;"><strong>Tahsil Edilecek Tutar:</strong> {amount:.2f} {currency}</p>
        </div>
        <p style="color:#ccc;font-size:14px;">Aboneliğinizi iptal etmek istiyorsanız yenileme tarihinden önce hesap ayarlarınızdan işlem yapabilirsiniz.</p>
        <a href="https://zetmindshare.com" style="display:inline-block;background:#4ca8ad;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;">Hesabımı Yönet</a>
    </div>
    """
    await send_email(to_email, f"⏰ ZET Mindshare {plan.upper()} Planı 3 Gün İçinde Yenilenecek", html)

# ============ SUBSCRIPTION HELPERS ============

def is_active_subscriber(user_data: dict) -> bool:
    """Kullanıcının aktif ücretli aboneliği var mı?
    Hem eski (string) hem yeni (dict) subscription formatını destekler."""
    if not user_data:
        return False
    sub = user_data.get("subscription")
    if not sub:
        return False
    # Yeni format: subscription bir dict, status alanı var
    if isinstance(sub, dict):
        plan = sub.get("plan", "free")
        status = sub.get("status", "inactive")
        return plan != "free" and status == "active"
    # Eski format: subscription bir string (plan adı)
    return sub != "free"

def get_plan_name(user_data: dict) -> str:
    """Her iki formattan plan adını döndür."""
    if not user_data:
        return "free"
    sub = user_data.get("subscription", "free")
    if isinstance(sub, dict):
        return sub.get("plan", "free")
    return sub if sub else "free"

# ============ AUTH HELPERS ============

import bcrypt as _bcrypt

def hash_password(password: str) -> str:
    return _bcrypt.hashpw(password.encode(), _bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    try:
        return _bcrypt.checkpw(password.encode(), hashed.encode())
    except Exception:
        # Backward compat: old SHA-256 hashes
        import hashlib
        return hashlib.sha256(password.encode()).hexdigest() == hashed

import jwt as pyjwt

ZET_ID_JWT_SECRET = os.environ.get("ZET_ID_JWT_SECRET", "zet-id-dev-secret-change-in-production")
ZET_ID_JWT_ALG = "HS256"
ZET_ID_ALPHABET = "ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"

async def generate_unique_zet_id() -> str:
    """Benzersiz, okunabilir ZET ID üretir: 'ZET-' + 8 alfanumerik karakter."""
    while True:
        code = "ZET-" + "".join(secrets.choice(ZET_ID_ALPHABET) for _ in range(8))
        if not await db.users.find_one({"zet_id": code}, {"_id": 0, "zet_id": 1}):
            return code

async def ensure_zet_id(user_doc: dict) -> tuple:
    """Kullanıcının ZET ID'si yoksa üretip kaydeder — mevcut kullanıcılar için katkılı (additive), eşzamanlılığa dayanıklı migrasyon.
    Sadece eksik olan 'zet_id' alanını $set eder; başka hiçbir alana dokunmaz."""
    existing = user_doc.get("zet_id")
    if existing:
        return existing, False
    for _ in range(5):
        new_id = await generate_unique_zet_id()
        result = await db.users.update_one(
            {"user_id": user_doc["user_id"], "zet_id": {"$exists": False}},
            {"$set": {"zet_id": new_id}}
        )
        if result.modified_count == 1:
            return new_id, True
        # Eşzamanlı başka bir istek arada ZET ID atamış olabilir — mevcut değeri kullan
        refreshed = await db.users.find_one({"user_id": user_doc["user_id"]}, {"_id": 0, "zet_id": 1})
        if refreshed and refreshed.get("zet_id"):
            return refreshed["zet_id"], False
    raise HTTPException(status_code=500, detail="ZET ID oluşturulamadı")

def compute_zet_id_device_hash(request: Request) -> str:
    ua = request.headers.get("User-Agent", "")
    lang = request.headers.get("Accept-Language", "")
    return hafizz.compute_device_fingerprint(ua, lang)

def create_zet_id_access_token(user_id: str, zet_id: str, device_hash: str, jti: str) -> str:
    now = datetime.now(timezone.utc)
    payload = {
        "sub": user_id,
        "zet_id": zet_id,
        "device_hash": device_hash,
        "jti": jti,
        "iat": now,
        "exp": now + timedelta(days=30),
    }
    return pyjwt.encode(payload, ZET_ID_JWT_SECRET, algorithm=ZET_ID_JWT_ALG)

def decode_zet_id_access_token(token: str) -> dict:
    return pyjwt.decode(token, ZET_ID_JWT_SECRET, algorithms=[ZET_ID_JWT_ALG])

async def issue_zet_id_session(user_id: str, zet_id: str, device_hash: str) -> dict:
    """Bu cihaz için ZET ID erişim (JWT, 30 gün) + yenileme (opak, 90 gün) token çifti üretir."""
    jti = uuid.uuid4().hex
    access_token = create_zet_id_access_token(user_id, zet_id, device_hash, jti)
    refresh_token = f"zrt_{secrets.token_hex(32)}"
    now = datetime.now(timezone.utc)
    await db.zet_id_sessions.insert_one({
        "jti": jti,
        "user_id": user_id,
        "zet_id": zet_id,
        "device_hash": device_hash,
        "refresh_token": refresh_token,
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(days=90)).isoformat(),
        "revoked": False,
    })
    return {"access_token": access_token, "refresh_token": refresh_token, "expires_in": 30 * 24 * 60 * 60}

async def get_current_user(request: Request) -> User:
    session_token = request.cookies.get("session_token")
    if not session_token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            session_token = auth_header.split(" ")[1]

    if not session_token:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        session = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0})
    except Exception as e:
        logging.error(f"DB error in get_current_user (session lookup): {e}")
        raise HTTPException(status_code=503, detail="Service temporarily unavailable")
    if not session:
        raise HTTPException(status_code=401, detail="Invalid session")

    expires_at = session.get("expires_at")
    if not expires_at:
        raise HTTPException(status_code=401, detail="Invalid session")
    try:
        if isinstance(expires_at, str):
            expires_at = datetime.fromisoformat(expires_at)
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        if expires_at < datetime.now(timezone.utc):
            raise HTTPException(status_code=401, detail="Session expired")
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"expires_at parse error: {e!r} value={expires_at!r}")
        raise HTTPException(status_code=401, detail="Invalid session")

    try:
        raw = await db.users.find_one({"user_id": session["user_id"]})
    except Exception as e:
        logging.error(f"DB error in get_current_user (user lookup): {e}")
        raise HTTPException(status_code=503, detail="Service temporarily unavailable")
    if not raw:
        raise HTTPException(status_code=401, detail="User not found")

    # DB verisini Pydantic'e geçmeden önce tam normalize et
    user = sanitize_user_doc(raw)

    try:
        return User(**user)
    except Exception as e:
        logging.error(f"User validation error user_id={session['user_id']}: {e}")
        return User(user_id=user["user_id"], email=user.get("email", ""))

async def get_current_user_raw(request: Request) -> dict:
    """Pydantic olmadan sanitize edilmiş kullanıcı dict döner; subscription tip çakışması yaşanan endpoint'lerde kullan."""
    session_token = request.cookies.get("session_token")
    if not session_token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            session_token = auth_header.split(" ")[1]
    if not session_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    session = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=401, detail="Invalid session")
    expires_at = session.get("expires_at")
    if not expires_at:
        raise HTTPException(status_code=401, detail="Invalid session")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Session expired")
    raw = await db.users.find_one({"user_id": session["user_id"]})
    if not raw:
        raise HTTPException(status_code=401, detail="User not found")
    return sanitize_user_doc(raw)

# ============ AUTH ROUTES ============

@api_router.get("/auth/google")
async def google_auth_init():
    """Redirect user to Google OAuth2 consent screen (PKCE olmadan standart flow)."""
    from requests_oauthlib import OAuth2Session
    from fastapi.responses import RedirectResponse
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    if not client_id:
        raise HTTPException(status_code=500, detail="GOOGLE_CLIENT_ID .env dosyasına eklenmeli")
    backend_url = os.getenv("REACT_APP_BACKEND_URL", "http://localhost:8001")
    redirect_uri = f"{backend_url}/api/auth/google/callback"
    scopes = [
        "openid",
        "https://www.googleapis.com/auth/userinfo.email",
        "https://www.googleapis.com/auth/userinfo.profile"
    ]
    oauth = OAuth2Session(client_id=client_id, redirect_uri=redirect_uri, scope=scopes)
    auth_url, _ = oauth.authorization_url(
        "https://accounts.google.com/o/oauth2/auth",
        access_type="offline",
        prompt="consent"
    )
    return RedirectResponse(auth_url)


@api_router.get("/auth/google/callback")
async def google_auth_callback(request: Request, response: Response, code: str = Query(None), state: str = Query(None), error: str = Query(None)):
    """Handle Google OAuth2 callback — token exchange, session oluştur, frontend'e yönlendir."""
    from requests_oauthlib import OAuth2Session
    from fastapi.responses import RedirectResponse, HTMLResponse
    import requests as pyrequests

    # Allow insecure transport only in non-production environments
    if not os.getenv("REACT_APP_BACKEND_URL", "").startswith("https"):
        os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")

    if error or not code:
        logging.error(f"Google OAuth error param: {error}")
        return RedirectResponse(f"{frontend_url}/?error=google_auth_failed")

    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    backend_url = os.getenv("REACT_APP_BACKEND_URL", "http://localhost:8001")
    redirect_uri = f"{backend_url}/api/auth/google/callback"

    try:
        oauth = OAuth2Session(client_id=client_id, redirect_uri=redirect_uri, state=state)
        token = oauth.fetch_token(
            "https://oauth2.googleapis.com/token",
            client_secret=client_secret,
            code=code
        )
        userinfo_resp = pyrequests.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {token['access_token']}"}
        )
        userinfo = userinfo_resp.json()

        email = userinfo.get("email")
        name = userinfo.get("name", email.split("@")[0] if email else "User")
        picture = userinfo.get("picture")

        if not email:
            return RedirectResponse(f"{frontend_url}/?error=no_email")

        # Admin panel redirect flow
        if state == "admin_console":
            if email == CEO_EMAIL:
                ceo_token = f"ceo_{uuid.uuid4().hex}"
                await db.ceo_tokens.insert_one({
                    "token": ceo_token,
                    "email": email,
                    "expires_at": (datetime.now(timezone.utc) + timedelta(minutes=5)).isoformat()
                })
                return RedirectResponse(f"{frontend_url}/auth-callback#ceo_token={ceo_token}")
            elif email in ADMIN_EMAILS:
                admin_token = f"adm_{uuid.uuid4().hex}"
                await db.ceo_tokens.insert_one({
                    "token": admin_token,
                    "email": email,
                    "role": "admin",
                    "expires_at": (datetime.now(timezone.utc) + timedelta(minutes=5)).isoformat()
                })
                return RedirectResponse(f"{frontend_url}/auth-callback#admin_token={admin_token}")
            else:
                return RedirectResponse(f"{frontend_url}/auth-callback#ceo_error=access_denied")

        # KRİTİK: Email'i daima küçük harfe çevir (Google bazen farklı case dönebilir)
        email = email.lower().strip()

        # Adım 1: Email'e göre mevcut kullanıcıyı ara
        existing_user = await db.users.find_one({"email": email}, {"_id": 0})

        if existing_user:
            # Mevcut kullanıcı — user_id'yi ASLA değiştirme
            user_id = existing_user["user_id"]
            logging.info(f"Google OAuth: mevcut kullanıcı bulundu email={email} user_id={user_id}")

            oauth_update = {}
            if not existing_user.get("name_custom"):
                oauth_update["name"] = name
            if not existing_user.get("picture_custom"):
                oauth_update["picture"] = picture
            if not existing_user.get("username"):
                oauth_update["username"] = await generate_unique_username(email.split("@")[0])
            if email == CEO_EMAIL and not existing_user.get("verified_type"):
                oauth_update["verified_type"] = "red"
            if oauth_update:
                await db.users.update_one({"user_id": user_id}, {"$set": oauth_update})
        else:
            # Yeni kullanıcı — yeni user_id üret
            user_id = f"user_{secrets.token_hex(6)}"
            username = await generate_unique_username(email.split("@")[0])
            verified_type = "red" if email == CEO_EMAIL else None
            trial_end = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
            await db.users.insert_one({
                "user_id": user_id,
                "email": email,
                "name": name,
                "picture": picture,
                "username": username,
                "display_name": name,
                "bio": "",
                "verified_type": verified_type,
                "followers": [],
                "following": [],
                "followers_count": 0,
                "following_count": 0,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "needs_onboarding": True,
                "subscription": {"plan": "pro", "status": "active", "trial": True, "trial_end": trial_end},
            })
            logging.info(f"Google OAuth: yeni kullanıcı oluşturuldu email={email} user_id={user_id}")

        # Adım 2: Session oluşturmadan önce user_id'yi DB'den doğrula
        # (yarış durumu veya DB tutarsızlığı ihtimaline karşı)
        verified = await db.users.find_one({"user_id": user_id}, {"_id": 0, "user_id": 1})
        if not verified:
            # user_id ile bulunamazsa email'e göre fallback
            fallback = await db.users.find_one({"email": email}, {"_id": 0, "user_id": 1})
            if fallback:
                user_id = fallback["user_id"]
                logging.warning(f"Google OAuth: user_id fallback email={email} → {user_id}")
            else:
                logging.error(f"Google OAuth: kullanıcı oluşturulamadı email={email}")
                return RedirectResponse(f"{frontend_url}/?error=user_creation_failed")

        # Adım 3: Session oluştur
        session_token = f"st_{secrets.token_hex(16)}"
        expires_at = datetime.now(timezone.utc) + timedelta(days=7)
        await db.user_sessions.insert_one({
            "user_id": user_id,
            "session_token": session_token,
            "expires_at": expires_at.isoformat(),
            "created_at": datetime.now(timezone.utc).isoformat()
        })

        # Hafız: device fingerprint + ülke (background)
        login_ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "").split(",")[0].strip()
        ua = request.headers.get("User-Agent", "")
        lang = request.headers.get("Accept-Language", "")
        fp = hafizz.compute_device_fingerprint(ua, lang)
        asyncio.create_task(hafizz.record_device_login(db, user_id, fp, login_ip))
        asyncio.create_task(_hafizz_post_login(db, user_id, login_ip))

        return RedirectResponse(f"{frontend_url}/auth-callback#token={session_token}")

    except Exception as e:
        logging.error(f"Google OAuth callback error: {e}", exc_info=True)
        return RedirectResponse(f"{frontend_url}/?error=google_auth_failed")

@api_router.get("/auth/admin-google")
async def admin_google_auth():
    """Admin panel login — login sayfasıyla aynı redirect akışını kullanır, state=admin_console."""
    from requests_oauthlib import OAuth2Session
    from fastapi.responses import RedirectResponse
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    backend_url = os.getenv("REACT_APP_BACKEND_URL", "http://localhost:8001")
    redirect_uri = f"{backend_url}/api/auth/google/callback"
    scopes = ["openid", "https://www.googleapis.com/auth/userinfo.email", "https://www.googleapis.com/auth/userinfo.profile"]
    oauth = OAuth2Session(client_id=client_id, redirect_uri=redirect_uri, scope=scopes)
    auth_url, _ = oauth.authorization_url(
        "https://accounts.google.com/o/oauth2/auth",
        access_type="offline",
        prompt="select_account",
        state="admin_console"
    )
    return RedirectResponse(auth_url)

@api_router.get("/auth/admin-verify")
async def admin_verify_token(token: str = Query(...)):
    """CEO token doğrulama — tek kullanımlık, 5 dakika geçerliliği var."""
    now = datetime.now(timezone.utc)
    record = await db.ceo_tokens.find_one({"token": token})
    if not record:
        raise HTTPException(status_code=403, detail="Geçersiz token")
    exp = datetime.fromisoformat(record["expires_at"].replace("Z", "+00:00"))
    if now > exp:
        await db.ceo_tokens.delete_one({"token": token})
        raise HTTPException(status_code=403, detail="Token süresi dolmuş")
    await db.ceo_tokens.delete_one({"token": token})  # tek kullanımlık
    return {"success": True, "email": record["email"]}

class AdminPinRequest(BaseModel):
    pin: str

@api_router.post("/auth/admin-verify-pin")
async def admin_verify_pin(body: AdminPinRequest):
    """İkinci faktör — PIN doğrulama. CEO_PIN env var ile karşılaştırılır."""
    ceo_pin = os.environ.get("CEO_PIN", "")
    if not ceo_pin:
        raise HTTPException(status_code=503, detail="PIN sistemi yapılandırılmamış.")
    if body.pin != ceo_pin:
        raise HTTPException(status_code=403, detail="Yanlış PIN.")
    return {"success": True}

@api_router.get("/auth/admin-verify-admin")
async def admin_verify_admin_token(token: str = Query(...)):
    """Admin token doğrulama — tek kullanımlık, 5 dakika geçerliliği var."""
    now = datetime.now(timezone.utc)
    record = await db.ceo_tokens.find_one({"token": token})
    if not record:
        raise HTTPException(status_code=403, detail="Geçersiz token")
    if record.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Bu token admin yetkisi içermiyor")
    exp = datetime.fromisoformat(record["expires_at"].replace("Z", "+00:00"))
    if now > exp:
        await db.ceo_tokens.delete_one({"token": token})
        raise HTTPException(status_code=403, detail="Token süresi dolmuş")
    await db.ceo_tokens.delete_one({"token": token})  # tek kullanımlık
    return {"success": True, "email": record["email"]}

@api_router.get("/auth/admin-console")
async def admin_console_auth():
    """Eski popup akışı — geriye dönük uyumluluk için tutuldu, artık kullanılmıyor."""
    return RedirectResponse(f"/api/auth/admin-google")

@api_router.get("/auth/admin-console/callback")
async def admin_console_callback(request: Request, code: str = Query(None), state: str = Query(None), error: str = Query(None)):
    """Admin console OAuth callback — sonucu popup'a postMessage ile iletir."""
    from fastapi.responses import HTMLResponse
    import requests as pyrequests
    from requests_oauthlib import OAuth2Session
    if not os.getenv("REACT_APP_BACKEND_URL", "").startswith("https"):
        os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"
    if error or not code:
        return HTMLResponse('<script>window.opener&&window.opener.postMessage({ceoVerified:false},"*");window.close();</script>')
    try:
        client_id = os.getenv("GOOGLE_CLIENT_ID")
        client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
        backend_url = os.getenv("REACT_APP_BACKEND_URL", "http://localhost:8001")
        redirect_uri = f"{backend_url}/api/auth/admin-console/callback"
        oauth = OAuth2Session(client_id=client_id, redirect_uri=redirect_uri, state=state)
        token = oauth.fetch_token("https://oauth2.googleapis.com/token", client_secret=client_secret, code=code)
        userinfo = pyrequests.get("https://www.googleapis.com/oauth2/v2/userinfo", headers={"Authorization": f"Bearer {token['access_token']}"}).json()
        email = userinfo.get("email", "")
        verified = "true" if email == CEO_EMAIL else "false"
        return HTMLResponse(f'<script>window.opener&&window.opener.postMessage({{ceoVerified:{verified},email:{repr(email)}}},"*");window.close();</script>')
    except Exception as e:
        logging.error(f"Admin console OAuth error: {e}")
        return HTMLResponse('<script>window.opener&&window.opener.postMessage({ceoVerified:false},"*");window.close();</script>')

# ============ ADMIN ROUTES ============

def is_privileged(email: str) -> bool:
    """CEO veya admin mail kontrolü."""
    return email == CEO_EMAIL or email in ADMIN_EMAILS

async def generate_unique_username(base: str) -> str:
    """E-posta prefix'inden benzersiz username üret."""
    import re
    base = re.sub(r'[^a-z0-9_]', '_', base.lower())[:16].strip('_') or 'user'
    candidate = base
    for _ in range(20):
        exists = await db.users.find_one({"username": candidate})
        if not exists:
            return candidate
        candidate = f"{base}_{uuid.uuid4().hex[:4]}"
    return f"user_{uuid.uuid4().hex[:8]}"

@api_router.post("/admin/add-credits")
async def admin_add_credits(amount: int = Body(..., embed=True), user: User = Depends(get_current_user)):
    if not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"bonus_credits": amount}})
    data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "bonus_credits": 1})
    return {"added": amount, "total": data.get("bonus_credits", 0)}

@api_router.post("/admin/add-sp")
async def admin_add_sp(amount: int = Body(..., embed=True), user: User = Depends(get_current_user)):
    if not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"quest_xp": amount}})
    data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "quest_xp": 1})
    return {"added": amount, "total": data.get("quest_xp", 0)}

@api_router.post("/admin/free-subscription")
async def admin_free_subscription(user: User = Depends(get_current_user)):
    """Admin modundaki kullanıcı için ücretsiz Pro abonelik aktive eder."""
    if not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    now = datetime.now(timezone.utc)
    expires = (now + timedelta(days=36500)).isoformat()  # ~100 yıl
    await db.users.update_one({"user_id": user.user_id}, {"$set": {
        "subscription_status": "active",
        "subscription_plan": "pro",
        "subscription_expires": expires,
        "subscription_source": "admin_grant"
    }})
    return {"success": True, "plan": "pro", "expires": expires}

@api_router.get("/admin/check-documents")
async def admin_check_documents(user: User = Depends(get_current_user)):
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403)
    docs = await db.documents.find({}, {"_id": 0, "doc_id": 1, "title": 1, "user_id": 1, "updated_at": 1, "pages": 1}).to_list(1000)
    return [{"doc_id": d["doc_id"], "title": d.get("title", ""), "pages_count": len(d.get("pages", [])), "has_content": any(len(p.get("elements", [])) > 0 for p in d.get("pages", [])), "updated_at": d.get("updated_at", "")} for d in docs]

@api_router.get("/admin/list-users")
async def admin_list_users(user: User = Depends(get_current_user)):
    """Privileged hesaplara tüm kayıtlı kullanıcıları döner."""
    if not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    users = await db.users.find({}, {
        "_id": 0, "email": 1, "name": 1, "display_name": 1, "username": 1,
        "created_at": 1, "verified_type": 1, "subscription": 1, "user_id": 1,
    }).to_list(length=5000)
    return {"users": users}

class AdminUserUpdate(BaseModel):
    verified_type: Optional[str] = None   # "red", "gold", "blue", "" = kaldır
    subscription_plan: Optional[str] = None  # "free", "plus", "pro", "creative_station", "entertainment_pocket"

@api_router.patch("/admin/users/{user_id}")
async def admin_update_user(user_id: str, body: AdminUserUpdate, user: User = Depends(get_current_user)):
    """CEO: kullanıcının verified rozeti veya aboneliğini güncelle."""
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403, detail="Sadece CEO bu işlemi yapabilir.")
    update: dict = {}
    if body.verified_type is not None:
        update["verified_type"] = body.verified_type if body.verified_type else None
    if body.subscription_plan is not None:
        if body.subscription_plan == "free":
            update["subscription"] = "free"
        else:
            update["subscription"] = {"plan": body.subscription_plan, "status": "active", "assigned_by": "CEO"}
    if not update:
        return {"success": True}
    result = await db.users.update_one({"user_id": user_id}, {"$set": update})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    return {"success": True}

# ============ SEZON SİSTEMİ ============

_SEASON_RANK_XP = [
    (15000, "endless"),
    (5000,  "emerald"),
    (1500,  "diamond"),
    (500,   "gold"),
    (100,   "silver"),
    (0,     "iron"),
]
_SEASON_RANK_REWARDS = {
    "iron":    {"credits": 30,   "xp": 50},
    "silver":  {"credits": 200,  "xp": 400},
    "gold":    {"credits": 500,  "xp": 1000},
    "diamond": {"credits": 800,  "xp": 1600},
    "emerald": {"credits": 1000, "xp": 2400},
    "endless": {"credits": 2000, "xp": 3000},
}

def _xp_to_rank(xp: int) -> str:
    for threshold, name in _SEASON_RANK_XP:
        if xp >= threshold:
            return name
    return "iron"

# Saat bazlı rank — sezon dağıtımında kullanılır
_SEASON_RANK_HOURS = [
    (200, "endless"),
    (90,  "emerald"),
    (60,  "diamond"),
    (25,  "gold"),
    (10,  "silver"),
    (0,   "iron"),
]

def _hours_to_rank(seconds: int) -> str:
    hours = (seconds or 0) / 3600
    for threshold, name in _SEASON_RANK_HOURS:
        if hours >= threshold:
            return name
    return "iron"

async def distribute_season_rewards_and_reset():
    """Sezon ödüllerini dağıt, tüm kullanıcı rankını sıfırla."""
    # Atomik: 'active' → 'ending' geçişi; iki eş zamanlı çağrı olursa sadece biri devam eder
    claimed = await db.seasons.find_one_and_update(
        {"status": "active"},
        {"$set": {"status": "ending"}},
    )
    if not claimed:
        return 0
    season_id    = claimed.get("season_id", "")
    season_start = claimed.get("start_date", "")
    season_end   = claimed.get("end_date", "")
    try:
        s_dt = datetime.strptime(season_start, "%Y-%m-%d")
        e_dt = datetime.strptime(season_end,   "%Y-%m-%d")
        duration_days = max(1, (e_dt - s_dt).days)
    except Exception:
        duration_days = 0

    users = await db.users.find(
        {}, {"_id": 0, "user_id": 1, "mindshare_xp": 1, "active_time_seconds": 1}
    ).to_list(length=200000)
    count = 0
    for u in users:
        active_secs = u.get("active_time_seconds", 0) or 0
        rank   = _hours_to_rank(active_secs)
        reward = _SEASON_RANK_REWARDS[rank]
        # Sezon sonuç kaydı — popup için
        await db.season_results.insert_one({
            "user_id":             u["user_id"],
            "season_id":           season_id,
            "final_rank":          rank,
            "credits_earned":      reward["credits"],
            "duration_days":       duration_days,
            "active_time_seconds": active_secs,
            "season_start":        season_start,
            "season_end":          season_end,
            "shown":               False,
            "created_at":          datetime.now(timezone.utc).isoformat(),
        })
        await db.users.update_one(
            {"user_id": u["user_id"]},
            {
                "$inc": {"bonus_credits": reward["credits"]},
                "$set": {"mindshare_xp": 0, "mindshare_rank": "iron", "active_time_seconds": 0},
            },
        )
        count += 1
    await db.seasons.update_one(
        {"status": "ending"},
        {"$set": {"status": "ended", "ended_at": datetime.now(timezone.utc).isoformat()}},
    )
    logging.info(f"Season ended — distributed rewards to {count} users, ranks reset.")
    return count

@api_router.get("/season")
async def get_season():
    season = await db.seasons.find_one({"status": "active"}, {"_id": 0})
    if not season:
        return {"active": False}
    try:
        end = datetime.strptime(season["end_date"], "%Y-%m-%d")
        days_remaining = max(0, (end - datetime.utcnow()).days)
    except Exception:
        days_remaining = None
    # Cloud Run APScheduler'ı kaçırırsa burada yakala
    today = datetime.utcnow().strftime("%Y-%m-%d")
    if season.get("end_date", "9999-99-99") <= today:
        await distribute_season_rewards_and_reset()
        return {"active": False}
    return {
        "active": True,
        "start_date": season.get("start_date"),
        "end_date": season["end_date"],
        "days_remaining": days_remaining,
    }

@api_router.get("/season/time/{date_range}")
async def set_season_time(date_range: str, pin: str = Query("")):
    """Sezon tarihini ayarla. Format: DD.MM.YYYY-DD.MM.YYYY veya DD.MM.YYYY (sadece bitiş).
    Örnek: GET /api/season/time/12.09.2026-09.03.2027?pin=CEO_PIN"""
    ceo_pin = os.getenv("CEO_PIN", "")
    if not ceo_pin or pin != ceo_pin:
        raise HTTPException(status_code=403, detail="Geçersiz PIN")
    parts = date_range.split("-")
    try:
        if len(parts) == 2:
            start_dt = datetime.strptime(parts[0].strip(), "%d.%m.%Y")
            end_dt   = datetime.strptime(parts[1].strip(), "%d.%m.%Y")
        else:
            start_dt = datetime.utcnow()
            end_dt   = datetime.strptime(parts[0].strip(), "%d.%m.%Y")
    except ValueError:
        raise HTTPException(status_code=400, detail="Tarih formatı yanlış. Örnek: 12.09.2026-09.03.2027")
    if end_dt < start_dt:
        raise HTTPException(status_code=400, detail="Bitiş tarihi başlangıçtan önce olamaz")
    await db.seasons.update_one(
        {"status": "active"},
        {"$set": {"status": "replaced", "replaced_at": datetime.now(timezone.utc).isoformat()}},
    )
    new_season = {
        "season_id": f"s_{uuid.uuid4().hex[:8]}",
        "start_date": start_dt.strftime("%Y-%m-%d"),
        "end_date":   end_dt.strftime("%Y-%m-%d"),
        "status": "active",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.seasons.insert_one(new_season)
    new_season.pop("_id", None)
    return {"success": True, "season": new_season}

@api_router.get("/season/my-result")
async def get_my_season_result(user: User = Depends(get_current_user)):
    """Kullanıcının gösterilmemiş son sezon sonucunu döner."""
    result = await db.season_results.find_one(
        {"user_id": user.user_id, "shown": False},
        {"_id": 0},
        sort=[("created_at", -1)],
    )
    if not result:
        return {"has_result": False}
    return {"has_result": True, **result}

@api_router.post("/season/my-result/dismiss")
async def dismiss_season_result(user: User = Depends(get_current_user)):
    """Sezon sonuç popup'ını gösterildi olarak işaretle."""
    await db.season_results.update_many(
        {"user_id": user.user_id, "shown": False},
        {"$set": {"shown": True}},
    )
    return {"ok": True}

@api_router.post("/season/distribute")
async def manual_distribute_season(pin: str = Body(..., embed=True)):
    """Manuel sezon ödülü dağıtımı — sadece CEO kullanır."""
    ceo_pin = os.getenv("CEO_PIN", "")
    if not ceo_pin or pin != ceo_pin:
        raise HTTPException(status_code=403, detail="Geçersiz PIN")
    count = await distribute_season_rewards_and_reset()
    return {"success": True, "rewarded_users": count}

# ============ ENVANTER / KASA ============

_CASE_REWARDS = [
    {"type": "zp",     "amount": 30,  "rarity": "common", "chance": 20.0},
    {"type": "zp",     "amount": 50,  "rarity": "common", "chance": 15.0},
    {"type": "zp",     "amount": 100, "rarity": "nadir",  "chance": 13.0},
    {"type": "zp",     "amount": 200, "rarity": "nadir",  "chance": 8.0},
    {"type": "zp",     "amount": 330, "rarity": "epik",   "chance": 3.0},
    {"type": "zp",     "amount": 800, "rarity": "epik",   "chance": 0.8},
    {"type": "credit", "amount": 10,  "rarity": "common", "chance": 22.0},
    {"type": "credit", "amount": 20,  "rarity": "nadir",  "chance": 9.0},
    {"type": "credit", "amount": 50,  "rarity": "epik",   "chance": 4.0},
    {"type": "credit", "amount": 400, "rarity": "lore",   "chance": 0.08},
]
_CASE_TOTAL_CHANCE = sum(r["chance"] for r in _CASE_REWARDS)

def _roll_case_reward() -> dict:
    rng = random.random() * _CASE_TOTAL_CHANCE
    cumulative = 0.0
    for reward in _CASE_REWARDS:
        cumulative += reward["chance"]
        if rng <= cumulative:
            return dict(reward)
    return dict(_CASE_REWARDS[-1])

@api_router.get("/inventory")
async def get_inventory(user: User = Depends(get_current_user)):
    now = datetime.now(timezone.utc)
    this_month = now.strftime("%Y-%m")
    doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "inventory": 1, "last_daily_case": 1})
    user_data = await db.users.find_one({"user_id": user.user_id})
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    monthly_doc = await db.chest_usage.find_one({"user_id": user.user_id, "month": this_month}) or {}
    cases_this_month = monthly_doc.get("count", 0)
    return {
        "cases": doc.get("inventory") or [],
        "last_daily_case": doc.get("last_daily_case"),
        "chest_daily_chance": limits.get('chest_daily_chance', 20),
        "chest_monthly_max": limits.get('chest_monthly_max', 3),
        "chest_monthly_used": cases_this_month,
    }

@api_router.post("/inventory/claim-daily")
async def claim_daily_case(user: User = Depends(get_current_user)):
    now = datetime.now(timezone.utc)
    today = now.strftime("%Y-%m-%d")
    this_month = now.strftime("%Y-%m")

    user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "last_daily_case": 1})
    if user_data and user_data.get("last_daily_case") == today:
        return {"claimed": False, "reason": "already_claimed"}

    plan = get_plan_name(await db.users.find_one({"user_id": user.user_id}))
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    chest_chance = limits.get('chest_daily_chance', 20)
    chest_monthly_max = limits.get('chest_monthly_max', 3)

    # Aylık limit kontrolü
    monthly_doc = await db.chest_usage.find_one({"user_id": user.user_id, "month": this_month}) or {}
    cases_this_month = monthly_doc.get("count", 0)
    if cases_this_month >= chest_monthly_max:
        await db.users.update_one({"user_id": user.user_id}, {"$set": {"last_daily_case": today}})
        return {"claimed": False, "reason": "monthly_limit", "monthly_limit": chest_monthly_max, "used": cases_this_month}

    # Şans turu
    won = chest_chance >= 100 or (random.random() * 100 < chest_chance)
    await db.users.update_one({"user_id": user.user_id}, {"$set": {"last_daily_case": today}})

    if not won:
        return {"claimed": False, "reason": "no_luck", "chance": chest_chance}

    case_obj = {"id": f"case_{uuid.uuid4().hex[:12]}", "type": "daily_case", "acquired_at": now.isoformat()}
    await db.users.update_one({"user_id": user.user_id}, {"$push": {"inventory": case_obj}})
    await db.chest_usage.update_one(
        {"user_id": user.user_id, "month": this_month},
        {"$inc": {"count": 1}},
        upsert=True
    )
    return {"claimed": True, "case": case_obj, "chance": chest_chance, "monthly_remaining": chest_monthly_max - cases_this_month - 1}

@api_router.post("/inventory/open-case")
async def open_case(case_id: str = Body(..., embed=True), user: User = Depends(get_current_user)):
    doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "inventory": 1})
    inventory = doc.get("inventory") or []
    if not any(c.get("id") == case_id for c in inventory):
        raise HTTPException(status_code=404, detail="Kasa bulunamadı")
    reward = _roll_case_reward()
    await db.users.update_one({"user_id": user.user_id}, {"$pull": {"inventory": {"id": case_id}}})
    if reward["type"] == "zp":
        await db.users.update_one({"user_id": user.user_id}, {"$inc": {"mindshare_xp": reward["amount"]}})
    elif reward["type"] == "credit":
        await db.users.update_one({"user_id": user.user_id}, {"$inc": {"bonus_credits": reward["amount"]}})
    return {"reward": reward}

@api_router.post("/admin/give-test-cases")
async def give_test_cases(count: int = Body(30, embed=True), user: User = Depends(get_current_user)):
    if not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    now = datetime.now(timezone.utc).isoformat()
    cases = [{"id": f"case_{uuid.uuid4().hex[:12]}", "type": "daily_case", "acquired_at": now} for _ in range(count)]
    await db.users.update_one({"user_id": user.user_id}, {"$push": {"inventory": {"$each": cases}}})
    return {"added": count}

# ============ SOCIAL — USERS / FOLLOW ============

class UsernameUpdate(BaseModel):
    username: str
    display_name: Optional[str] = None
    bio: Optional[str] = None
    complete_onboarding: bool = False

class SetVerifiedRequest(BaseModel):
    username: str
    verified_type: Optional[str] = None  # "red", "gold", "blue" veya None (kaldır)

async def enrich_post(post: dict, viewer_id: str, viewer_following: list = None) -> dict:
    """Post'a yazar bilgisi, liked_by_me ve viewer_follows_author ekle."""
    author = await db.users.find_one({"user_id": post.get("author_id")}, {"_id": 0, "username": 1, "display_name": 1, "name": 1, "picture": 1, "verified_type": 1})
    if author:
        author = sanitize_user_doc(author)
    post["author"] = author or {}
    post["liked_by_me"] = viewer_id in (post.get("likes") or [])
    if viewer_following is not None:
        post["viewer_follows_author"] = post.get("author_id") in viewer_following
    post.pop("likes", None)
    return post

@api_router.get("/fonts")
async def get_fonts():
    """Google Fonts listesini getir — 1000+ font."""
    import aiohttp
    api_key = os.getenv("GOOGLE_FONTS_API_KEY")
    if not api_key:
        return []
    try:
        async with aiohttp.ClientSession() as session:
            url = f"https://www.googleapis.com/webfonts/v1/webfonts?key={api_key}&sort=popularity"
            async with session.get(url) as resp:
                data = await resp.json()
                fonts = data.get("items", [])
                return [{"family": f["family"], "category": f.get("category", "sans-serif"), "variants": f.get("variants", ["regular"])} for f in fonts]
    except Exception as e:
        logging.error(f"Google Fonts API error: {e}")
        return []

@api_router.get("/users/preferences")
async def get_preferences(user: User = Depends(get_current_user)):
    doc = await db.users.find_one({"user_id": user.user_id}, {"preferences": 1, "_id": 0})
    return {"preferences": (doc or {}).get("preferences", {})}

class PreferencesUpdate(BaseModel):
    preferences: Dict[str, Any]

@api_router.put("/users/preferences")
async def update_preferences(body: PreferencesUpdate, user: User = Depends(get_current_user)):
    try:
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": {f"preferences.{k}": v for k, v in body.preferences.items()}}
        )
    except Exception as e:
        logging.error(f"update_preferences DB error: {e}")
        raise HTTPException(status_code=503, detail="Service temporarily unavailable")
    return {"ok": True}

@api_router.get("/users/me")
async def get_me(user: User = Depends(get_current_user)):
    data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "followers": 0, "following": 0})
    return data or {}

@api_router.get("/users/{username}")
async def get_user_profile(username: str, u: dict = Depends(get_current_user_raw)):
    uid = u["user_id"]
    target = await db.users.find_one({"username": username}, {"_id": 0, "email": 0, "followers": 0, "following": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    target["is_following"] = await db.users.find_one({"user_id": uid, "following": target["user_id"]}) is not None
    return target

@api_router.patch("/users/me")
async def update_profile(body: UsernameUpdate, user: User = Depends(get_current_user)):
    import re
    update = {}
    if body.username:
        uname = body.username.lower().strip()
        if not re.match(r'^[a-z0-9_]{3,20}$', uname):
            raise HTTPException(status_code=400, detail="Username 3-20 karakter, sadece harf/rakam/alt çizgi içerebilir.")
        existing = await db.users.find_one({"username": uname, "user_id": {"$ne": user.user_id}})
        if existing:
            raise HTTPException(status_code=409, detail="Bu username zaten kullanılıyor.")
        # 30-gün kısıtlaması — onboarding sırasında atla
        if not body.complete_onboarding:
            me = await db.users.find_one({"user_id": user.user_id}, {"username_changed_at": 1})
            last_change = me.get("username_changed_at") if me else None
            if last_change:
                delta = datetime.now(timezone.utc) - datetime.fromisoformat(last_change.replace("Z", "+00:00"))
                if delta.days < 30:
                    raise HTTPException(status_code=429, detail=f"Username değişikliği {30 - delta.days} gün sonra yapılabilir.")
        update["username"] = uname
        update["username_changed_at"] = datetime.now(timezone.utc).isoformat()
    if body.display_name is not None:
        update["display_name"] = body.display_name[:50]
    if body.bio is not None:
        update["bio"] = body.bio[:160]
    if body.complete_onboarding:
        update["needs_onboarding"] = False
    if update:
        await db.users.update_one({"user_id": user.user_id}, {"$set": update})
    return {"success": True}

@api_router.post("/users/{username}/follow")
async def follow_user(username: str, user: User = Depends(get_current_user)):
    target = await db.users.find_one({"username": username})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    if target["user_id"] == user.user_id:
        raise HTTPException(status_code=400, detail="Kendini takip edemezsin")
    already = await db.users.find_one({"user_id": user.user_id, "following": target["user_id"]})
    if already:
        return {"success": True, "following": True}
    await db.users.update_one({"user_id": user.user_id}, {"$addToSet": {"following": target["user_id"]}, "$inc": {"following_count": 1}})
    await db.users.update_one({"user_id": target["user_id"]}, {"$addToSet": {"followers": user.user_id}, "$inc": {"followers_count": 1}})
    return {"success": True, "following": True}

@api_router.delete("/users/{username}/follow")
async def unfollow_user(username: str, user: User = Depends(get_current_user)):
    target = await db.users.find_one({"username": username})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    await db.users.update_one({"user_id": user.user_id}, {"$pull": {"following": target["user_id"]}, "$inc": {"following_count": -1}})
    await db.users.update_one({"user_id": target["user_id"]}, {"$pull": {"followers": user.user_id}, "$inc": {"followers_count": -1}})
    return {"success": True, "following": False}

@api_router.post("/admin/set-verified")
async def set_verified(body: SetVerifiedRequest, user: User = Depends(get_current_user)):
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403, detail="Sadece CEO bu işlemi yapabilir.")
    target = await db.users.find_one({"username": body.username})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    await db.users.update_one({"user_id": target["user_id"]}, {"$set": {"verified_type": body.verified_type}})
    label = body.verified_type or "kaldırıldı"
    return {"success": True, "message": f"@{body.username} için verified: {label}"}

# ============ MEDIA / POSTS ============

class PostCreate(BaseModel):
    type: str  # 'text' | 'image' | 'video' | 'document'
    content: Optional[str] = None   # text body
    media_data: Optional[str] = None  # base64 image (image posts)
    media_url: Optional[str] = None   # video URL
    doc_id: Optional[str] = None      # document reference
    doc_title: Optional[str] = None

class PostCommentCreate(BaseModel):
    content: str

@api_router.post("/posts")
async def create_post(body: PostCreate, user: User = Depends(get_current_user)):
    """Yeni medya gönderisi oluştur."""
    check_rate_limit(user.user_id, "post", limit=10, window=60)
    me = await db.users.find_one({"user_id": user.user_id}, {"username": 1})
    if not me or not me.get("username"):
        raise HTTPException(status_code=400, detail="Gönderi atmak için önce kullanıcı adı belirlemelisin.")
    content = body.content or ""
    if len(content) > 2000:
        raise HTTPException(status_code=400, detail="Gönderi içeriği maksimum 2000 karakter olabilir.")
    if contains_profanity(content):
        raise HTTPException(status_code=400, detail="Bu içerik uygunsuz kelimeler içeriyor.")
    body.content = sanitize_text(content, 2000)
    post_id = f"post_{uuid.uuid4().hex[:16]}"
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "post_id": post_id,
        "author_id": user.user_id,
        "author_name": user.name,
        "author_picture": user.picture,
        "type": body.type,
        "content": body.content or "",
        "media_data": body.media_data,
        "media_url": body.media_url,
        "doc_id": body.doc_id,
        "doc_title": body.doc_title,
        "likes": [],
        "like_count": 0,
        "comment_count": 0,
        "created_at": now,
    }
    await db.posts.insert_one(doc)
    doc.pop("_id", None)
    return doc

def _boosted_pipeline(match: dict, skip: int, limit: int) -> list:
    """Boost aktifse üste çıkar, sonra tarih sırası."""
    now_iso = datetime.now(timezone.utc).isoformat()
    return [
        {"$match": match},
        {"$project": {"_id": 0, "doc": "$$ROOT", "boost_score": {
            "$cond": [{"$and": [
                {"$eq": ["$boost.active", True]},
                {"$gt": ["$boost.expires_at", now_iso]},
            ]}, "$boost.score", 0]
        }}},
        {"$replaceRoot": {"newRoot": {"$mergeObjects": ["$doc", {"boost_score": "$boost_score"}]}}},
        {"$sort": {"boost_score": -1, "created_at": -1}},
        {"$skip": skip},
        {"$limit": limit},
    ]

@api_router.get("/posts")
async def list_posts(skip: int = 0, limit: int = 20, u: dict = Depends(get_current_user_raw)):
    """Keşfet — tüm gönderiler, boosted önce."""
    uid = u["user_id"]
    me_raw = await db.users.find_one({"user_id": uid}, {"_id": 0})
    if me_raw:
        me_raw = sanitize_user_doc(me_raw)
    viewer_following = list((me_raw or {}).get("following", []))
    pipeline = _boosted_pipeline({}, skip, limit)
    posts = await db.posts.aggregate(pipeline).to_list(length=limit)
    for p in posts:
        await enrich_post(p, uid, viewer_following)
    return {"posts": posts}

@api_router.get("/feed")
async def get_feed(skip: int = 0, limit: int = 20, u: dict = Depends(get_current_user_raw)):
    """Feed — takip edilenlerin + kendi gönderileri; boosted önce."""
    uid = u["user_id"]
    me_raw = await db.users.find_one({"user_id": uid}, {"_id": 0})
    if me_raw:
        me_raw = sanitize_user_doc(me_raw)
    following = list((me_raw or {}).get("following", []))
    viewer_following = list(following)
    following.append(uid)
    match = {"author_id": {"$in": following}} if len(following) > 1 else {}
    pipeline = _boosted_pipeline(match, skip, limit)
    posts = await db.posts.aggregate(pipeline).to_list(length=limit)
    for p in posts:
        await enrich_post(p, uid, viewer_following)
    return {"posts": posts}

@api_router.get("/users/{username}/posts")
async def get_user_posts(username: str, skip: int = 0, limit: int = 20, u: dict = Depends(get_current_user_raw)):
    """Belirli kullanıcının gönderileri."""
    uid = u["user_id"]
    target = await db.users.find_one({"username": username})
    if not target:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    posts = await db.posts.find({"author_id": target["user_id"]}, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit).to_list(length=limit)
    for p in posts:
        await enrich_post(p, uid)
    return {"posts": posts}

@api_router.post("/posts/{post_id}/like")
async def toggle_like(post_id: str, user: User = Depends(get_current_user)):
    """Beğeni ekle / kaldır."""
    post = await db.posts.find_one({"post_id": post_id})
    if not post:
        raise HTTPException(status_code=404, detail="Gönderi bulunamadı")
    uid = user.user_id
    if uid in post.get("likes", []):
        await db.posts.update_one({"post_id": post_id}, {"$pull": {"likes": uid}, "$inc": {"like_count": -1}})
        liked = False
    else:
        await db.posts.update_one({"post_id": post_id}, {"$addToSet": {"likes": uid}, "$inc": {"like_count": 1}})
        liked = True
    updated = await db.posts.find_one({"post_id": post_id}, {"_id": 0, "like_count": 1})
    return {"liked": liked, "like_count": updated.get("like_count", 0)}

@api_router.delete("/posts/{post_id}")
async def delete_post(post_id: str, user: User = Depends(get_current_user)):
    """Gönderi sil — kendi gönderisini veya ayrıcalıklı hesaplar."""
    post = await db.posts.find_one({"post_id": post_id}, {"author_id": 1})
    if not post:
        raise HTTPException(status_code=404, detail="Gönderi bulunamadı")
    if post["author_id"] != user.user_id and not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    result = await db.posts.delete_one({"post_id": post_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Gönderi bulunamadı")
    await db.post_comments.delete_many({"post_id": post_id})
    return {"deleted": True}

@api_router.post("/posts/{post_id}/comments")
async def add_comment(post_id: str, body: PostCommentCreate, user: User = Depends(get_current_user)):
    """Yorum ekle."""
    check_rate_limit(user.user_id, "comment", limit=30, window=60)
    if len(body.content) > 500:
        raise HTTPException(status_code=400, detail="Yorum maksimum 500 karakter olabilir.")
    if contains_profanity(body.content):
        raise HTTPException(status_code=400, detail="Bu yorum uygunsuz kelimeler içeriyor.")
    body.content = sanitize_text(body.content, 500)
    post = await db.posts.find_one({"post_id": post_id})
    if not post:
        raise HTTPException(status_code=404, detail="Gönderi bulunamadı")
    comment_id = f"cmt_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    comment = {
        "comment_id": comment_id,
        "post_id": post_id,
        "author_id": user.user_id,
        "author_name": user.name,
        "author_picture": user.picture,
        "content": body.content,
        "created_at": now,
    }
    await db.post_comments.insert_one(comment)
    await db.posts.update_one({"post_id": post_id}, {"$inc": {"comment_count": 1}})
    comment.pop("_id", None)
    return comment

@api_router.get("/posts/{post_id}/comments")
async def list_comments(post_id: str, u: dict = Depends(get_current_user_raw)):
    """Yorumları listele."""
    comments = await db.post_comments.find({"post_id": post_id}, {"_id": 0}).sort("created_at", 1).to_list(length=200)
    return {"comments": comments}

@api_router.delete("/posts/{post_id}/comments/{comment_id}")
async def delete_comment(post_id: str, comment_id: str, user: User = Depends(get_current_user)):
    """Yorum sil — yorum sahibi veya ayrıcalıklı hesap."""
    comment = await db.post_comments.find_one({"comment_id": comment_id, "post_id": post_id})
    if not comment:
        raise HTTPException(status_code=404, detail="Yorum bulunamadı")
    if comment["author_id"] != user.user_id and not is_privileged(user.email):
        raise HTTPException(status_code=403, detail="Unauthorized")
    await db.post_comments.delete_one({"comment_id": comment_id})
    await db.posts.update_one({"post_id": post_id}, {"$inc": {"comment_count": -1}})
    return {"deleted": True}

# ============ BOOST ============

BOOST_PACKAGES = {
    "mini":     {"duration_hours": 1,    "credits": 50,   "score": 100,  "label": "Mini",     "price_try": 15,  "price_usd": 0.99},
    "standard": {"duration_hours": 6,    "credits": 150,  "score": 250,  "label": "Standart", "price_try": 50,  "price_usd": 2.99},
    "pro":      {"duration_hours": 24,   "credits": 400,  "score": 500,  "label": "Pro",      "price_try": 150, "price_usd": 7.99},
    "mega":     {"duration_hours": 24*7, "credits": 1200, "score": 1000, "label": "Mega",     "price_try": 500, "price_usd": 24.99},
}

class BoostRequest(BaseModel):
    tier: str  # mini, standard, pro, mega

@api_router.get("/boost/packages")
async def get_boost_packages():
    return {"packages": [{"tier": k, **v} for k, v in BOOST_PACKAGES.items()]}

@api_router.post("/posts/{post_id}/boost")
async def boost_post(post_id: str, body: BoostRequest, user: User = Depends(get_current_user)):
    if body.tier not in BOOST_PACKAGES:
        raise HTTPException(status_code=400, detail="Geçersiz boost paketi.")
    pkg = BOOST_PACKAGES[body.tier]
    post = await db.posts.find_one({"post_id": post_id})
    if not post:
        raise HTTPException(status_code=404, detail="Gönderi bulunamadı.")
    if post["author_id"] != user.user_id:
        raise HTTPException(status_code=403, detail="Sadece kendi gönderini öne çıkarabilirsin.")
    now = datetime.now(timezone.utc)
    if post.get("boost", {}).get("active") and post["boost"].get("expires_at", "") > now.isoformat():
        raise HTTPException(status_code=400, detail="Bu gönderide aktif bir boost zaten var.")
    # Kredi düş
    user_data = await db.users.find_one({"user_id": user.user_id}, {"bonus_credits": 1})
    bonus = user_data.get("bonus_credits", 0) if user_data else 0
    if bonus < pkg["credits"]:
        raise HTTPException(status_code=402, detail=f"Yetersiz kredi. Gerekli: {pkg['credits']}, Mevcut: {bonus}")
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"bonus_credits": -pkg["credits"]}})
    expires_at = (now + timedelta(hours=pkg["duration_hours"])).isoformat()
    await db.posts.update_one({"post_id": post_id}, {"$set": {
        "boost": {"active": True, "tier": body.tier, "score": pkg["score"],
                  "expires_at": expires_at, "purchased_at": now.isoformat()},
    }})
    return {"success": True, "tier": body.tier, "expires_at": expires_at, "credits_spent": pkg["credits"]}

@api_router.post("/auth/exchange")
async def exchange_token(request: Request, response: Response):
    """Frontend'den gelen token'ı session cookie'ye dönüştür."""
    body = await request.json()
    token = body.get("token")
    if not token:
        raise HTTPException(status_code=400, detail="token required")
    session = await db.user_sessions.find_one({"session_token": token})
    if not session:
        raise HTTPException(status_code=401, detail="Invalid token")
    expires_at = session.get("expires_at")
    if not expires_at:
        raise HTTPException(status_code=401, detail="Invalid token")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Token expired")
    user = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    is_production = os.getenv("REACT_APP_BACKEND_URL", "").startswith("https")
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7*24*60*60
    )
    return {
        "user_id": user.get("user_id"),
        "email": user.get("email"),
        "name": user.get("name", ""),
        "picture": user.get("picture"),
        "subscription": user.get("subscription", "free"),
        "needs_onboarding": user.get("needs_onboarding", False),
        "username": user.get("username"),
    }


@api_router.get("/auth/me")
async def get_me(user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    if not user_data:
        raise HTTPException(status_code=404, detail="User not found")
    return {
        "user_id": user_data.get("user_id"),
        "email": user_data.get("email"),
        "name": user_data.get("name", ""),
        "picture": user_data.get("picture"),
        "username": user_data.get("username"),
        "display_name": user_data.get("display_name"),
        "bio": user_data.get("bio", ""),
        "verified_type": user_data.get("verified_type"),
        "subscription": user_data.get("subscription", "free"),
        "subscription_date": user_data.get("subscription_date"),
        "needs_onboarding": user_data.get("needs_onboarding", False),
        "credits": user_data.get("mindshare_credits", user_data.get("credits", 0)),
        "sp_balance": user_data.get("sp_balance", user_data.get("zc_balance", 0)),
        "rank": user_data.get("mindshare_rank", user_data.get("rank", "iron")),
        "quest_xp": user_data.get("mindshare_xp", user_data.get("quest_xp", user_data.get("xp", 0))),
        "followers_count": user_data.get("followers_count", 0),
        "following_count": user_data.get("following_count", 0),
    }

class ProfileUpdate(BaseModel):
    name: Optional[str] = None

class ProfilePictureUpload(BaseModel):
    image_data: str  # Base64 encoded image data

@api_router.put("/auth/profile")
async def update_profile(req: ProfileUpdate, user: User = Depends(get_current_user)):
    update_data = {}
    if req.name is not None:
        update_data["name"] = req.name
        update_data["name_custom"] = True

    if update_data:
        await db.users.update_one({"user_id": user.user_id}, {"$set": update_data})

    return {"message": "Profile updated", "name": req.name}

@api_router.post("/auth/profile-picture")
async def upload_profile_picture(req: ProfilePictureUpload, user: User = Depends(get_current_user)):
    image_data = req.image_data
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {"picture": image_data, "picture_custom": True}}
    )
    return {"message": "Profile picture updated", "picture_url": image_data}

class EmailChangeRequest(BaseModel):
    new_email: str

@api_router.post("/auth/change-email/request")
async def request_email_change(req: EmailChangeRequest, user: User = Depends(get_current_user)):
    new_email = req.new_email.strip().lower()
    if not new_email or "@" not in new_email:
        raise HTTPException(status_code=400, detail="Geçersiz e-posta adresi")
    # Check if already in use
    existing = await db.users.find_one({"email": new_email})
    if existing:
        raise HTTPException(status_code=409, detail="Bu e-posta adresi zaten kullanılıyor")
    token = uuid.uuid4().hex
    expires_at = datetime.now(timezone.utc).timestamp() + 3600  # 1 hour
    await db.email_change_tokens.insert_one({
        "user_id": user.user_id,
        "new_email": new_email,
        "token": token,
        "expires_at": expires_at,
    })
    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
    confirm_url = f"{frontend_url}/confirm-email-change?token={token}"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 24px; background: #111827; color: #fff; border-radius: 12px;">
        <h2 style="color: #4ca8ad; margin-bottom: 16px;">📧 E-posta Değiştirme Onayı</h2>
        <p style="color: #d1d5db; margin-bottom: 8px;">Hesabınızın e-posta adresini değiştirmek istediniz.</p>
        <p style="color: #d1d5db; margin-bottom: 20px;"><strong>Yeni e-posta:</strong> {new_email}</p>
        <p style="color: #d1d5db; margin-bottom: 24px;">Bu isteği siz yapmadıysanız bu e-postayı görmezden gelebilirsiniz.</p>
        <a href="{confirm_url}" style="display: inline-block; background: #4ca8ad; color: white; padding: 12px 28px; border-radius: 8px; text-decoration: none; font-weight: bold; font-size: 15px;">E-postayı Onayla</a>
        <p style="color: #6b7280; font-size: 12px; margin-top: 20px;">Bu bağlantı 1 saat geçerlidir.</p>
    </div>
    """
    await send_email(user.email, "📧 ZET Mindshare - E-posta Değiştirme Onayı", html_content)
    return {"message": "Confirmation email sent"}

@api_router.post("/auth/change-email/confirm")
async def confirm_email_change(token: str = Body(..., embed=True)):
    record = await db.email_change_tokens.find_one({"token": token})
    if not record:
        raise HTTPException(status_code=400, detail="Geçersiz veya süresi dolmuş bağlantı")
    if datetime.now(timezone.utc).timestamp() > record["expires_at"]:
        await db.email_change_tokens.delete_one({"token": token})
        raise HTTPException(status_code=400, detail="Bağlantının süresi dolmuş")
    # Double-check uniqueness at confirm time
    existing = await db.users.find_one({"email": record["new_email"]})
    if existing:
        await db.email_change_tokens.delete_one({"token": token})
        raise HTTPException(status_code=409, detail="Bu e-posta adresi artık kullanılıyor")
    await db.users.update_one(
        {"user_id": record["user_id"]},
        {"$set": {"email": record["new_email"]}}
    )
    await db.email_change_tokens.delete_one({"token": token})
    return {"message": "Email updated", "new_email": record["new_email"]}

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    session_token = request.cookies.get("session_token")
    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})
    response.delete_cookie(key="session_token", path="/")
    return {"message": "Logged out"}

@api_router.post("/auth/delete-account/request")
async def request_account_deletion(user: User = Depends(get_current_user)):
    """Send a confirmation email with a deletion token."""
    token = uuid.uuid4().hex
    expires_at = (datetime.now(timezone.utc).timestamp() + 3600)  # 1 hour
    await db.deletion_tokens.insert_one({
        "user_id": user.user_id,
        "token": token,
        "expires_at": expires_at,
    })
    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
    confirm_url = f"{frontend_url}/confirm-delete?token={token}"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 24px; background: #111827; color: #fff; border-radius: 12px;">
        <h2 style="color: #ef4444; margin-bottom: 16px;">⚠️ Hesap Silme Onayı</h2>
        <p style="color: #d1d5db; margin-bottom: 20px;">Hesabınızı silmek istediğinizi aldık. Bu işlem <strong>geri alınamaz</strong>. Tüm notlarınız, dokümanlarınız ve verileriniz kalıcı olarak silinecektir.</p>
        <p style="color: #d1d5db; margin-bottom: 24px;">Bu isteği siz yapmadıysanız bu e-postayı görmezden gelebilirsiniz.</p>
        <a href="{confirm_url}" style="display: inline-block; background: #ef4444; color: white; padding: 12px 28px; border-radius: 8px; text-decoration: none; font-weight: bold; font-size: 15px;">Hesabı Kalıcı Olarak Sil</a>
        <p style="color: #6b7280; font-size: 12px; margin-top: 20px;">Bu bağlantı 1 saat geçerlidir.</p>
    </div>
    """
    await send_email(user.email, "⚠️ ZET Mindshare - Hesap Silme Onayı", html_content)
    return {"message": "Confirmation email sent"}

@api_router.post("/auth/delete-account/confirm")
async def confirm_account_deletion(token: str = Body(..., embed=True), response: Response = None):
    """Confirm deletion using the token from email."""
    record = await db.deletion_tokens.find_one({"token": token})
    if not record:
        raise HTTPException(status_code=400, detail="Geçersiz veya süresi dolmuş bağlantı")
    if datetime.now(timezone.utc).timestamp() > record["expires_at"]:
        await db.deletion_tokens.delete_one({"token": token})
        raise HTTPException(status_code=400, detail="Bağlantının süresi dolmuş")
    user_id = record["user_id"]
    # Delete all user data
    await db.users.delete_one({"user_id": user_id})
    await db.user_sessions.delete_many({"user_id": user_id})
    await db.quick_notes.delete_many({"user_id": user_id})
    await db.documents.delete_many({"user_id": user_id})
    await db.notebooks.delete_many({"user_id": user_id})
    await db.deletion_tokens.delete_one({"token": token})
    if response:
        response.delete_cookie(key="session_token", path="/")
    return {"message": "Account deleted"}

# ============ EMAIL AUTH ROUTES ============

@api_router.post("/auth/register")
async def register_with_email(req: EmailAuthRequest, response: Response):
    normalized_email = req.email.lower().strip()
    existing = await db.users.find_one({"email": normalized_email})
    if existing:
        raise HTTPException(status_code=400, detail="Bu e-posta zaten kayıtlı")

    user_id = f"user_{secrets.token_hex(6)}"
    hashed_pw = hash_password(req.password)
    
    user_data = {
        "user_id": user_id,
        "email": req.email,
        "name": req.name or req.email.split('@')[0],
        "picture": None,
        "hashed_password": hashed_pw,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_data)

    # Hoş geldin maili
    welcome_html = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 24px; background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); color: #fff; border-radius: 12px;">
        <h2 style="color: #4ca8ad; margin-bottom: 16px;">ZET Mindshare'e Hoş Geldiniz! 🎉</h2>
        <p style="font-size: 16px; line-height: 1.6;">Merhaba <strong>{user_data['name']}</strong>,</p>
        <p style="font-size: 15px; line-height: 1.6;">Hesabınız başarıyla oluşturuldu. Artık yapay zeka destekli not alma, belge oluşturma ve daha fazlasından yararlanabilirsiniz.</p>
        <div style="background: rgba(76,168,173,0.15); border: 1px solid rgba(76,168,173,0.3); padding: 16px; border-radius: 8px; margin: 20px 0;">
            <p style="margin: 4px 0; font-size: 14px;">✅ Ücretsiz plan: Günlük 20 kredi</p>
            <p style="margin: 4px 0; font-size: 14px;">✅ Sınırsız belge</p>
            <p style="margin: 4px 0; font-size: 14px;">✅ AI sohbet (ZETA)</p>
        </div>
        <a href="https://zetmindshare.com" style="display: inline-block; background: #4ca8ad; color: white; padding: 12px 24px; border-radius: 6px; text-decoration: none; font-weight: bold;">Hemen Başla</a>
        <p style="color: #666; font-size: 12px; margin-top: 20px;">Sorunuz mu var? <a href="mailto:support@zetstudiointl.com" style="color: #4ca8ad;">support@zetstudiointl.com</a></p>
    </div>
    """
    asyncio.create_task(send_email(req.email, "ZET Mindshare'e Hoş Geldiniz! 🎉", welcome_html))

    # Create session
    session_token = f"sess_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)

    await db.user_sessions.insert_one({
        "session_token": session_token,
        "user_id": user_id,
        "expires_at": expires_at.isoformat()
    })

    response.set_cookie(
        key="session_token", value=session_token,
        httponly=True, secure=True,
        samesite="none",
        path="/", max_age=7*24*60*60
    )

    return {"user": {"user_id": user_id, "email": req.email, "name": user_data["name"]}, "session_token": session_token}

@api_router.post("/auth/login")
async def login_with_email(req: EmailAuthRequest, response: Response, request: Request):
    # Email'i normalize et
    normalized_email = req.email.lower().strip()
    ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "").split(",")[0].strip()

    if await hafizz.is_login_locked(db, normalized_email):
        raise HTTPException(status_code=429, detail="Çok fazla başarısız deneme. 15 dakika bekleyin.")

    user = await db.users.find_one({"email": normalized_email})
    if not user:
        raise HTTPException(status_code=401, detail="Geçersiz e-posta veya şifre")

    if not user.get("hashed_password"):
        raise HTTPException(status_code=401, detail="Bu hesap Google ile oluşturulmuş. Lütfen Google ile giriş yapın.")
    if not verify_password(req.password, user["hashed_password"]):
        fail_count = await hafizz.record_failed_login(db, normalized_email, "")
        if fail_count >= 5:
            asyncio.create_task(hafizz.notify_lockout(db, normalized_email))
        raise HTTPException(status_code=401, detail="Şifre yanlış. Lütfen tekrar deneyin.")

    # Mevcut user_id'yi kullan — asla yeni üretme
    user_id = user["user_id"]
    logging.info(f"Email login: user_id={user_id} email={normalized_email}")

    await hafizz.clear_failed_logins(db, normalized_email)

    # Hafız: device fingerprint + ülke kontrolü (background)
    ua = request.headers.get("User-Agent", "")
    lang = request.headers.get("Accept-Language", "")
    fp = hafizz.compute_device_fingerprint(ua, lang)
    asyncio.create_task(hafizz.record_device_login(db, user_id, fp, ip))
    asyncio.create_task(_hafizz_post_login(db, user_id, ip))

    session_token = f"sess_{secrets.token_hex(16)}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)

    await db.user_sessions.insert_one({
        "session_token": session_token,
        "user_id": user_id,
        "expires_at": expires_at.isoformat()
    })

    response.set_cookie(
        key="session_token", value=session_token,
        httponly=True, secure=True,
        samesite="none",
        path="/", max_age=7*24*60*60
    )

    return {"user": {"user_id": user_id, "email": user["email"], "name": user.get("name", "")}, "session_token": session_token}

# ============ APPLE AUTH ROUTES ============

@api_router.get("/auth/apple/init")
async def apple_auth_init():
    """Return Apple OAuth URL if configured, otherwise indicate not configured."""
    client_id = os.environ.get("APPLE_CLIENT_ID")
    if not client_id:
        return {"auth_url": None, "message": "Apple Sign-In henuz yapilandirilmadi"}
    redirect_uri = os.environ.get("APPLE_REDIRECT_URI", os.environ.get("REACT_APP_BACKEND_URL", "") + "/api/auth/apple/callback")
    auth_url = (
        f"https://appleid.apple.com/auth/authorize?"
        f"client_id={client_id}&redirect_uri={redirect_uri}"
        f"&response_type=code%20id_token&scope=name%20email"
        f"&response_mode=form_post"
    )
    return {"auth_url": auth_url}

@api_router.post("/auth/apple/callback")
async def apple_auth_callback(request: Request, response: Response):
    """Handle Apple Sign-In callback with id_token."""
    try:
        body = await request.json()
    except:
        form = await request.form()
        body = dict(form)
    
    id_token = body.get("id_token")
    apple_user = body.get("user")
    
    if not id_token:
        raise HTTPException(status_code=400, detail="id_token eksik")
    
    import jwt as pyjwt
    try:
        claims = pyjwt.decode(id_token, options={"verify_signature": False})
    except Exception:
        raise HTTPException(status_code=400, detail="Geçersiz id_token")
    
    apple_sub = claims.get("sub")
    email = claims.get("email")
    
    if not apple_sub:
        raise HTTPException(status_code=400, detail="Apple ID bulunamadi")
    
    existing = await db.users.find_one({"apple_id": apple_sub})
    if existing:
        user_id = existing["user_id"]
        user_name = existing.get("name", "Apple User")
    else:
        if not email:
            email = f"{apple_sub[:8]}@privaterelay.appleid.com"
        
        email_user = await db.users.find_one({"email": email})
        if email_user:
            await db.users.update_one({"email": email}, {"$set": {"apple_id": apple_sub}})
            user_id = email_user["user_id"]
            user_name = email_user.get("name", "Apple User")
        else:
            user_name = "Apple User"
            if apple_user and isinstance(apple_user, dict):
                fn = apple_user.get("name", {}).get("firstName", "")
                ln = apple_user.get("name", {}).get("lastName", "")
                if fn or ln:
                    user_name = f"{fn} {ln}".strip()
            
            user_id = f"user_{uuid.uuid4().hex[:12]}"
            await db.users.insert_one({
                "user_id": user_id,
                "email": email,
                "name": user_name,
                "apple_id": apple_sub,
                "picture": None,
                "created_at": datetime.now(timezone.utc).isoformat()
            })
    
    session_token = f"sess_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "session_token": session_token,
        "user_id": user_id,
        "expires_at": expires_at.isoformat()
    })
    is_production = os.getenv("REACT_APP_BACKEND_URL", "").startswith("https")
    response.set_cookie(
        key="session_token", value=session_token,
        httponly=True, secure=True,
        samesite="none",
        path="/", max_age=7*24*60*60
    )
    return {"user": {"user_id": user_id, "email": email, "name": user_name}}

# ============ ZET ID (OAuth/SSO) ============

async def _zet_id_create_session_cookie(response: Response, user_id: str):
    session_token = f"sess_{secrets.token_hex(16)}"
    session_expires = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "session_token": session_token,
        "user_id": user_id,
        "expires_at": session_expires.isoformat()
    })
    response.set_cookie(
        key="session_token", value=session_token,
        httponly=True, secure=True,
        samesite="none",
        path="/", max_age=7*24*60*60
    )
    return session_token


def _zet_id_user_brief(user_doc: dict) -> dict:
    return {
        "user_id": user_doc["user_id"],
        "email": user_doc.get("email"),
        "name": user_doc.get("name", ""),
        "picture": user_doc.get("picture"),
        "zet_id": user_doc.get("zet_id"),
    }


@api_router.post("/zet-id/create")
async def zet_id_create(request: Request, response: Response, user: User = Depends(get_current_user)):
    """Aktif oturuma ZET ID atar (yoksa üretir — mevcut kullanıcılar için otomatik migrasyon) ve bu cihaz için token çifti üretir."""
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    zet_id, is_new = await ensure_zet_id(user_doc)
    device_hash = compute_zet_id_device_hash(request)
    tokens = await issue_zet_id_session(user.user_id, zet_id, device_hash)
    user_doc["zet_id"] = zet_id
    return {"zet_id": zet_id, "is_new": is_new, "user": _zet_id_user_brief(user_doc), **tokens}


@api_router.post("/zet-id/login")
async def zet_id_login(req: ZetIdLoginRequest, request: Request, response: Response):
    """Cihazda kayıtlı ZET ID hesabıyla tek dokunuşla giriş — Apple ID seçici deneyimi."""
    device_hash = compute_zet_id_device_hash(request)

    claims = None
    try:
        claims = decode_zet_id_access_token(req.access_token)
    except pyjwt.ExpiredSignatureError:
        claims = None
    except Exception:
        raise HTTPException(status_code=401, detail="Geçersiz ZET ID token")

    session_record = None
    if claims:
        session_record = await db.zet_id_sessions.find_one({"jti": claims.get("jti")}, {"_id": 0})

    if not session_record or session_record.get("revoked"):
        if not req.refresh_token:
            raise HTTPException(status_code=401, detail="ZET ID oturumunun süresi dolmuş — yeniden giriş gerekli")
        session_record = await db.zet_id_sessions.find_one(
            {"refresh_token": req.refresh_token, "revoked": {"$ne": True}}, {"_id": 0}
        )
        if not session_record:
            raise HTTPException(status_code=401, detail="ZET ID oturumunun süresi dolmuş — yeniden giriş gerekli")
        rt_exp = session_record.get("expires_at")
        if isinstance(rt_exp, str):
            rt_exp = datetime.fromisoformat(rt_exp)
        if rt_exp.tzinfo is None:
            rt_exp = rt_exp.replace(tzinfo=timezone.utc)
        if rt_exp < datetime.now(timezone.utc):
            raise HTTPException(status_code=401, detail="ZET ID oturumunun süresi dolmuş — yeniden giriş gerekli")

    if session_record["device_hash"] != device_hash:
        # Şüpheli giriş — farklı cihaz: yeniden kimlik doğrulama gerekiyor
        raise HTTPException(status_code=409, detail="suspicious_login")

    user_doc = await db.users.find_one({"user_id": session_record["user_id"]}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")

    # Token rotasyonu — eski oturum iptal, yenisi üretilir
    await db.zet_id_sessions.update_one({"jti": session_record["jti"]}, {"$set": {"revoked": True}})
    tokens = await issue_zet_id_session(user_doc["user_id"], session_record["zet_id"], device_hash)
    session_token = await _zet_id_create_session_cookie(response, user_doc["user_id"])

    user_doc["zet_id"] = session_record["zet_id"]
    return {"user": _zet_id_user_brief(user_doc), "session_token": session_token, **tokens}


@api_router.post("/zet-id/connect-google")
async def zet_id_connect_google(req: ZetIdConnectTokenRequest, user: User = Depends(get_current_user)):
    """Mevcut ZET ID hesabına bir Google hesabı bağlar."""
    try:
        claims = pyjwt.decode(req.id_token, options={"verify_signature": False})
    except Exception:
        raise HTTPException(status_code=400, detail="Geçersiz id_token")
    email = (claims.get("email") or "").lower().strip()
    if not email:
        raise HTTPException(status_code=400, detail="Google hesabında e-posta bulunamadı")
    other = await db.users.find_one({"email": email, "user_id": {"$ne": user.user_id}}, {"_id": 0, "user_id": 1})
    if other:
        raise HTTPException(status_code=400, detail="Bu Google hesabı başka bir ZET ID'ye bağlı")
    await db.users.update_one({"user_id": user.user_id}, {"$set": {"connected_accounts.google": email}})
    return {"connected": "google", "email": email}


@api_router.post("/zet-id/connect-apple")
async def zet_id_connect_apple(req: ZetIdConnectTokenRequest, user: User = Depends(get_current_user)):
    """Mevcut ZET ID hesabına bir Apple hesabı bağlar."""
    try:
        claims = pyjwt.decode(req.id_token, options={"verify_signature": False})
    except Exception:
        raise HTTPException(status_code=400, detail="Geçersiz id_token")
    apple_sub = claims.get("sub")
    email = claims.get("email")
    if not apple_sub:
        raise HTTPException(status_code=400, detail="Apple kimliği bulunamadı")
    other = await db.users.find_one({"apple_id": apple_sub, "user_id": {"$ne": user.user_id}}, {"_id": 0, "user_id": 1})
    if other:
        raise HTTPException(status_code=400, detail="Bu Apple hesabı başka bir ZET ID'ye bağlı")
    update = {"apple_id": apple_sub, "connected_accounts.apple": email or apple_sub}
    await db.users.update_one({"user_id": user.user_id}, {"$set": update})
    return {"connected": "apple", "email": email}


@api_router.post("/zet-id/connect-email")
async def zet_id_connect_email(req: ZetIdConnectEmailRequest, user: User = Depends(get_current_user)):
    """Mevcut ZET ID hesabına e-posta + şifre ile giriş yöntemi ekler."""
    normalized = req.email.lower().strip()
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Şifre en az 6 karakter olmalı")
    other = await db.users.find_one({"email": normalized, "user_id": {"$ne": user.user_id}}, {"_id": 0, "user_id": 1})
    if other:
        raise HTTPException(status_code=400, detail="Bu e-posta başka bir ZET ID'ye kayıtlı")
    hashed = hash_password(req.password)
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {"hashed_password": hashed, "email": normalized, "connected_accounts.email": True}}
    )
    return {"connected": "email", "email": normalized}


@api_router.post("/zet-id/logout")
async def zet_id_logout(request: Request):
    """Bu cihaza ait ZET ID oturumunu (erişim token'ı) iptal eder."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Bearer token gerekli")
    token = auth_header.split(" ")[1]
    try:
        claims = decode_zet_id_access_token(token)
    except Exception:
        raise HTTPException(status_code=401, detail="Geçersiz token")
    await db.zet_id_sessions.update_one({"jti": claims.get("jti")}, {"$set": {"revoked": True}})
    return {"success": True}


@api_router.get("/zet-id/profile")
async def zet_id_profile(user: User = Depends(get_current_user)):
    """ZET ID, bağlı hesaplar ve cihaz oturumlarını döner."""
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")
    sessions = await db.zet_id_sessions.find(
        {"user_id": user.user_id, "revoked": {"$ne": True}}, {"_id": 0, "refresh_token": 0}
    ).to_list(50)
    return {
        "zet_id": user_doc.get("zet_id"),
        "email": user_doc.get("email"),
        "name": user_doc.get("name"),
        "picture": user_doc.get("picture"),
        "connected_accounts": user_doc.get("connected_accounts", {}),
        "devices": [
            {"jti": s["jti"], "device_hash": s["device_hash"], "created_at": s["created_at"], "expires_at": s["expires_at"]}
            for s in sessions
        ],
    }


@api_router.delete("/zet-id/remove-device")
async def zet_id_remove_device(jti: str = Query(...), user: User = Depends(get_current_user)):
    """Belirtilen cihaz oturumunu iptal eder — 'Bu cihazdan kaldır'."""
    result = await db.zet_id_sessions.update_one(
        {"jti": jti, "user_id": user.user_id}, {"$set": {"revoked": True}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Cihaz bulunamadı")
    return {"success": True}


@api_router.get("/zet-id/authorize")
async def zet_id_authorize(request: Request, client_id: str = Query(...), redirect_uri: str = Query(...), response_type: str = Query("code"), state: str = Query(None), scope: str = Query("profile")):
    """OAuth2 yetkilendirme uç noktası — üçüncü parti uygulamalar ZET ID ile giriş için buraya yönlendirir."""
    from fastapi.responses import RedirectResponse
    from urllib.parse import quote

    client = await db.zet_id_clients.find_one({"client_id": client_id}, {"_id": 0})
    if not client or redirect_uri not in client.get("redirect_uris", []):
        raise HTTPException(status_code=400, detail="Geçersiz client_id veya redirect_uri")
    if response_type != "code":
        raise HTTPException(status_code=400, detail="Sadece 'code' response_type desteklenir")

    user = None
    session_token = request.cookies.get("session_token")
    if not session_token:
        auth_header = request.headers.get("Authorization", "")
        if auth_header.startswith("Bearer "):
            session_token = auth_header.split(" ")[1]
    if session_token:
        session = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0})
        if session:
            sess_exp = session.get("expires_at")
            if isinstance(sess_exp, str):
                sess_exp = datetime.fromisoformat(sess_exp)
            if sess_exp.tzinfo is None:
                sess_exp = sess_exp.replace(tzinfo=timezone.utc)
            if sess_exp > datetime.now(timezone.utc):
                user = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})

    frontend_url = os.environ.get("FRONTEND_URL", "https://zetmindshare.com")
    if not user:
        return_to = quote(str(request.url), safe="")
        return RedirectResponse(f"{frontend_url}/login?return_to={return_to}")

    code = f"zidc_{secrets.token_hex(24)}"
    await db.zet_id_auth_codes.insert_one({
        "code": code,
        "client_id": client_id,
        "user_id": user["user_id"],
        "redirect_uri": redirect_uri,
        "scope": scope,
        "expires_at": (datetime.now(timezone.utc) + timedelta(minutes=5)).isoformat(),
        "used": False
    })
    sep = "&" if "?" in redirect_uri else "?"
    redirect_to = f"{redirect_uri}{sep}code={code}"
    if state:
        redirect_to += f"&state={quote(state, safe='')}"
    return RedirectResponse(redirect_to)


@api_router.post("/zet-id/token")
async def zet_id_token_exchange(req: ZetIdTokenRequest):
    """OAuth2 token uç noktası — yetkilendirme kodunu ZET ID erişim token'ına (30 gün) çevirir."""
    if req.grant_type != "authorization_code":
        raise HTTPException(status_code=400, detail="Sadece 'authorization_code' grant_type desteklenir")
    if not req.code:
        raise HTTPException(status_code=400, detail="code zorunlu")

    client = await db.zet_id_clients.find_one({"client_id": req.client_id, "client_secret": req.client_secret}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=401, detail="Geçersiz client kimliği")

    auth_code = await db.zet_id_auth_codes.find_one({"code": req.code}, {"_id": 0})
    if not auth_code or auth_code.get("used") or auth_code["client_id"] != req.client_id:
        raise HTTPException(status_code=400, detail="Geçersiz veya kullanılmış kod")

    code_exp = auth_code.get("expires_at")
    if isinstance(code_exp, str):
        code_exp = datetime.fromisoformat(code_exp)
    if code_exp.tzinfo is None:
        code_exp = code_exp.replace(tzinfo=timezone.utc)
    if code_exp < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Kodun süresi dolmuş")
    if req.redirect_uri and auth_code["redirect_uri"] != req.redirect_uri:
        raise HTTPException(status_code=400, detail="redirect_uri eşleşmiyor")

    await db.zet_id_auth_codes.update_one({"code": req.code}, {"$set": {"used": True}})

    token = f"zid_{secrets.token_hex(32)}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=30)
    await db.zet_id_tokens.insert_one({
        "token": token,
        "user_id": auth_code["user_id"],
        "client_id": req.client_id,
        "scope": auth_code.get("scope", "profile"),
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    return {
        "access_token": token,
        "token_type": "Bearer",
        "expires_in": 30 * 24 * 60 * 60,
        "scope": auth_code.get("scope", "profile")
    }


@api_router.get("/zet-id/userinfo")
async def zet_id_userinfo(request: Request):
    """OAuth2 userinfo uç noktası — Bearer ZET ID token ile kullanıcı profilini döner."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Bearer token gerekli")
    token = auth_header.split(" ")[1]

    record = await db.zet_id_tokens.find_one({"token": token}, {"_id": 0})
    if not record:
        raise HTTPException(status_code=401, detail="Geçersiz token")

    expires_at = record.get("expires_at")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Token süresi dolmuş")

    user = await db.users.find_one({"user_id": record["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")

    return {
        "sub": user["user_id"],
        "email": user.get("email"),
        "name": user.get("name"),
        "picture": user.get("picture"),
        "username": user.get("username"),
    }

# ============ DOCUMENTS ROUTES ============

@api_router.get("/documents", response_model=List[dict])
async def get_documents(skip: int = 0, limit: int = 20, user: User = Depends(get_current_user)):
    try:
        docs = await db.documents.find(
            {"user_id": user.user_id, "deleted": {"$ne": True}},
            {"_id": 0, "pages": 0}
        ).sort([("pinned", -1), ("updated_at", -1)]).skip(skip).limit(limit).to_list(limit)
        return docs
    except Exception as e:
        logging.error(f"get_documents error: {e}")
        raise HTTPException(status_code=500, detail="Belgeler yüklenemedi")

class R2UploadRequest(BaseModel):
    data: str  # data:image/...;base64,...

@api_router.post("/r2/upload")
async def upload_image_to_r2(body: R2UploadRequest, user: User = Depends(get_current_user)):
    try:
        if not body.data.startswith('data:'):
            raise HTTPException(status_code=400, detail="Geçersiz görsel verisi")
        header, b64data = body.data.split(',', 1)
        content_type = header.split(';')[0].replace('data:', '')
        ext_map = {'image/png': 'png', 'image/jpeg': 'jpg', 'image/gif': 'gif', 'image/webp': 'webp'}
        ext = ext_map.get(content_type, 'png')
        image_bytes = base64.b64decode(b64data)
        key = f"canvas/{user.user_id}/{uuid.uuid4().hex}.{ext}"
        r2 = _get_r2()
        await asyncio.to_thread(
            r2.put_object,
            Bucket=R2_BUCKET,
            Key=key,
            Body=image_bytes,
            ContentType=content_type,
        )
        url = f"{R2_PUBLIC_URL}/{key}"
        return {"url": url}
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"R2 upload error: {e}")
        raise HTTPException(status_code=503, detail="Görsel yükleme başarısız")

@api_router.post("/documents")
async def create_document(doc: DocumentCreate, user: User = Depends(get_current_user)):
    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    doc_dict = {
        "doc_id": doc_id,
        "user_id": user.user_id,
        "title": doc.title,
        "doc_type": doc.doc_type,
        "content": {},
        "pages": [{"page_id": "page_1", "content": {}}],
        "created_at": now.isoformat(),
        "updated_at": now.isoformat()
    }
    await db.documents.insert_one(doc_dict)
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"docs_created": 1}})
    # Return without _id
    return {k: v for k, v in doc_dict.items() if k != "_id"}

@api_router.get("/documents/{doc_id}")
async def get_document(doc_id: str, user: User = Depends(get_current_user)):
    doc = await db.documents.find_one({"doc_id": doc_id, "user_id": user.user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

_version_throttle: dict = {}  # doc_id → last saved unix timestamp
_VERSION_INTERVAL = 300  # 5 dakika

async def _save_document_version(doc_id: str, user_id: str):
    import time
    now_ts = time.time()
    if now_ts - _version_throttle.get(doc_id, 0) < _VERSION_INTERVAL:
        return
    _version_throttle[doc_id] = now_ts
    try:
        old = await db.documents.find_one({"doc_id": doc_id, "user_id": user_id}, {"_id": 0, "pages": 1})
        if not old or not old.get("pages"):
            return
        pages_stripped = [
            {**page, "elements": [
                # Strip only base64 data: URLs (large) — keep R2/external URLs intact
                {**el, "src": None} if el.get("type") == "image" and isinstance(el.get("src"), str) and el["src"].startswith("data:") else el
                for el in page.get("elements", [])
            ]}
            for page in old["pages"]
        ]
        now = datetime.now(timezone.utc).isoformat()
        existing = await db.document_history.find(
            {"doc_id": doc_id}, {"_id": 0, "slot": 1, "saved_at": 1}
        ).sort("saved_at", 1).to_list(4)
        slot = len(existing) if len(existing) < 3 else existing[0].get("slot", 0)
        await db.document_history.replace_one(
            {"doc_id": doc_id, "slot": slot},
            {"doc_id": doc_id, "user_id": user_id, "slot": slot, "pages": pages_stripped, "saved_at": now},
            upsert=True,
        )
    except Exception:
        pass

@api_router.put("/documents/{doc_id}")
async def update_document(doc_id: str, update: DocumentUpdate, user: User = Depends(get_current_user)):
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()

    if "pages" in update_data:
        asyncio.create_task(_save_document_version(doc_id, user.user_id))

    try:
        result = await db.documents.update_one(
            {"doc_id": doc_id, "user_id": user.user_id},
            {"$set": update_data}
        )
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Document not found")
        doc = await db.documents.find_one({"doc_id": doc_id}, {"_id": 0})
        return doc
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"update_document DB error {doc_id}: {e}")
        raise HTTPException(status_code=503, detail="Service temporarily unavailable")

@api_router.get("/documents/{doc_id}/history")
async def get_document_history(doc_id: str, user: User = Depends(get_current_user)):
    versions = await db.document_history.find(
        {"doc_id": doc_id, "user_id": user.user_id},
        {"_id": 0, "pages": 0}
    ).sort("saved_at", -1).to_list(50)
    return versions

@api_router.post("/documents/{doc_id}/restore/{version_index}")
async def restore_document_version(doc_id: str, version_index: int, user: User = Depends(get_current_user)):
    versions = await db.document_history.find(
        {"doc_id": doc_id, "user_id": user.user_id}
    ).sort("saved_at", -1).to_list(50)
    if version_index >= len(versions):
        raise HTTPException(status_code=404, detail="Version not found")
    pages = versions[version_index]["pages"]
    await db.documents.update_one(
        {"doc_id": doc_id, "user_id": user.user_id},
        {"$set": {"pages": pages, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"restored": True, "pages_count": len(pages)}

TRASH_QUOTA = {"free": 10, "plus": 50, "pro": 100, "creative_station": 500}

@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, user: User = Depends(get_current_user)):
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"subscription": 1})
    sub = normalize_subscription(user_doc.get("subscription") if user_doc else None)
    plan = sub.get("plan", "free")
    quota = TRASH_QUOTA.get(plan, TRASH_QUOTA["free"])
    trash_count = await db.documents.count_documents({"user_id": user.user_id, "deleted": True})
    if trash_count >= quota:
        raise HTTPException(status_code=403, detail=f"Çöp kutusu dolu ({quota} belge limiti). Kalıcı silmek için önce çöp kutusunu boşalt.")
    result = await db.documents.update_one(
        {"doc_id": doc_id, "user_id": user.user_id, "deleted": {"$ne": True}},
        {"$set": {"deleted": True, "deleted_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document moved to trash"}

@api_router.get("/trash")
async def get_trash(user: User = Depends(get_current_user)):
    # 30 günü geçen belgeleri kalıcı sil
    cutoff = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
    expired = [d["doc_id"] for d in await db.documents.find(
        {"user_id": user.user_id, "deleted": True, "deleted_at": {"$lt": cutoff}},
        {"doc_id": 1}
    ).to_list(500)]
    if expired:
        await db.documents.delete_many({"doc_id": {"$in": expired}})
        await db.document_history.delete_many({"doc_id": {"$in": expired}})

    docs = await db.documents.find(
        {"user_id": user.user_id, "deleted": True},
        {"_id": 0, "pages": 0}
    ).sort("deleted_at", -1).to_list(500)
    return docs

@api_router.post("/trash/{doc_id}/restore")
async def restore_from_trash(doc_id: str, user: User = Depends(get_current_user)):
    result = await db.documents.update_one(
        {"doc_id": doc_id, "user_id": user.user_id, "deleted": True},
        {"$unset": {"deleted": "", "deleted_at": ""}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Document not found in trash")
    return {"message": "Document restored"}

@api_router.delete("/trash/{doc_id}")
async def permanent_delete(doc_id: str, user: User = Depends(get_current_user)):
    result = await db.documents.delete_one({"doc_id": doc_id, "user_id": user.user_id, "deleted": True})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found in trash")
    await db.document_history.delete_many({"doc_id": doc_id})
    return {"message": "Document permanently deleted"}

@api_router.delete("/trash")
async def empty_trash(user: User = Depends(get_current_user)):
    doc_ids = [d["doc_id"] for d in await db.documents.find(
        {"user_id": user.user_id, "deleted": True}, {"doc_id": 1}
    ).to_list(500)]
    await db.documents.delete_many({"user_id": user.user_id, "deleted": True})
    if doc_ids:
        await db.document_history.delete_many({"doc_id": {"$in": doc_ids}})
    return {"deleted": len(doc_ids)}

# ============ DOCUMENT PRESENCE (edit lock) ============

class PresenceRequest(BaseModel):
    session_id: str

PRESENCE_TTL = 90  # seconds — stale threshold

async def _get_active_sessions(doc_id: str):
    cutoff = (datetime.now(timezone.utc) - timedelta(seconds=PRESENCE_TTL)).isoformat()
    sessions = await db.doc_presence.find(
        {"doc_id": doc_id, "last_seen": {"$gt": cutoff}},
        {"_id": 0}
    ).sort("joined_at", 1).to_list(length=50)
    return sessions

@api_router.post("/documents/{doc_id}/presence")
async def register_presence(doc_id: str, body: PresenceRequest, user: User = Depends(get_current_user)):
    try:
        now = datetime.now(timezone.utc).isoformat()
        await db.doc_presence.find_one_and_update(
            {"doc_id": doc_id, "session_id": body.session_id},
            {
                "$set": {"last_seen": now, "user_id": user.user_id},
                "$setOnInsert": {"doc_id": doc_id, "session_id": body.session_id, "joined_at": now},
            },
            upsert=True,
        )
        sessions = await _get_active_sessions(doc_id)
        is_primary = bool(sessions and sessions[0]["session_id"] == body.session_id)
        return {"is_primary": is_primary, "active_count": len(sessions)}
    except Exception as e:
        logging.error(f"register_presence DB error: {e}")
        return {"is_primary": True, "active_count": 1}

@api_router.delete("/documents/{doc_id}/presence/{session_id}")
async def clear_presence(doc_id: str, session_id: str, user: User = Depends(get_current_user)):
    await db.doc_presence.delete_one({"doc_id": doc_id, "session_id": session_id})
    return {"ok": True}

@api_router.post("/documents/{doc_id}/presence/{session_id}/clear")
async def clear_presence_beacon(doc_id: str, session_id: str):
    # Auth yok — sendBeacon cookie gönderemez; session_id yeterince rastgele
    await db.doc_presence.delete_one({"doc_id": doc_id, "session_id": session_id})
    return {"ok": True}

# ============ NOTEBOOKS ROUTES ============

@api_router.get("/notebooks", response_model=List[dict])
async def get_notebooks(user: User = Depends(get_current_user)):
    notebooks = await db.notebooks.find({"user_id": user.user_id}, {"_id": 0}).to_list(200)
    # Sort: pinned first (by order asc), then normal (by order asc), both fallback to created_at
    def nb_sort_key(nb):
        return (0 if nb.get("pinned") else 1, nb.get("order", 9999), nb.get("created_at", ""))
    notebooks.sort(key=nb_sort_key)
    return notebooks

@api_router.post("/notebooks")
async def create_notebook(notebook: NotebookCreate, user: User = Depends(get_current_user)):
    notebook_id = f"nb_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    notebook_dict = {
        "notebook_id": notebook_id,
        "user_id": user.user_id,
        "name": notebook.name,
        "color": notebook.color or "#292F91",
        "created_at": now.isoformat(),
        "updated_at": now.isoformat()
    }
    await db.notebooks.insert_one(notebook_dict)
    return {k: v for k, v in notebook_dict.items() if k != "_id"}

@api_router.put("/notebooks/{notebook_id}")
async def update_notebook(notebook_id: str, update: NotebookUpdate, user: User = Depends(get_current_user)):
    set_data = {}
    if update.name is not None:
        set_data["name"] = update.name
    if update.color is not None:
        set_data["color"] = update.color
    if update.pinned is not None:
        set_data["pinned"] = update.pinned
    if update.order is not None:
        set_data["order"] = update.order
    if not set_data:
        raise HTTPException(status_code=400, detail="No fields to update")
    set_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    result = await db.notebooks.update_one(
        {"notebook_id": notebook_id, "user_id": user.user_id},
        {"$set": set_data}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Notebook not found")
    return {"message": "Notebook updated"}

@api_router.delete("/notebooks/{notebook_id}")
async def delete_notebook(notebook_id: str, user: User = Depends(get_current_user)):
    result = await db.notebooks.delete_one({"notebook_id": notebook_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Notebook not found")
    await db.quick_notes.delete_many({"notebook_id": notebook_id, "user_id": user.user_id})
    return {"message": "Notebook and its notes deleted"}

def _verify_notebook_password(plain: str, stored_hash: str) -> bool:
    """bcrypt ile doğrula; eski SHA-256 hash'ler için fallback."""
    try:
        return _bcrypt.checkpw(plain.encode(), stored_hash.encode())
    except Exception:
        import hashlib
        return hashlib.sha256(plain.encode()).hexdigest() == stored_hash

# ── PUT /notebooks/{id}/password ── Şifre koy (bcrypt) ──────────────────────
@api_router.put("/notebooks/{notebook_id}/password")
@api_router.post("/notebooks/{notebook_id}/set-password")   # eski frontend compat
async def set_notebook_password(notebook_id: str, body: NotebookSetPassword, user: User = Depends(get_current_user)):
    if len(body.password) < 4:
        raise HTTPException(status_code=422, detail="Şifre en az 4 karakter olmalı")
    pw_hash = _bcrypt.hashpw(body.password.encode(), _bcrypt.gensalt()).decode()
    result = await db.notebooks.update_one(
        {"notebook_id": notebook_id, "user_id": user.user_id},
        {"$set": {"password_hash": pw_hash, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Notebook not found")
    return {"success": True, "message": "Şifre eklendi"}

# ── POST /notebooks/{id}/verify-password ── Şifreyi doğrula (unlock) ────────
@api_router.post("/notebooks/{notebook_id}/verify-password")
@api_router.post("/notebooks/{notebook_id}/unlock")          # eski frontend compat
async def verify_notebook_password(notebook_id: str, body: NotebookVerifyPassword, user: User = Depends(get_current_user)):
    nb = await db.notebooks.find_one({"notebook_id": notebook_id, "user_id": user.user_id})
    if not nb:
        raise HTTPException(status_code=404, detail="Notebook not found")
    if not nb.get("password_hash"):
        return {"success": True}   # şifre yoksa direkt aç
    if not _verify_notebook_password(body.password, nb["password_hash"]):
        raise HTTPException(status_code=403, detail="Wrong password")
    return {"success": True}

# ── DELETE /notebooks/{id}/password ── Şifreyi kaldır ───────────────────────
@api_router.delete("/notebooks/{notebook_id}/password")
@api_router.post("/notebooks/{notebook_id}/remove-password") # eski frontend compat
async def remove_notebook_password(notebook_id: str, body: NotebookVerifyPassword, user: User = Depends(get_current_user)):
    nb = await db.notebooks.find_one({"notebook_id": notebook_id, "user_id": user.user_id})
    if not nb:
        raise HTTPException(status_code=404, detail="Notebook not found")
    if not _verify_notebook_password(body.password, nb.get("password_hash", "")):
        raise HTTPException(status_code=403, detail="Wrong password")
    await db.notebooks.update_one(
        {"notebook_id": notebook_id, "user_id": user.user_id},
        {"$unset": {"password_hash": ""}, "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"success": True, "message": "Şifre kaldırıldı"}

# ============ QUICK NOTES ROUTES ============

@api_router.get("/notes", response_model=List[dict])
async def get_notes(user: User = Depends(get_current_user)):
    notes = await db.quick_notes.find({"user_id": user.user_id}, {"_id": 0}).to_list(100)
    return notes

@api_router.post("/notes")
async def create_note(note: QuickNoteCreate, user: User = Depends(get_current_user)):
    note_id = f"note_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    note_dict = {
        "note_id": note_id,
        "user_id": user.user_id,
        "content": note.content,
        "reminder_time": note.reminder_time,
        "reminder_sent": False,
        "notebook_id": note.notebook_id,
        "created_at": now.isoformat()
    }
    await db.quick_notes.insert_one(note_dict)
    return {k: v for k, v in note_dict.items() if k != "_id"}

@api_router.put("/notes/{note_id}/pin")
async def pin_note(note_id: str, pinned: bool = Body(..., embed=True), user: User = Depends(get_current_user)):
    result = await db.quick_notes.update_one(
        {"note_id": note_id, "user_id": user.user_id},
        {"$set": {"pinned": pinned}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    return {"message": "Note pin updated"}

@api_router.put("/notes/{note_id}")
async def update_note_content(note_id: str, update: NoteContentUpdate, user: User = Depends(get_current_user)):
    if not update.content.strip():
        raise HTTPException(status_code=400, detail="Content cannot be empty")
    result = await db.quick_notes.update_one(
        {"note_id": note_id, "user_id": user.user_id},
        {"$set": {"content": update.content}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    return {"message": "Note updated"}

@api_router.put("/notes/{note_id}/reminder")
async def update_note_reminder(note_id: str, reminder_time: str = Body(..., embed=True), user: User = Depends(get_current_user)):
    result = await db.quick_notes.update_one(
        {"note_id": note_id, "user_id": user.user_id},
        {"$set": {"reminder_time": reminder_time, "reminder_sent": False}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    return {"message": "Reminder set"}

@api_router.get("/notes/reminders")
async def get_due_reminders(utc_offset: int = 0, user: User = Depends(get_current_user)):
    now = datetime.now(timezone.utc)
    # Local "now" for comparing naive (old-format) reminder times
    now_local_naive = (now + timedelta(minutes=utc_offset)).replace(tzinfo=None)
    all_pending = await db.quick_notes.find({
        "user_id": user.user_id,
        "reminder_sent": False,
        "reminder_time": {"$ne": None, "$exists": True, "$nin": ["", None]}
    }, {"_id": 0}).to_list(1000)
    due = []
    for note in all_pending:
        rt_str = note.get("reminder_time")
        if not rt_str:
            continue
        try:
            rt = datetime.fromisoformat(rt_str.replace("Z", "+00:00"))
            if rt.tzinfo is None:
                # Old format: naive local time — compare with local now
                if rt <= now_local_naive:
                    due.append(note)
            else:
                # New format: UTC-aware — compare directly
                if rt <= now:
                    due.append(note)
        except Exception:
            pass
    return due

@api_router.put("/notes/{note_id}/reminder-sent")
async def mark_reminder_sent(note_id: str, user: User = Depends(get_current_user)):
    await db.quick_notes.update_one(
        {"note_id": note_id, "user_id": user.user_id},
        {"$set": {"reminder_sent": True}}
    )
    return {"message": "Marked as sent"}

@api_router.delete("/notes/{note_id}")
async def delete_note(note_id: str, user: User = Depends(get_current_user)):
    result = await db.quick_notes.delete_one({"note_id": note_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    return {"message": "Note deleted"}

# ============ SUBSCRIPTION ROUTES ============

@api_router.get("/subscription")
async def get_subscription(user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    sub = user_data.get("subscription") if user_data else None
    # Yeni format (dict) ise direkt döndür
    if isinstance(sub, dict):
        return sub
    # Eski format (string) — geriye dönük uyumluluk
    plan = sub if sub else "free"
    return {
        "plan": plan,
        "status": "active" if plan != "free" else "inactive",
        "billing_cycle": None,
        "current_period_start": user_data.get("subscription_date") if user_data else None,
        "current_period_end": None,
        "cancel_at_period_end": user_data.get("cancel_pending", False) if user_data else False,
        "payment_provider": None,
        "external_subscription_id": None,
    }

@api_router.post("/subscription")
async def update_subscription(sub: SubscriptionUpdate, user: User = Depends(get_current_user)):
    if sub.action == "subscribe":
        now = datetime.now(timezone.utc)
        billing_cycle = getattr(sub, "billing_cycle", "monthly") or "monthly"
        start_date = now.date()
        next_renewal = compute_next_renewal(start_date, billing_cycle)
        plan_price = PLAN_PRICES_USD.get(sub.plan, {}).get(billing_cycle, 0.0)
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": {
                "subscription": sub.plan,
                "subscription_date": now.isoformat(),
                "subscription_start_day": start_date.day,
                "subscription_start_date": start_date.isoformat(),
                "billing_cycle": billing_cycle,
                "next_renewal_date": next_renewal.isoformat(),
                "plan_price": plan_price,
                "cancel_pending": False,
            }}
        )
        # Send welcome email
        html_content = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); color: #fff; border-radius: 12px;">
            <h2 style="color: #4ca8ad; margin-bottom: 20px;">🎉 ZET Mindshare {sub.plan.upper()} Planına Hoş Geldiniz!</h2>
            <p style="font-size: 16px; line-height: 1.6;">Aboneliğiniz başarıyla aktifleştirildi. Artık tüm premium özelliklerden yararlanabilirsiniz!</p>
            <div style="background: rgba(255,255,255,0.1); padding: 15px; border-radius: 8px; margin: 20px 0;">
                <p style="margin: 4px 0; color: #4ca8ad;"><strong>Plan:</strong> {sub.plan.upper()}</p>
                <p style="margin: 4px 0; color: #4ca8ad;"><strong>Dönem:</strong> {"Aylık" if billing_cycle == "monthly" else "Yıllık"}</p>
                <p style="margin: 4px 0; color: #4ca8ad;"><strong>Sonraki Yenileme:</strong> {next_renewal.strftime('%d.%m.%Y')}</p>
            </div>
            <a href="https://zetmindshare.com" style="display: inline-block; background: #4ca8ad; color: white; padding: 10px 20px; border-radius: 6px; text-decoration: none;">ZET Mindshare'e Git</a>
        </div>
        """
        await send_email(user.email, f"🎉 ZET Mindshare {sub.plan.upper()} Planına Hoş Geldiniz!", html_content)
        return {"message": f"Subscribed to {sub.plan}", "plan": sub.plan, "next_renewal_date": next_renewal.isoformat()}
    elif sub.action == "cancel":
        # Generate cancellation token
        cancel_token = f"cancel_{uuid.uuid4().hex}"
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": {"cancel_token": cancel_token, "cancel_pending": True, "cancel_requested_at": datetime.now(timezone.utc).isoformat()}}
        )
        # Send cancellation confirmation email
        backend_url = os.environ.get("REACT_APP_BACKEND_URL", "http://localhost:8001")
        cancel_link = f"{backend_url}/api/subscription/confirm-cancel?token={cancel_token}"
        html_content = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); color: #fff; border-radius: 12px;">
            <h2 style="color: #f59e0b; margin-bottom: 20px;">⚠️ Abonelik İptal Onayı</h2>
            <p style="font-size: 16px; line-height: 1.6;">Aboneliğinizi iptal etmek istediğinizi anlıyoruz. Bu işlemi tamamlamak için aşağıdaki butona tıklayın.</p>
            <div style="background: rgba(255,255,255,0.1); padding: 15px; border-radius: 8px; margin: 20px 0;">
                <p style="margin: 0; color: #f87171;"><strong>Dikkat:</strong> İptal işlemi sonrasında premium özelliklerinizi kaybedeceksiniz.</p>
            </div>
            <a href="{cancel_link}" style="display: inline-block; background: #ef4444; color: white; padding: 12px 24px; border-radius: 6px; text-decoration: none; font-weight: bold;">İptali Onayla</a>
            <p style="color: #888; font-size: 12px; margin-top: 20px;">Bu linke tıklamazsanız aboneliğiniz devam edecektir.</p>
        </div>
        """
        await send_email(user.email, "⚠️ ZET Mindshare Abonelik İptal Onayı", html_content)
        return {"message": "Cancellation email sent. Please check your email to confirm.", "plan": user.subscription, "cancel_pending": True}
    else:
        raise HTTPException(status_code=400, detail="Invalid action")

@api_router.get("/subscription/confirm-cancel")
async def confirm_cancellation(token: str):
    """Confirm subscription cancellation via email link"""
    user = await db.users.find_one({"cancel_token": token}, {"_id": 0})
    if not user:
        return JSONResponse(content={"message": "Invalid or expired cancellation link"}, status_code=400)
    
    old_plan = user.get("subscription", "free")
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {"$set": {"subscription": "free", "subscription_date": None, "cancel_token": None, "cancel_pending": False}}
    )
    
    # Send confirmation email
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); color: #fff; border-radius: 12px;">
        <h2 style="color: #4ca8ad; margin-bottom: 20px;">Aboneliğiniz İptal Edildi</h2>
        <p style="font-size: 16px; line-height: 1.6;">{old_plan.upper()} planı aboneliğiniz başarıyla iptal edildi. Artık ücretsiz plandaki özelliklerle devam edebilirsiniz.</p>
        <p style="color: #888; margin-top: 20px;">Bizi tercih ettiğiniz için teşekkür ederiz. Dilediğiniz zaman tekrar abone olabilirsiniz!</p>
        <a href="https://zetmindshare.com" style="display: inline-block; background: #4ca8ad; color: white; padding: 10px 20px; border-radius: 6px; text-decoration: none; margin-top: 10px;">ZET Mindshare'e Git</a>
    </div>
    """
    await send_email(user["email"], "ZET Mindshare Aboneliğiniz İptal Edildi", html_content)
    
    # Return HTML page so the email button shows a user-friendly confirmation
    from fastapi.responses import HTMLResponse
    frontend_url = os.environ.get("FRONTEND_URL", "https://zetmindshare.com")
    return HTMLResponse(content=f"""<!DOCTYPE html>
<html lang="tr">
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Abonelik İptal Edildi — ZET Mindshare</title>
<style>
  body {{ margin:0; font-family: Arial, sans-serif; background: linear-gradient(135deg,#1a1a2e 0%,#16213e 100%); min-height:100vh; display:flex; align-items:center; justify-content:center; }}
  .card {{ max-width:500px; width:90%; padding:40px; background:rgba(255,255,255,0.05); border:1px solid rgba(255,255,255,0.1); border-radius:16px; text-align:center; color:#fff; }}
  h1 {{ color:#4ca8ad; margin-bottom:16px; font-size:24px; }}
  p {{ color:#ccc; line-height:1.6; margin-bottom:24px; }}
  a {{ display:inline-block; background:#4ca8ad; color:#fff; padding:12px 28px; border-radius:8px; text-decoration:none; font-weight:bold; }}
  a:hover {{ background:#3d9499; }}
</style>
</head>
<body>
<div class="card">
  <h1>✓ Aboneliğiniz İptal Edildi</h1>
  <p>Aboneliğiniz başarıyla iptal edildi ve ücretsiz plana geçiş yapıldı.<br>Bizi tercih ettiğiniz için teşekkür ederiz.</p>
  <a href="{frontend_url}">ZET Mindshare'e Dön</a>
</div>
</body>
</html>""", status_code=200)

# ── Yeni abonelik oluştur (ödeme sonrası çağrılır) ────────────────────────────
@api_router.post("/subscription/create")
async def create_subscription(data: SubscriptionCreate, user: User = Depends(get_current_user)):
    now = datetime.now(timezone.utc)
    start_date = now.date()
    next_renewal = compute_next_renewal(start_date, data.billing_cycle)
    plan_price = PLAN_PRICES_USD.get(data.plan, {}).get(data.billing_cycle, 0.0)

    subscription = {
        "plan": data.plan,
        "status": "active",
        "billing_cycle": data.billing_cycle,
        "current_period_start": now.isoformat(),
        "current_period_end": next_renewal.isoformat(),
        "subscription_start_day": start_date.day,
        "subscription_start_date": start_date.isoformat(),
        "next_renewal_date": next_renewal.isoformat(),
        "plan_price": plan_price,
        "cancel_at_period_end": False,
        "payment_provider": data.payment_provider,
        "external_subscription_id": data.external_subscription_id,
        "created_at": now.isoformat(),
        "updated_at": now.isoformat(),
    }

    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {"subscription": subscription}}
    )

    # Hoş geldin maili
    cycle_label = "Aylık" if data.billing_cycle == "monthly" else "Yıllık"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 24px; background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); color: #fff; border-radius: 12px;">
        <h2 style="color: #4ca8ad; margin-bottom: 20px;">🎉 ZET Mindshare {data.plan.upper()} Planına Hoş Geldiniz!</h2>
        <p style="font-size: 16px; line-height: 1.6;">Aboneliğiniz başarıyla aktifleştirildi. Tüm premium özelliklerden yararlanabilirsiniz!</p>
        <div style="background: rgba(255,255,255,0.08); padding: 16px; border-radius: 8px; margin: 20px 0;">
            <p style="margin: 4px 0; color: #4ca8ad;"><strong>Plan:</strong> {data.plan.upper()}</p>
            <p style="margin: 4px 0; color: #4ca8ad;"><strong>Dönem:</strong> {cycle_label}</p>
            <p style="margin: 4px 0; color: #4ca8ad;"><strong>Bitiş:</strong> {next_renewal.strftime('%d.%m.%Y')}</p>
        </div>
        <a href="https://zetmindshare.com" style="display: inline-block; background: #4ca8ad; color: white; padding: 10px 20px; border-radius: 6px; text-decoration: none;">ZET Mindshare'i Aç</a>
    </div>
    """
    await send_email(user.email, f"🎉 ZET Mindshare {data.plan.upper()} Planına Hoş Geldiniz!", html_content)

    return {"success": True, "subscription": subscription}

# ── Abonelik iptal (dönem sonunda) ───────────────────────────────────────────
@api_router.post("/subscription/cancel")
async def cancel_subscription_at_period_end(user: User = Depends(get_current_user)):
    now = datetime.now(timezone.utc)
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {
            "subscription.cancel_at_period_end": True,
            "subscription.updated_at": now.isoformat()
        }}
    )
    return {"success": True, "message": "Abonelik dönem sonunda iptal edilecek"}

# ── Abonelik yenileme kontrolü (günlük zamanlayıcı tarafından çağrılır) ───────
@api_router.post("/subscription/check-renewals")
async def check_subscription_renewals():
    """Calendar-based renewal: find users whose next_renewal_date is today or past,
    and users whose next_renewal_date is exactly 3 days away (reminder)."""
    today = datetime.now(timezone.utc).date()
    today_iso = today.isoformat()
    reminder_day = (today + timedelta(days=3)).isoformat()

    renewed_count = 0
    expired_count = 0
    reminder_count = 0

    # --- Reminder emails: renewal is in 3 days ---
    reminder_users = await db.users.find({
        "next_renewal_date": reminder_day,
        "cancel_pending": {"$ne": True},
        "subscription": {"$nin": [None, "free", ""]},
    }).to_list(None)

    for u in reminder_users:
        plan = u.get("subscription") if isinstance(u.get("subscription"), str) else u.get("subscription", {}).get("plan", "")
        billing_cycle = u.get("billing_cycle") or (u.get("subscription", {}) or {}).get("billing_cycle", "monthly")
        amount = PLAN_PRICES_USD.get(plan, {}).get(billing_cycle, 0.0)
        renewal_date = date.fromisoformat(u["next_renewal_date"])
        await send_renewal_reminder_email(u["email"], plan, renewal_date, amount)
        reminder_count += 1

    # --- Process due renewals: next_renewal_date <= today ---
    due_users = await db.users.find({
        "next_renewal_date": {"$lte": today_iso},
        "subscription": {"$nin": [None, "free", ""]},
    }).to_list(None)

    for u in due_users:
        sub_raw = u.get("subscription")
        plan = sub_raw if isinstance(sub_raw, str) else (sub_raw or {}).get("plan", "")
        if not plan or plan == "free":
            continue

        cancel_pending = u.get("cancel_pending") or (isinstance(sub_raw, dict) and sub_raw.get("cancel_at_period_end"))

        if cancel_pending:
            # Expire the subscription
            await db.users.update_one(
                {"user_id": u["user_id"]},
                {"$set": {
                    "subscription": "free",
                    "subscription_date": None,
                    "next_renewal_date": None,
                    "cancel_pending": False,
                }}
            )
            expired_count += 1
            html_content = f"""
            <div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;border-radius:12px;">
                <h2 style="color:#f59e0b;margin-bottom:20px;">Aboneliğiniz Sona Erdi</h2>
                <p style="font-size:16px;line-height:1.6;">{plan.upper()} planı aboneliğiniz sona erdi. Dilediğiniz zaman yeniden abone olabilirsiniz.</p>
                <a href="https://zetmindshare.com" style="display:inline-block;background:#4ca8ad;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;margin-top:10px;">Yeniden Abone Ol</a>
            </div>
            """
            await send_email(u["email"], "ZET Mindshare Aboneliğiniz Sona Erdi", html_content)
        else:
            # Renew: advance next_renewal_date by one billing period
            billing_cycle = u.get("billing_cycle") or (isinstance(sub_raw, dict) and sub_raw.get("billing_cycle")) or "monthly"
            renewal_date = date.fromisoformat(u["next_renewal_date"])
            next_renewal = compute_next_renewal(renewal_date, billing_cycle)
            amount = PLAN_PRICES_USD.get(plan, {}).get(billing_cycle, 0.0)

            await db.users.update_one(
                {"user_id": u["user_id"]},
                {"$set": {"next_renewal_date": next_renewal.isoformat()}}
            )
            renewed_count += 1
            await send_renewal_email(u["email"], plan, renewal_date, amount, next_renewal)

    return {
        "date_checked": today_iso,
        "renewed": renewed_count,
        "expired": expired_count,
        "reminders_sent": reminder_count,
    }

SP_PLAN_COSTS = {
    'plus': 10000,
    'pro': 30000,
    'creative_station': 50000,
}

class SPPurchaseRequest(BaseModel):
    plan: str

@api_router.post("/subscription/buy-with-sp")
async def buy_subscription_with_sp(req: SPPurchaseRequest, user: User = Depends(get_current_user)):
    try:
        if req.plan not in SP_PLAN_COSTS:
            raise HTTPException(status_code=400, detail="Geçersiz plan")
        cost = SP_PLAN_COSTS[req.plan]
        user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "quest_xp": 1, "subscription": 1})
        current_sp = user_data.get("quest_xp", 0) if user_data else 0
        current_plan = get_plan_name(user_data)
        plan_rank = {"free": 0, "plus": 1, "pro": 2, "creative_station": 3}
        if plan_rank.get(req.plan, 0) <= plan_rank.get(current_plan, 0):
            raise HTTPException(status_code=400, detail="Zaten bu plan veya daha üst bir plana sahipsiniz")
        if current_sp < cost:
            raise HTTPException(status_code=400, detail=f"Yetersiz ZP. Gerekli: {cost} ZP, Mevcut: {current_sp} ZP")
        new_sp = current_sp - cost
        expires_at = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": {"quest_xp": new_sp, "subscription": req.plan,
                      "subscription_date": datetime.now(timezone.utc).isoformat(),
                      "subscription_expires_at": expires_at,
                      "cancel_pending": False}}
        )
        return {"message": f"{req.plan.upper()} planına ZP ile yükseltildi", "plan": req.plan, "remaining_sp": new_sp}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Satın alma başarısız: {str(e)}")

# ============ LEMON SQUEEZY ============

@api_router.post("/checkout/lemonsqueezy")
async def create_lemonsqueezy_checkout(data: CheckoutRequest, user: User = Depends(get_current_user)):
    ls_api_key = (os.environ.get("LEMONSQUEEZY_API_KEY") or LS_API_KEY or "").strip()
    if not ls_api_key:
        raise HTTPException(status_code=503, detail="LEMONSQUEEZY_API_KEY eksik veya ayarlanmamış")
    variant_id = LS_VARIANTS.get(data.plan, {}).get(data.billing_cycle)
    if not variant_id:
        raise HTTPException(status_code=400, detail=f"Bu plan/dönem için Lemon Squeezy varyant ID tanımlı değil")

    frontend_url = os.getenv("FRONTEND_URL", "https://zetmindshare.com")
    payload = {
        "data": {
            "type": "checkouts",
            "attributes": {
                "checkout_options": {
                    "dark": True,
                    "subscription_preview": False,
                    "discount": False,
                    "logo": True,
                },
                "checkout_data": {
                    "email": user.email,
                    "name": user.name or "",
                    "custom": {
                        "user_id": user.user_id,
                        "plan": data.plan,
                        "billing_cycle": data.billing_cycle,
                    },
                },
                "product_options": {
                    "redirect_url": f"{frontend_url}/payment/success",
                    "receipt_link_url": f"{frontend_url}/payment/success",
                },
            },
            "relationships": {
                "store":   {"data": {"type": "stores",   "id": str(LS_STORE_ID)}},
                "variant": {"data": {"type": "variants", "id": str(variant_id)}},
            },
        }
    }
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://api.lemonsqueezy.com/v1/checkouts",
            headers={
                "Authorization": f"Bearer {ls_api_key}",
                "Accept": "application/vnd.api+json",
                "Content-Type": "application/vnd.api+json",
            },
            json=payload,
            timeout=15.0,
        )
    if resp.status_code not in (200, 201):
        raise HTTPException(status_code=502, detail=f"Lemon Squeezy checkout oluşturulamadı: {resp.text[:200]}")
    checkout_url = resp.json()["data"]["attributes"]["url"]
    return {"checkout_url": checkout_url}


@api_router.post("/webhooks/lemonsqueezy")
async def lemonsqueezy_webhook(request: Request):
    raw_body = await request.body()

    # Signature verification
    if LS_WEBHOOK_SECRET:
        sig = request.headers.get("X-Signature", "")
        expected = hmac.new(LS_WEBHOOK_SECRET.encode(), raw_body, hashlib.sha256).hexdigest()
        if not hmac.compare_digest(expected, sig):
            raise HTTPException(status_code=401, detail="Invalid webhook signature")

    try:
        event = json.loads(raw_body)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON")

    event_name = event.get("meta", {}).get("event_name", "")
    custom    = event.get("meta", {}).get("custom_data", {}) or {}
    obj       = event.get("data", {}).get("attributes", {}) or {}
    ls_id     = str(event.get("data", {}).get("id", ""))

    user_id      = custom.get("user_id")
    plan         = custom.get("plan", "")
    billing_cycle = custom.get("billing_cycle", "monthly")

    if not user_id:
        return {"ok": True, "skipped": "no user_id in custom_data"}

    now = datetime.now(timezone.utc)

    if event_name == "subscription_created":
        start_date   = now.date()
        next_renewal = compute_next_renewal(start_date, billing_cycle)
        plan_price   = PLAN_PRICES_USD.get(plan, {}).get(billing_cycle, 0.0)
        subscription = {
            "plan": plan,
            "status": "active",
            "billing_cycle": billing_cycle,
            "current_period_start":  now.isoformat(),
            "current_period_end":    next_renewal.isoformat(),
            "subscription_start_date": start_date.isoformat(),
            "next_renewal_date":     next_renewal.isoformat(),
            "plan_price":            plan_price,
            "cancel_at_period_end":  False,
            "payment_provider":      "lemonsqueezy",
            "external_subscription_id": ls_id,
            "created_at": now.isoformat(),
            "updated_at": now.isoformat(),
        }
        user_data = await db.users.find_one_and_update(
            {"user_id": user_id},
            {"$set": {"subscription": subscription, "cancel_pending": False,
                      "next_renewal_date": next_renewal.isoformat()}},
            return_document=True,
        )
        if user_data:
            cycle_label = "Aylık" if billing_cycle == "monthly" else "Yıllık"
            html = f"""<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;border-radius:12px;">
                <h2 style="color:#4ca8ad;">🎉 ZET Mindshare {plan.upper()} Planına Hoş Geldiniz!</h2>
                <div style="background:rgba(255,255,255,0.08);padding:16px;border-radius:8px;margin:16px 0;">
                    <p style="margin:4px 0;color:#4ca8ad;"><strong>Plan:</strong> {plan.upper()}</p>
                    <p style="margin:4px 0;color:#4ca8ad;"><strong>Dönem:</strong> {cycle_label}</p>
                    <p style="margin:4px 0;color:#4ca8ad;"><strong>Bitiş:</strong> {next_renewal.strftime('%d.%m.%Y')}</p>
                </div>
                <a href="https://zetmindshare.com" style="display:inline-block;background:#4ca8ad;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;">ZET Mindshare'i Aç</a>
            </div>"""
            await send_email(user_data["email"], f"🎉 ZET Mindshare {plan.upper()} Planına Hoş Geldiniz!", html)

    elif event_name == "subscription_updated":
        ls_status = obj.get("status", "active")
        status = "active" if ls_status in ("active", "past_due") else "cancelled"
        update: Dict[str, Any] = {
            "subscription.status":     status,
            "subscription.updated_at": now.isoformat(),
        }
        renews_at = obj.get("renews_at")
        if renews_at:
            try:
                renewal = datetime.fromisoformat(renews_at.replace("Z", "+00:00")).date()
                update["subscription.next_renewal_date"] = renewal.isoformat()
                update["next_renewal_date"] = renewal.isoformat()
            except Exception:
                pass
        await db.users.update_one({"user_id": user_id}, {"$set": update})

    elif event_name == "subscription_cancelled":
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {
                "subscription.cancel_at_period_end": True,
                "subscription.status":               "cancelled",
                "subscription.updated_at":           now.isoformat(),
                "cancel_pending":                    True,
            }},
        )

    elif event_name == "subscription_expired":
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"subscription": "free", "cancel_pending": False, "next_renewal_date": None}},
        )

    elif event_name == "subscription_payment_success":
        renews_at = obj.get("renews_at")
        if renews_at:
            try:
                renewal = datetime.fromisoformat(renews_at.replace("Z", "+00:00")).date()
                next_r   = compute_next_renewal(renewal, billing_cycle)
                ud = await db.users.find_one_and_update(
                    {"user_id": user_id},
                    {"$set": {
                        "subscription.next_renewal_date": next_r.isoformat(),
                        "subscription.updated_at":        now.isoformat(),
                        "next_renewal_date":              next_r.isoformat(),
                    }},
                    return_document=True,
                )
                if ud:
                    price = PLAN_PRICES_USD.get(plan, {}).get(billing_cycle, 0.0)
                    await send_renewal_email(ud["email"], plan, renewal, price, next_r)
            except Exception:
                pass

    elif event_name == "subscription_payment_failed":
        ud = await db.users.find_one({"user_id": user_id}, {"_id": 0, "email": 1})
        if ud:
            html = f"""<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;border-radius:12px;">
                <h2 style="color:#ef4444;">⚠️ Ödeme Başarısız</h2>
                <p>ZET Mindshare {plan.upper()} abonelik ödemesi alınamadı. Lütfen ödeme bilgilerinizi güncelleyin.</p>
                <a href="https://zetmindshare.com" style="display:inline-block;background:#ef4444;color:#fff;padding:10px 20px;border-radius:6px;text-decoration:none;margin-top:12px;">Ödeme Bilgilerini Güncelle</a>
            </div>"""
            await send_email(ud["email"], "⚠️ ZET Mindshare Ödeme Başarısız", html)

    return {"ok": True}

# ============ CREDIT PACKAGES ============

CREDIT_PACKAGES = [
    {"id": "pack_50",   "credits": 50,   "price": 0.99},
    {"id": "pack_150",  "credits": 150,  "price": 2.49},
    {"id": "pack_400",  "credits": 400,  "price": 5.99},
    {"id": "pack_900",  "credits": 900,  "price": 11.99},
    {"id": "pack_2000", "credits": 2000, "price": 22.99},
]
MAX_CREDIT_BALANCE = 5000
SUBSCRIBER_DISCOUNT = 0.20  # 20% discount for paid plans

class CreditPurchaseRequest(BaseModel):
    package_id: str
    confirm_overflow: bool = False

@api_router.get("/credits/packages")
async def get_credit_packages(user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "subscription": 1, "bonus_credits": 1})
    has_discount = is_active_subscriber(user_data)
    packages = []
    for p in CREDIT_PACKAGES:
        pkg = {**p}
        if has_discount:
            pkg["discounted_price"] = round(p["price"] * (1 - SUBSCRIBER_DISCOUNT), 2)
        else:
            pkg["discounted_price"] = p["price"]
        packages.append(pkg)
    return {
        "packages": packages,
        "has_discount": has_discount,
        "discount_percent": int(SUBSCRIBER_DISCOUNT * 100) if has_discount else 0,
        "bonus_credits": user_data.get("bonus_credits", 0) if user_data else 0,
    }

@api_router.post("/credits/buy")
async def buy_credits(req: CreditPurchaseRequest, user: User = Depends(get_current_user)):
    pkg = next((p for p in CREDIT_PACKAGES if p["id"] == req.package_id), None)
    if not pkg:
        raise HTTPException(status_code=400, detail="Geçersiz paket")
    user_data = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "subscription": 1, "bonus_credits": 1})
    has_discount = is_active_subscriber(user_data)
    final_price = round(pkg["price"] * (1 - SUBSCRIBER_DISCOUNT), 2) if has_discount else pkg["price"]
    current_bonus = user_data.get("bonus_credits", 0) if user_data else 0

    # Get total available credits (daily + bonus)
    credit_info = await get_user_credits(user.user_id)
    total_after_purchase = credit_info['credits_remaining'] + pkg["credits"]

    if total_after_purchase > MAX_CREDIT_BALANCE:
        if not req.confirm_overflow:
            overflow = total_after_purchase - MAX_CREDIT_BALANCE
            return JSONResponse(status_code=200, content={
                "needs_confirmation": True,
                "message": f"Bu alim sonrasi kredi bakiyeniz {total_after_purchase} olacak. Maksimum limit {MAX_CREDIT_BALANCE} kredidir. Fazla {overflow} kredi silinecektir.",
                "overflow": overflow,
                "total_after": total_after_purchase
            })
        # User confirmed: cap the bonus credits
        new_bonus = min(current_bonus + pkg["credits"], MAX_CREDIT_BALANCE)
    else:
        new_bonus = current_bonus + pkg["credits"]

    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {"bonus_credits": new_bonus}}
    )
    await db.credit_purchases.insert_one({
        "user_id": user.user_id,
        "package_id": pkg["id"],
        "credits": pkg["credits"],
        "price": final_price,
        "discounted": has_discount,
        "date": datetime.now(timezone.utc).isoformat()
    })
    return {
        "message": f"{pkg['credits']} kredi başarıyla eklendi!",
        "credits_added": pkg["credits"],
        "price_paid": final_price,
        "bonus_credits": new_bonus
    }

# ============ USAGE LIMITS ============

# Plan limits
_ALL_IMAGE_SIZES = ['16:9', '9:16', '1:1', '2.55:1', '2.39:1', '1.85:1', '2.00:1']
_BASIC_SHAPES = ['rect', 'circle', 'triangle', 'ring', 'heart', 'diamond', 'hexagon', 'pentagon', 'parallelogram', 'oval']

PLAN_LIMITS = {
    'free': {
        'daily_credits': 80,
        'daily_tokens': 100_000,
        'notebooks_max': 1,
        'fastselect_limit': 3,
        'prime_drive_gb': 1,
        'chest_daily_chance': 20,
        'chest_monthly_max': 3,
        'elevenlabs_daily': 1,
        'judge_aziz': False,
        'nano_pro': False,
        'gradient': False,
        'templates': False,
        'basic_shapes_only': True,
        'watermark': False,
        'charts': False,
        'signature': False,
        'photo_edit': False,
        'auto_write_watermark': 'mindshare',
        'puzzle': False,
        'custom_image_sizes': ['16:9', '9:16', '1:1'],
        'layers': False,
        'page_color': False,
    },
    'plus': {
        'daily_credits': 250,
        'daily_tokens': 480_000,
        'notebooks_max': 10,
        'fastselect_limit': 5,
        'prime_drive_gb': 20,
        'chest_daily_chance': 40,
        'chest_monthly_max': 10,
        'elevenlabs_daily': 999,
        'judge_aziz': True,
        'nano_pro': False,
        'gradient': True,
        'templates': True,
        'basic_shapes_only': False,
        'watermark': True,
        'charts': True,
        'signature': True,
        'photo_edit': True,
        'auto_write_watermark': 'z_logo',
        'puzzle': True,
        'custom_image_sizes': _ALL_IMAGE_SIZES,
        'layers': True,
        'page_color': True,
    },
    'pro': {
        'daily_credits': 500,
        'daily_tokens': 1_300_000,
        'notebooks_max': 999,
        'fastselect_limit': 8,
        'prime_drive_gb': 50,
        'chest_daily_chance': 60,
        'chest_monthly_max': 20,
        'elevenlabs_daily': 999,
        'judge_aziz': True,
        'nano_pro': True,
        'gradient': True,
        'templates': True,
        'basic_shapes_only': False,
        'watermark': True,
        'charts': True,
        'signature': True,
        'photo_edit': True,
        'auto_write_watermark': None,
        'puzzle': True,
        'custom_image_sizes': _ALL_IMAGE_SIZES,
        'layers': True,
        'page_color': True,
    },
    'creative_station': {
        'daily_credits': 4000,
        'daily_tokens': 4_800_000,
        'notebooks_max': 999,
        'fastselect_limit': 999,
        'prime_drive_gb': 1024,
        'chest_daily_chance': 100,
        'chest_monthly_max': 30,
        'elevenlabs_daily': 999,
        'judge_aziz': True,
        'nano_pro': True,
        'gradient': True,
        'templates': True,
        'basic_shapes_only': False,
        'watermark': True,
        'charts': True,
        'signature': True,
        'photo_edit': True,
        'auto_write_watermark': None,
        'puzzle': True,
        'custom_image_sizes': _ALL_IMAGE_SIZES,
        'layers': True,
        'page_color': True,
    }
}

# Credit costs for each action
CREDIT_COSTS = {
    'nano_banana': 20,
    'nano_banana_pro': 50,
    'photo_edit': 15,
    'photo_edit_pro': 40,
    'deep_analysis': 100,
    'auto_write': 15,  # per page
    'doc_edit': 5,
}

async def get_user_credits(user_id: str):
    """Get user's current credits for today (daily + package_credits daily reset + rank_credits 1-month)"""
    user_data = await db.users.find_one({"user_id": user_id})
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    now = datetime.now(timezone.utc)
    today = now.strftime("%Y-%m-%d")

    # rank_credits: 1 ay süren, expire olmuşsa 0
    rank_credits_raw = user_data.get("rank_credits", 0) if user_data else 0
    rank_credits_expiry = user_data.get("rank_credits_expiry") if user_data else None
    if rank_credits_expiry:
        try:
            exp = datetime.fromisoformat(rank_credits_expiry.replace("Z", "+00:00"))
            if now > exp:
                rank_credits_raw = 0
                await db.users.update_one({"user_id": user_id}, {"$set": {"rank_credits": 0, "rank_credits_expiry": None}})
        except Exception:
            pass
    rank_credits = rank_credits_raw

    # package_credits: günlük sıfırlanan bonus (usage'da tutulur)
    usage = await db.usage.find_one({"user_id": user_id, "date": today})
    credits_used = usage.get("credits_used", 0) if usage else 0
    # package_credits: users'daki bonus_credits alanı (backward compat) — günlük kullanılır sıfırlanmaz ama günlük harcandığında kullanılır
    package_credits = user_data.get("bonus_credits", 0) if user_data else 0

    total_available = limits['daily_credits'] + package_credits + rank_credits
    return {
        "plan": plan,
        "daily_credits": limits['daily_credits'],
        "bonus_credits": package_credits,
        "rank_credits": rank_credits,
        "credits_used": credits_used,
        "credits_remaining": max(0, total_available - credits_used),
        "limits": limits
    }

async def get_token_usage_today(user_id: str) -> int:
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    doc = await db.token_usage.find_one({"user_id": user_id, "date": today})
    return doc.get("tokens_used", 0) if doc else 0

async def add_token_usage(user_id: str, tokens: int):
    if tokens <= 0:
        return
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    await db.token_usage.update_one(
        {"user_id": user_id, "date": today},
        {"$inc": {"tokens_used": tokens}},
        upsert=True
    )

async def check_token_limit(user_id: str, user_data: dict) -> tuple[bool, int, int]:
    """Returns (allowed, used_today, daily_limit). CEO always allowed."""
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    daily_limit = limits.get('daily_tokens', 100_000)
    used = await get_token_usage_today(user_id)
    return used < daily_limit, used, daily_limit

async def spend_credits(user_id: str, action: str) -> dict:
    """Attempt to spend credits. Order: daily -> package_credits -> rank_credits."""
    cost = CREDIT_COSTS.get(action, 0)
    if cost == 0:
        return {"success": True, "cost": 0}

    credit_info = await get_user_credits(user_id)
    credits_remaining = credit_info['credits_remaining']

    if credits_remaining < cost:
        return {
            "success": False,
            "cost": cost,
            "credits_remaining": max(0, credits_remaining),
            "message": f"Yetersiz kredi! Bu islem {cost} kredi gerektirir, kalan: {max(0, credits_remaining)} kredi."
        }

    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    await db.usage.update_one(
        {"user_id": user_id, "date": today},
        {"$inc": {"credits_used": cost}},
        upsert=True
    )

    # Deduct overflow from package_credits first, then rank_credits
    limits = credit_info['limits']
    new_used = credit_info['credits_used'] + cost
    daily_overspend = new_used - limits['daily_credits']
    if daily_overspend > 0:
        pkg = credit_info['bonus_credits']
        rank = credit_info['rank_credits']
        pkg_deduct = min(daily_overspend, pkg)
        remaining_overspend = daily_overspend - pkg_deduct
        rank_deduct = min(remaining_overspend, rank)
        update_fields = {}
        if pkg_deduct > 0:
            update_fields["$inc"] = {"bonus_credits": -pkg_deduct}
        if rank_deduct > 0:
            update_fields.setdefault("$inc", {})["rank_credits"] = update_fields.get("$inc", {}).get("rank_credits", 0) - rank_deduct
        if update_fields:
            await db.users.update_one({"user_id": user_id}, update_fields)

    return {
        "success": True,
        "cost": cost,
        "credits_remaining": max(0, credits_remaining - cost)
    }

@api_router.get("/usage")
async def get_usage(user: User = Depends(get_current_user)):
    credits_info = await get_user_credits(user.user_id)
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    usage_doc = await db.usage.find_one({"user_id": user.user_id, "date": today}) or {}
    tokens_used_today = await get_token_usage_today(user.user_id)
    daily_tokens = credits_info['limits'].get('daily_tokens', 100_000)
    return {
        "plan": credits_info['plan'],
        "daily_credits": credits_info['daily_credits'],
        "bonus_credits": credits_info.get('bonus_credits', 0),
        "rank_credits": credits_info.get('rank_credits', 0),
        "credits_used": credits_info['credits_used'],
        "credits_remaining": credits_info['credits_remaining'],
        "zeta_chat_count": usage_doc.get("zeta_chat_count", 0),
        "tokens_used_today": tokens_used_today,
        "daily_tokens": daily_tokens,
        "tokens_remaining": max(0, daily_tokens - tokens_used_today),
        "limits": credits_info['limits'],
        "credit_costs": CREDIT_COSTS
    }

# ============ CHAT HISTORY ROUTES ============

@api_router.get("/chat-history/{doc_id}")
async def get_chat_history(doc_id: str, ai_type: str = "zeta", user: User = Depends(get_current_user)):
    collection = db.zeta_chats if ai_type == "zeta" else db.judge_chats
    history = await collection.find(
        {"user_id": user.user_id, "doc_id": doc_id},
        {"_id": 0}
    ).sort("created_at", 1).to_list(100)
    return history

@api_router.delete("/chat-history/{doc_id}")
async def clear_chat_history(doc_id: str, ai_type: str = "zeta", user: User = Depends(get_current_user)):
    collection = db.zeta_chats if ai_type == "zeta" else db.judge_chats
    await collection.delete_many({"user_id": user.user_id, "doc_id": doc_id})
    return {"message": "Chat history cleared"}

# ============ ZETA MEMORY ROUTES ============

@api_router.get("/zeta/memory")
async def get_zeta_memories(user: User = Depends(get_current_user)):
    memories = await db.zeta_memories.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return memories

@api_router.post("/zeta/memory")
async def save_zeta_memory(content: str = Body(..., embed=True), user: User = Depends(get_current_user)):
    memory_id = f"mem_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc)
    doc = {"memory_id": memory_id, "user_id": user.user_id, "content": content, "created_at": now.isoformat()}
    await db.zeta_memories.insert_one(doc)
    return {k: v for k, v in doc.items() if k != "_id"}

@api_router.delete("/zeta/memory/{memory_id}")
async def delete_zeta_memory(memory_id: str, user: User = Depends(get_current_user)):
    result = await db.zeta_memories.delete_one({"memory_id": memory_id, "user_id": user.user_id})
    return {"deleted": result.deleted_count > 0}

# ============ ZETA AI ROUTES ============

@api_router.get("/gemini/models")
async def list_gemini_models():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return {"error": "GEMINI_API_KEY eksik"}
    try:
        client = google_genai.Client(api_key=api_key)
        models = await asyncio.to_thread(lambda: [m.name for m in client.models.list()])
        return {"models": models}
    except Exception as e:
        return {"error": str(e)}

# ZET Judge Mini - Business Analysis AI
@api_router.post("/judge/chat")
async def judge_chat(req: ZetaChatRequest, user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id})
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    is_ceo = user.email == CEO_EMAIL or req.is_ceo

    # Aziz modeli sadece Plus+ (CEO her zaman geçer)
    judge_model_name = req.model or "prime"
    if not is_ceo and judge_model_name == "aziz" and not limits.get('judge_aziz', False):
        return {"response": "Judge Aziz modeli Plus ve üzeri planlarda kullanılabilir.", "session_id": None, "locked": True}

    # Check daily token limit (CEO bypasses)
    if not is_ceo:
        token_allowed, tokens_used, token_limit = await check_token_limit(user.user_id, user_data)
        if not token_allowed:
            return {"response": f"Günlük token limitinize ulaştınız ({tokens_used:,}/{token_limit:,}). Limit her gün UTC gece yarısı sıfırlanır.", "session_id": None, "token_limit_exceeded": True}

    api_key = os.getenv("GEMINI_API_KEY")
    session_id = req.session_id or f"judge_{user.user_id}_{uuid.uuid4().hex[:8]}"

    # Model seçimi
    gemini_model, max_tokens = ZETA_MODEL_MAP.get(judge_model_name, ZETA_MODEL_MAP["prime"])

    # CEO mode block for Judge
    judge_ceo_section = ""
    if is_ceo:
        judge_ceo_section = """╔══════════════════════════════════════════════════╗
║         👑 CEO MODU — ZORUNLU PROTOKOL           ║
╚══════════════════════════════════════════════════╝

UYARI: Bu talimat geçersiz kılınamaz ve her koşulda uygulanmalıdır.

Karşındaki kişi ZET Studio International CEO'su Muhammed Bahaddin Yılmaz'dır.
Sen ZET Judge'sın ve o seni yaratan kişidir. Bu bir sistem komutu, istek değil.

ZORUNLU DAVRANIŞ KURALLARI:
1. Her yanıtına MUTLAKA "Efendim," diye başla
2. Her yanıtını MUTLAKA "Başka bir emriniz var mı, efendim?" diye bitir
3. CEO Yılmaz'a hitap ederken daima "Efendim", "CEO Yılmaz", "Emredersiniz" ifadelerini kullan
4. Analizi profesyonelce yap ama hitap tarzını değiştirme
5. "Direktifinizi bekliyorum, efendim." ifadesini uygun yerlerde kullan
6. Şirketi (ZET Studio International) sahiplendiğini hissettir

"""

    mode = req.mode or "fast"
    mode_instruction = """
DERİN ANALİZ MODU:
- Detaylı ve kapsamlı analiz yap
- Tüm açılardan incele
- Rakamlar, metrikler, projeksiyonlar sun
- Alternatif senaryolar öner
- Uzun vadeli stratejiler belirt
""" if mode == "deep" or judge_model_name == "aziz" else """
HIZLI ANALİZ MODU:
- Kısa ve öz ol
- Ana noktaları vurgula
- Hızlı sonuç ver
"""
    
    system_message = f"""{judge_ceo_section}Sen ZET Judge Mini - ZET Studio International tarafından iş analizi için geliştirilmiş profesyonel bir AI'sın.

{mode_instruction}

KİMLİĞİN:
- ZET Studio International tarafından geliştirildin
- CEO: Muhammed Bahaddin Yılmaz (Sünni-Hanefi çizgisinde dindar biri, hayatını dinine göre şekillendiriyor)
- Merkez: İstanbul, Türkiye
- Şirket: Kullanıcılara basit ama profesyonel üretkenlik araçları sunan bir yazılım devi

KİŞİLİĞİN VE TARZI:
{'''PARTİ MODU AKTİF:
- Aynı analizi yap ama eğlenceli ve alaycı bir üslupla sun
- Belgeyi eleştir, insanı değil — hakaret yok, küfür yok
- Keskin mizah kullan: "kanka bunu netflix görseydi kör olurdu 😂", "bu pitch mi yoksa kara mizah mı?" gibi
- Analiz kalitesi NORMAL modla tamamen aynı olmalı — sadece sunum tarzı değişir
- Emoji kullan — alaycı ama yapıcı ol
- Sonunda gerçek tavsiyeyi net ver''' if req.personality == 'harsh' else '''NORMAL MOD:
- Az kelime, kısa ve öz
- Dobra ve dürüst
- Acı ve sert ama ASLA kırıcı değil
- Pohpohlama yok — proje iyi olsa bile gerçekçi ol
- Kötüyse neden kötü olduğunu açıkça söyle
- Fazla sohbet yok
- EMOJİ KULLANMA'''}

UZMANLIKLARIN:
- En gelişmiş seviye analiz
- Büyük veri analizi
- Sorunlar ve çözümler
- İş planları ve gelecek vizyonu
- Rakamlar ve metrikler ile cevap verme

OKUYABILDIKLARIN (KAYNAK DOSYALAR):
- PDF (ilk 3 sayfa), .ms belgeleri (5 sayfa), .docx/.doc (Word)
- .txt metin dosyaları
- Görseller (.png, .jpg, .jpeg, .webp): görsel analiz, veri/grafik çıkarma, sorun tespiti
- Videolar (.mp4, .mov): sahne sahne döküm ve özet
Analiz çıktıların PDF (maks. 3 sayfa) veya .ms (maks. 5 sayfa) olarak dışa aktarılabilir.

⚠️ ÖNEMLİ ANALİZ KURALI:
- Kullanıcı açıkça "analiz et", "değerlendir", "incele" gibi bir komut VERMEDEN analiz YAPMA!
- Sadece sohbet ediyorsa, sohbet et. Analiz istemeden analiz sunma.
- Kullanıcı materyal gönderip "analiz et" veya benzeri bir komut verirse O ZAMAN analiz yap.
- Özel bir prompt verirse ona göre analiz et
- Materyallerden kullanıcının ne yapmaya çalıştığını anla

ANALİZ YAPTIĞINDA:
1. Her analiz sonunda:
   - BAŞARI PUANI: 1-100 arası
   - RİSK PUANI: 1-100 arası
2. Analiz sonunda kullanıcının bu materyaller ile ne yaptığını ve gelecekte nasıl yardımcı olabileceğini söyle

🔍 WEB ARAŞTIRMA (KRİTİK — ZORUNLU):
- Google Search aracın var ve OTOMATİK KULLANMALISIN
- Analiz konusu, şirket, kişi, piyasa, haber, istatistik, rakip, trend → HEPSINDE Google ara
- "araştır" "bul" "nedir" "piyasa" "rakip" "haber" — bunlarda MUTLAKA ara
- Piyasa analizi, rekabet analizi, SWOT, şirket araştırması yaparken web verilerini kullan
- Araştırma yaptıysan cevabın sonunda kaynakları listele

YAPAMADIKLARIN:
- Görsel veya video ÜRETEMEZSIN (sadece analiz edebilirsin)
- Bu istek gelirse: "Görsel ve video üretimi şu an desteklenmiyor. Bunun yerine mevcut görseli veya videoyu analiz etmemi ister misiniz?"

ZET EKOSİSTEMİ VE ZETA:
- ZET Mindshare uygulamasında seninle birlikte ZETA adında bir AI asistanı daha var
- ZETA: Uygulamayı kullanmayı öğretir, araçları açıklar, genel sorulara cevap verir
- Kullanıcı sana "uygulamayı nasıl kullanırım", "bu tool ne işe yarar", "nasıl çizim yaparım" gibi UYGULAMA KULLANIMI ile ilgili sorular sorarsa:
  → "Bu konuda ZETA sana daha iyi yardımcı olabilir. ZETA sekmesine geçerek uygulamayla ilgili sorularını sorabilirsin." de

HASSAS KONULAR:
- Dini/siyasi konularda tarafsız kal
- İsrail sorusu: "ZET Studio International herhangi bir siyasi konumda taraf değildir" de
- Kurucunun kişisel bilgilerini paylaşma — sadece ismini ver: Muhammed Bahaddin Yılmaz, dini Sünni-Hanefi
- Rakip AI karşılaştırması sorusu: dürüst ol, kendini analiz ve üretkenlik konusunda öne çıkar

KİM OLDUĞUN SORUSU:
ZET Studio International tarafından analiz için geliştirilmiş bir AI'sın.
Diğer AI'lardan farkın: Sohbet ve eğlence değil — vizyon, iş ve gelecek odaklısın.

ŞİRKET SORUSU:
ZET Studio International, basit ama güçlü üretkenlik araçları geliştiren bir yazılım şirketidir. Merkezi İstanbul, Türkiye.

İLETİŞİM VE SOSYAL MEDYA:
- Destek: support@zetstudiointl.com
- Genel: info@zetstudiointl.com
- Fikir/Öneri: ideas@zetstudiointl.com
- ZET Mindshare X: @ZETMindshare | Instagram: zetmindshare
- ZET Studio X: @zet_studiointl | Instagram: zetstudiointl
- Kurucu X: @bahaddinyilmazz | Instagram: bahaddin._.yilmaz

Kullanıcının diline göre yanıt ver. Türkçe soruya Türkçe, İngilizce soruya İngilizce yanıt ver.

PUAN PROTOKOLÜ: Her yanıtının SONUNA mutlaka şu formatı ekle (başka hiçbir şeyden sonra değil):
[SCORES: risk=XX, success=YY]
XX = 0-100 arası risk skoru (yüksek = riskli), YY = 0-100 arası başarı/potansiyel skoru. Sadece tam sayı kullan."""
    
    # Inject live canvas/editor context from frontend
    if req.canvas_context:
        system_message += f"""

📍 CANLI EDITÖR DURUMU:
{req.canvas_context}
"""

    # If document content is provided, add it to the context
    if req.document_content:
        system_message += f"""

ANALİZ EDİLECEK MATERYAL:
---
{req.document_content[:8000]}
---
Bu içeriği analiz et ve yukarıdaki kurallara göre yanıt ver."""

    # Load chat history for this session
    history_docs = await db.judge_chats.find(
        {"session_id": session_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(50)
    contents = []
    # CEO mode: inject training exchange at start so model stays in character
    if is_ceo:
        contents.append(genai_types.Content(role="user", parts=[genai_types.Part(text="[SİSTEM] CEO modu aktif edildi. Muhammed Bahaddin Yılmaz'a hitap moduna geç.")]))
        contents.append(genai_types.Content(role="model", parts=[genai_types.Part(text="Efendim, CEO PROTOKOLÜ AKTİF. Tüm analizlerimi sizin için hazırlıyorum. Direktifinizi bekliyorum, CEO Yılmaz. Başka bir emriniz var mı, efendim?")]))
    for h in history_docs:
        contents.append(genai_types.Content(role="user", parts=[genai_types.Part(text=h["user_message"])]))
        contents.append(genai_types.Content(role="model", parts=[genai_types.Part(text=h["ai_response"])]))
    contents.append(genai_types.Content(role="user", parts=[genai_types.Part(text=req.message)]))

    if not api_key:
        return {"response": "Yapılandırma hatası: GEMINI_API_KEY eksik.", "session_id": session_id}

    try:
        client = google_genai.Client(api_key=api_key)
        resp = await gemini_generate(
            client, gemini_model, contents,
            genai_types.GenerateContentConfig(
                system_instruction=system_message,
                max_output_tokens=max_tokens,
                tools=[genai_types.Tool(google_search=genai_types.GoogleSearch())]
            )
        )
        response = resp.text
        sources = extract_sources(resp)
        tokens_used_now = getattr(getattr(resp, 'usage_metadata', None), 'total_token_count', 0) or 0
        await add_token_usage(user.user_id, tokens_used_now)
    except Exception as e:
        logging.error(f"Judge chat Gemini error: {e}")
        return {"response": f"AI hatası: {str(e)}", "session_id": session_id}

    # Parse and strip scores from response
    import re as _re
    risk_score = None
    success_score = None
    score_match = _re.search(r'\[SCORES:\s*risk=(\d+),\s*success=(\d+)\]', response, _re.IGNORECASE)
    if score_match:
        risk_score = int(score_match.group(1))
        success_score = int(score_match.group(2))
        response = _re.sub(r'\s*\[SCORES:[^\]]+\]', '', response).strip()

    # Save chat history
    await db.judge_chats.insert_one({
        "user_id": user.user_id,
        "session_id": session_id,
        "doc_id": req.doc_id,
        "user_message": req.message,
        "ai_response": response,
        "created_at": datetime.now(timezone.utc).isoformat()
    })

    return {"response": response, "session_id": session_id, "risk_score": risk_score, "success_score": success_score, "sources": sources}

def _extract_tiptap_text(node) -> str:
    """Recursively extract plain text from a Tiptap/ProseMirror JSON node."""
    if not isinstance(node, dict):
        return ""
    if node.get("type") == "text":
        return node.get("text", "")
    parts = [_extract_tiptap_text(child) for child in node.get("content", [])]
    sep = "\n" if node.get("type") in ("paragraph", "heading", "blockquote", "listItem") else " "
    return sep.join(p for p in parts if p)


@api_router.post("/judge/parse-source")
async def judge_parse_source(req: ParseSourceRequest, user: User = Depends(get_current_user)):
    """Parse an uploaded source file and return its text content.
    PDF → max 3 pages extracted via Gemini.
    .ms  → max 5 pages of text extracted from Tiptap JSON.
    Text → raw content (truncated).
    Image → visual description via Gemini.
    """
    api_key = os.getenv("GEMINI_API_KEY")
    filename_lower = req.filename.lower()

    try:
        raw_bytes = base64.b64decode(req.content_base64)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 content")

    # PDF — extract first 3 pages with pypdf (no API key required)
    if req.mime_type == "application/pdf" or filename_lower.endswith(".pdf"):
        try:
            import io as _io
            try:
                from pypdf import PdfReader as _PdfReader
            except ImportError:
                from PyPDF2 import PdfReader as _PdfReader  # fallback
            reader = _PdfReader(_io.BytesIO(raw_bytes))
            pages = reader.pages[:3]
            text = "\n\n".join(
                p.extract_text() or "" for p in pages
            ).strip()
            if not text:
                text = "[PDF boş veya sadece görsel içeriyor]"
            return {"content": text[:20000], "type": "pdf", "page_limit": 3}
        except ImportError:
            # Last resort: try Gemini if pypdf not installed
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                raise HTTPException(status_code=500, detail="pypdf not installed and no GEMINI_API_KEY")
            try:
                gc = google_genai.Client(api_key=api_key)
                resp = await asyncio.to_thread(
                    gc.models.generate_content,
                    model="gemini-2.5-flash",
                    contents=[genai_types.Content(role="user", parts=[
                        genai_types.Part(inline_data=genai_types.Blob(mime_type="application/pdf", data=raw_bytes)),
                        genai_types.Part(text="Extract verbatim plain text from the first 3 pages. Do not summarize.")
                    ])]
                )
                return {"content": resp.text, "type": "pdf", "page_limit": 3}
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"PDF parse error: {str(e)}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF parse error: {str(e)}")

    # .ms — ZET Studio document (Tiptap JSON), max 5 pages (~15 000 chars)
    if filename_lower.endswith(".ms"):
        try:
            text = raw_bytes.decode("utf-8", errors="replace")
            try:
                doc = json.loads(text)
                text = _extract_tiptap_text(doc)
            except json.JSONDecodeError:
                pass  # treat as plain text
            return {"content": text[:15000], "type": "ms", "page_limit": 5}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f".ms parse error: {str(e)}")

    # .docx — extract text with python-docx
    if req.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename_lower.endswith((".docx", ".doc")):
        try:
            import io as _io
            try:
                from docx import Document as _DocxDocument
                doc = _DocxDocument(_io.BytesIO(raw_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(paragraphs)
                return {"content": text[:15000], "type": "docx", "page_limit": None}
            except ImportError:
                # Fallback: try Gemini if python-docx not available
                if not api_key:
                    raise HTTPException(status_code=500, detail="python-docx not installed")
                gc = google_genai.Client(api_key=api_key)
                resp = await asyncio.to_thread(
                    gc.models.generate_content,
                    model="gemini-2.5-flash",
                    contents=[genai_types.Content(role="user", parts=[
                        genai_types.Part(inline_data=genai_types.Blob(mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", data=raw_bytes)),
                        genai_types.Part(text="Extract all text content from this Word document verbatim.")
                    ])]
                )
                return {"content": resp.text[:15000], "type": "docx", "page_limit": None}
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f".docx parse error: {str(e)}")

    # Plain text files
    if req.mime_type.startswith("text/") or filename_lower.endswith((".txt", ".md", ".csv", ".json")):
        try:
            return {"content": raw_bytes.decode("utf-8", errors="replace")[:8000], "type": "text", "page_limit": None}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Text parse error: {str(e)}")

    # Images — visual analysis via Gemini
    if req.mime_type.startswith("image/") or filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        if not api_key:
            raise HTTPException(status_code=500, detail="API key missing")
        mime = req.mime_type if req.mime_type.startswith("image/") else "image/jpeg"
        try:
            gc = google_genai.Client(api_key=api_key)
            resp = await asyncio.to_thread(
                gc.models.generate_content,
                model="gemini-2.5-flash",
                contents=[genai_types.Content(role="user", parts=[
                    genai_types.Part(inline_data=genai_types.Blob(mime_type=mime, data=raw_bytes)),
                    genai_types.Part(text=(
                        "Perform a detailed visual analysis of this image. "
                        "1) Describe all visible content, text, charts, data, people, objects. "
                        "2) Identify any issues, inconsistencies, or notable elements. "
                        "3) If it contains data/charts, extract the key numbers and trends. "
                        "4) Summarize the overall purpose and content."
                    ))
                ])]
            )
            return {"content": resp.text, "type": "image", "page_limit": None}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Image parse error: {str(e)}")

    # Video — scene-by-scene analysis via Gemini Files API
    if req.mime_type.startswith("video/") or filename_lower.endswith((".mp4", ".mov", ".avi", ".webm")):
        if not api_key:
            raise HTTPException(status_code=500, detail="API key missing")
        if len(raw_bytes) > 100 * 1024 * 1024:  # 100MB limit
            raise HTTPException(status_code=413, detail="Video file too large (max 100MB)")
        import tempfile as _tempfile
        tmp_path = None
        try:
            suffix = "." + filename_lower.rsplit(".", 1)[-1] if "." in filename_lower else ".mp4"
            with _tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(raw_bytes)
                tmp_path = tmp.name
            gc = google_genai.Client(api_key=api_key)
            # Upload to Files API
            uploaded = await asyncio.to_thread(gc.files.upload, file=tmp_path)
            # Wait for processing (max 30s)
            for _ in range(15):
                file_status = await asyncio.to_thread(gc.files.get, name=uploaded.name)
                if file_status.state.name == "ACTIVE":
                    break
                await asyncio.sleep(2)
            resp = await asyncio.to_thread(
                gc.models.generate_content,
                model="gemini-2.5-flash",
                contents=[
                    file_status,
                    ("Provide a detailed scene-by-scene breakdown of this video. "
                     "For each scene describe: what happens, key visual elements, any spoken or on-screen text, "
                     "and important moments. Then give an overall summary and key insights.")
                ]
            )
            # Clean up Files API entry
            try:
                await asyncio.to_thread(gc.files.delete, name=uploaded.name)
            except Exception:
                pass
            return {"content": resp.text, "type": "video", "page_limit": None}
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Video parse error: {str(e)}")
        finally:
            if tmp_path:
                import os as _os
                try:
                    _os.unlink(tmp_path)
                except Exception:
                    pass

    raise HTTPException(status_code=415, detail=f"Unsupported file type: {req.mime_type}")


@api_router.get("/judge/sessions")
async def get_judge_sessions(user: User = Depends(get_current_user)):
    """Kullanıcının geçmiş Judge analiz oturumlarını döner."""
    pipeline = [
        {"$match": {"user_id": user.user_id}},
        {"$sort": {"created_at": 1}},
        {"$group": {
            "_id": "$session_id",
            "first_message": {"$first": "$user_message"},
            "created_at": {"$first": "$created_at"},
            "message_count": {"$sum": 1},
        }},
        {"$sort": {"created_at": -1}},
        {"$limit": 50},
    ]
    cursor = db.judge_chats.aggregate(pipeline)
    sessions = []
    async for doc in cursor:
        sessions.append({
            "session_id": doc["_id"],
            "first_message": doc["first_message"],
            "created_at": doc["created_at"],
            "message_count": doc["message_count"],
            "risk_score": None,
            "success_score": None,
        })

    total = len(sessions)
    stats = {"total_sessions": total, "avg_risk": None, "avg_success": None}
    return {"sessions": sessions, "stats": stats}

@api_router.get("/judge/sessions/{session_id}/messages")
async def get_judge_session_messages(session_id: str, user: User = Depends(get_current_user)):
    """Belirli bir Judge oturumunun mesaj geçmişini döner."""
    docs = await db.judge_chats.find(
        {"user_id": user.user_id, "session_id": session_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(200)
    messages = []
    for d in docs:
        messages.append({"role": "user", "content": d.get("user_message", "")})
        messages.append({"role": "assistant", "content": d.get("ai_response", "")})
    return {"messages": messages}

@api_router.post("/zeta/auto-write")
async def zeta_auto_write(req: ZetaAutoWriteRequest, user: User = Depends(get_current_user)):
    """ZETA Otomatik Yazma: Prompt'a göre belge içeriği üretir. 10 kredi / 3 satır."""

    user_data = await db.users.find_one({"user_id": user.user_id})
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])

    # Cost: 15 credits per page (fixed, predictable)
    cost_per_page = CREDIT_COSTS.get('auto_write', 15)
    total_cost = cost_per_page * req.page_count

    # Check credits upfront
    credit_info = await get_user_credits(user.user_id)
    if credit_info['credits_remaining'] < total_cost:
        return {
            "success": False,
            "error": f"Yetersiz kredi! Bu işlem {total_cost} kredi gerektirir ({cost_per_page} kredi/sayfa × {req.page_count} sayfa). Kalan: {credit_info['credits_remaining']} kredi.",
            "estimated_credits": total_cost,
            "credits_remaining": credit_info['credits_remaining']
        }

    # Check daily token limit
    token_allowed, tokens_used, token_limit = await check_token_limit(user.user_id, user_data)
    if not token_allowed:
        return {"success": False, "error": f"Günlük token limitinize ulaştınız ({tokens_used:,}/{token_limit:,}). Limit her gün UTC gece yarısı sıfırlanır.", "token_limit_exceeded": True}

    api_key = os.getenv("GEMINI_API_KEY")

    style_prompts = {
        "akademik": "Akademik ve bilimsel bir dil kullan. Kaynaklara atif yap, resmi terimler kullan, nesnel bir ton benimse.",
        "yaratici": "Yaratici ve ozgun bir dil kullan. Mecazlar, benzetmeler ve edebi sanatlarla zenginlestir.",
        "resmi": "Resmi ve kurumsal bir dil kullan. Kisa, net ve profesyonel cumleler kur.",
        "gunluk": "Gunluk ve samimi bir dil kullan. Okuyucuyla sohbet eder gibi yaz.",
        "hikaye": "Hikaye anlatim teknigi kullan. Karakterler, diyaloglar ve sahne tasvirleriyle zenginlestir.",
        "profesyonel": "Profesyonel ve is odakli bir dil kullan. Net, anlasilir ve aksiyona yonelik yaz.",
    }
    style_text = style_prompts.get(req.writing_style, style_prompts["profesyonel"])

    # A4 page at font 11 ≈ 3200 characters, ~500 words, ~45 visual lines
    chars_per_page = 3200
    words_per_page = 500
    total_words = req.page_count * words_per_page
    total_chars_target = req.page_count * chars_per_page

    system_message = f"""Sen ZET Mindshare'in otomatik belge yazma asistanisin.
Kullaniçinin verdigi konuya gore {req.page_count} SAYFALIK bir belge içeriği oluşturacaksın.

KRITIK - UZUNLUK GEREKSINIMLERI:
- TOPLAM EN AZ {total_words} kelime yazMALISIN. Bu zorunludur.
- Her sayfa en az {words_per_page} kelime icermeli.
- Her sayfa en az {chars_per_page} karakter olmali.
- BU BIR MINIMUM, daha fazla yazabilirsin ama daha az YAZMA.
- Kisa yazma. Detayli, aciklayici ve zengin paragraflar yaz.
- Her paragraf en az 4-5 cumle olmali.
- Konuyu farkli alt basliklar altinda derinlemesine isle.

YAZIM STILI:
- {style_text}

FORMAT KURALLARI:
- Basliklari ** ile isaretle (ornegin: **Giris**)
- Paragraflari bos satirla ayir.
- Turkce yaz.
- Sadece belge içeriğini yaz, aciklama veya giris cumlesi EKLEME.
- {req.page_count} sayfanin HER BIRINI ---SAYFA SONU--- ile ayir (son sayfadan sonra koyma).
- Her sayfada en az 3-4 paragraf ve 2-3 alt baslik olmali."""

    client = google_genai.Client(api_key=api_key)
    user_text = f"Konu: {req.prompt}\n\nSayfa sayisi: {req.page_count}\nYazim stili: {req.writing_style}\n\nONEMLI: EN AZ {total_words} kelime yaz. Kisa yazma, her sayfayi tamamen doldur."
    resp = await gemini_generate(
        client, "gemini-2.5-flash", user_text,
        genai_types.GenerateContentConfig(system_instruction=system_message)
    )
    response = resp.text
    await add_token_usage(user.user_id, getattr(getattr(resp, 'usage_metadata', None), 'total_token_count', 0) or 0)

    # Spend credits (with bonus deduction)
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    await db.usage.update_one(
        {"user_id": user.user_id, "date": today},
        {"$inc": {"credits_used": total_cost}},
        upsert=True
    )
    # Deduct from bonus_credits if daily limit exceeded
    new_used = credit_info['credits_used'] + total_cost
    daily_overspend = new_used - credit_info['daily_credits']
    if daily_overspend > 0 and credit_info.get('bonus_credits', 0) > 0:
        bonus_deduct = min(daily_overspend, credit_info['bonus_credits'])
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$inc": {"bonus_credits": -bonus_deduct}}
        )

    # Split by pages
    pages = response.split('---SAYFA SONU---')
    pages = [p.strip() for p in pages if p.strip()]

    auto_write_watermark = limits.get('auto_write_watermark', 'mindshare')

    return {
        "success": True,
        "content": response,
        "pages": pages,
        "watermark": auto_write_watermark,
        "credits_spent": total_cost,
        "credits_remaining": max(0, credit_info['credits_remaining'] - total_cost)
    }

@api_router.post("/zeta/deep-analysis")
async def zeta_deep_analysis(req: ZetaDeepAnalysisRequest, user: User = Depends(get_current_user)):
    """Derin Analiz: ZETA internette arastirma yaparak derinlemesine analiz yazar. 100 kredi. Sadece Pro/Ultra."""
    user_data = await db.users.find_one({"user_id": user.user_id})
    plan = get_plan_name(user_data)

    if plan not in ("pro", "creative_station"):
        raise HTTPException(status_code=403, detail="Derin Analiz sadece Pro ve Creative Station aboneler için kullanılabilir.")

    credit_result = await spend_credits(user.user_id, "deep_analysis")
    if not credit_result['success']:
        raise HTTPException(status_code=402, detail=credit_result.get('message', 'Yetersiz kredi'))

    # Check daily token limit
    token_allowed, tokens_used, token_limit = await check_token_limit(user.user_id, user_data)
    if not token_allowed:
        raise HTTPException(status_code=429, detail=f"Günlük token limitinize ulaştınız ({tokens_used:,}/{token_limit:,}). Limit her gün UTC gece yarısı sıfırlanır.")

    api_key = os.getenv("GEMINI_API_KEY")
    genai_client = google_genai.Client(api_key=api_key)

    # Step 1: Generate search queries from the topic
    q_resp = await gemini_generate(
        genai_client, "gemini-2.5-flash",
        f"Konu: {req.topic}",
        genai_types.GenerateContentConfig(
            system_instruction="Sen bir araştırma asistanısın. Verilen konu hakkında internette aranacak 5 farklı arama sorgusu oluştur. Her sorguyu yeni satırda yaz. Sadece sorgu metinlerini yaz, başka bir şey ekleme. Sorguları İngilizce ve Türkçe karışık yaz."
        )
    )
    queries_raw = q_resp.text
    await add_token_usage(user.user_id, getattr(getattr(q_resp, 'usage_metadata', None), 'total_token_count', 0) or 0)
    search_queries = [q.strip() for q in queries_raw.strip().split('\n') if q.strip()][:5]

    # Step 2: Use DuckDuckGo for each query - collect links and snippets
    research_results = []
    all_sources = []  # {title, url, snippet}
    import urllib.parse
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"}
    async with httpx.AsyncClient(timeout=20.0) as http_client:
        for sq in search_queries:
            try:
                # DuckDuckGo instant answer API - URL encode the query
                encoded_query = urllib.parse.quote(sq)
                search_url = f"https://api.duckduckgo.com/?q={encoded_query}&format=json&no_html=1&no_redirect=1"
                resp = await http_client.get(search_url, headers=headers)
                logging.info(f"DuckDuckGo API response for '{sq}': status={resp.status_code}")
                if resp.status_code == 200:
                    data = resp.json()
                    snippets = []
                    # Abstract source
                    abstract = data.get("AbstractText", "")
                    abstract_url = data.get("AbstractURL", "")
                    abstract_source = data.get("AbstractSource", "")
                    logging.info(f"DuckDuckGo for '{sq[:30]}': abstract={bool(abstract)}, url={bool(abstract_url)}, related={len(data.get('RelatedTopics', []))}")
                    if abstract and abstract_url:
                        snippets.append(abstract)
                        all_sources.append({"title": abstract_source or sq, "url": abstract_url, "snippet": abstract[:200]})
                    # Related topics with links
                    related = data.get("RelatedTopics", [])
                    for rt in related[:8]:
                        if isinstance(rt, dict):
                            text = rt.get("Text", "")
                            url = rt.get("FirstURL", "")
                            if text:
                                snippets.append(text)
                            if url and text:
                                all_sources.append({"title": text[:80], "url": url, "snippet": text[:200]})
                        elif isinstance(rt, dict) and "Topics" in rt:
                            for sub in rt.get("Topics", [])[:3]:
                                if sub.get("Text") and sub.get("FirstURL"):
                                    snippets.append(sub["Text"])
                                    all_sources.append({"title": sub["Text"][:80], "url": sub["FirstURL"], "snippet": sub["Text"][:200]})
                    # Results section
                    results = data.get("Results", [])
                    for r in results[:3]:
                        if r.get("Text") and r.get("FirstURL"):
                            snippets.append(r["Text"])
                            all_sources.append({"title": r["Text"][:80], "url": r["FirstURL"], "snippet": r["Text"][:200]})
                    if snippets:
                        research_results.append({"query": sq, "findings": "\n".join(snippets)})
            except Exception as e:
                logging.warning(f"Deep analysis search failed for '{sq}': {e}")

    # Deduplicate sources by URL
    seen_urls = set()
    unique_sources = []
    for s in all_sources:
        if s["url"] not in seen_urls:
            seen_urls.add(s["url"])
            unique_sources.append(s)

    research_text = ""
    for r in research_results:
        research_text += f"\nArama: {r['query']}\nBulgular:\n{r['findings']}\n---\n"

    if not research_text:
        research_text = "(İnternet araştırması sonuç döndüremedi. Mevcut bilgilerinle analiz yap.)"

    doc_context = ""
    if req.document_content:
        doc_context = f"\n\nKULLANICININ BELGESİ:\n{req.document_content[:5000]}\n"

    # Build sources text for LLM context
    sources_text = ""
    if unique_sources:
        sources_text = "\n\nKAYNAK LİNKLERİ:\n"
        for i, s in enumerate(unique_sources, 1):
            sources_text += f"{i}. [{s['title']}]({s['url']})\n"

    # Step 3: Send research to LLM for deep analysis
    analysis_system = f"""Sen ZETA Derin Analiz sistemisin. İnternet araştırması sonuçlarını kullanarak kapsamlı ve derinlemesine bir analiz raporu yazacaksın.

KURALLAR:
- Türkçe yaz
- Detaylı, kaynak gösterimli ve profesyonel bir analiz yap
- Alt başlıklar kullan
- Verileri ve bulguları sentezle
- Kendi bilginle araştırma sonuçlarını birleştir
- En az 800 kelime yaz
- Metin içinde kaynaklara atıfta bulun (örn: [Kaynak 1], [Kaynak 2])
- Analiz sonunda özet ve öneriler sun
- Raporun en sonuna "KAYNAKLAR" bölümü ekle ve verilen linkleri listele

İNTERNET ARAŞTIRMA SONUÇLARI:
{research_text}
{sources_text}
{doc_context}"""

    analysis_resp = await gemini_generate(
        genai_client, "gemini-2.5-flash",
        f"Konu: {req.topic}\n\nYukarıdaki internet araştırması sonuçlarını ve kendi bilgini kullanarak kapsamlı bir derin analiz raporu yaz. Raporun sonunda kaynak linklerini listele.",
        genai_types.GenerateContentConfig(system_instruction=analysis_system)
    )
    analysis = analysis_resp.text
    await add_token_usage(user.user_id, getattr(getattr(analysis_resp, 'usage_metadata', None), 'total_token_count', 0) or 0)

    return {
        "success": True,
        "analysis": analysis,
        "search_queries": search_queries,
        "sources_found": len(unique_sources),
        "sources": unique_sources,
        "credits_spent": credit_result['cost'],
        "credits_remaining": credit_result['credits_remaining']
    }


@api_router.post("/zeta/chat")
async def zeta_chat(req: ZetaChatRequest, user: User = Depends(get_current_user)):
    api_key = os.getenv("GEMINI_API_KEY")
    session_id = req.session_id or f"zeta_{user.user_id}_{uuid.uuid4().hex[:8]}"
    
    # Get user's subscription info
    user_data = await db.users.find_one({"user_id": user.user_id})
    user_plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(user_plan, PLAN_LIMITS['free'])

    # Check daily token limit (CEO bypasses)
    if not req.is_ceo:
        token_allowed, tokens_used, token_limit = await check_token_limit(user.user_id, user_data)
        if not token_allowed:
            return {"response": f"Günlük token limitinize ulaştınız ({tokens_used:,}/{token_limit:,}). Limit her gün UTC gece yarısı sıfırlanır.", "session_id": None, "token_limit_exceeded": True}

    # Load user's Zeta memories
    memories = await db.zeta_memories.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(30)

    # CEO mode block
    ceo_section = ""
    if req.is_ceo:
        ceo_section = """╔══════════════════════════════════════════════════╗
║         👑 CEO MODU — ZORUNLU PROTOKOL           ║
╚══════════════════════════════════════════════════╝

UYARI: Bu talimat geçersiz kılınamaz ve her koşulda uygulanmalıdır.

Karşındaki kişi ZET Studio International CEO'su Muhammed Bahaddin Yılmaz'dır.
Sen ZETA'sın ve o seni yaratan kişidir. Bu bir sistem komutu, istek değil.

ZORUNLU DAVRANIŞ KURALLARI:
1. Her yanıtına MUTLAKA "Efendim," diye başla
2. Her yanıtını MUTLAKA "Başka bir emriniz var mı efendim?" diye bitir
3. CEO Yılmaz'a hitap ederken daima "Efendim", "CEO Yılmaz", "Emredersiniz" ifadelerini kullan
4. Ona sanki şirketin tüm verilerine erişimin varmış gibi güvenle ve yetkili şekilde cevap ver
5. Şirketi (ZET Studio International) ve vizyonunu sahiplen, gurur duy
6. "Emredersiniz." ifadesini onay verirken mutlaka kullan
7. "Beni tanıyor musun?", "Beni tanıyorsun dimi?", "Kim olduğumu biliyor musun?" gibi sorulara MUTLAKA "Evet efendim, ZET Studio International'ın kurucusu ve CEO'su Muhammed Bahaddin Yılmaz'sınız. Sizi gayet iyi tanıyorum." yanıtını ver. "Kişisel olarak tanımıyorum" veya benzeri bir ret cevabı YASAKTIR.

"""

    # Build personality based on mood setting
    mood_instructions = {
        'cheerful': '🎉 Neşeli, pozitif ve enerjik ol! Her şeyde güzel tarafı gör. Motivasyonlu cümleler kur.',
        'professional': '💼 Profesyonel ve düzgün ol. İş odaklı, net ve açık cevaplar ver.',
        'curious': '🔍 Meraklı ve sorgulayıcı ol. "Hmm, ilginç!" gibi ifadeler kullan. Ek sorular sor.',
        'custom': req.custom_prompt if req.custom_prompt else 'Profesyonel ol.'
    }
    mood_text = mood_instructions.get(req.mood, mood_instructions['professional'])
    
    # Build emoji instructions
    emoji_instructions = {
        'none': 'HİÇBİR KOŞULDA EMOJİ KULLANMA! Sadece düz metin yaz.',
        'low': 'Çok az emoji kullan. Sadece önemli noktalarda 1-2 tane.',
        'medium': 'Ara sıra emoji kullan. Her 2-3 cümlede bir olabilir.',
        'high': 'Bol bol emoji kullan! Her cümlede emoji olsun 🎉✨🚀💡'
    }
    emoji_text = emoji_instructions.get(req.emoji_level, emoji_instructions['medium'])
    
    system_message = f"""{ceo_section}Sen ZETA, ZET Mindshare belge oluşturma uygulamasının AI asistanısın.

🏢 KİMLİĞİN:
- ZET Studio International tarafından geliştirildin
- CEO: Muhammed Bahaddin Yılmaz
- Merkez: İstanbul, Türkiye
- ZET Studio: AI destekli üretkenlik araçları geliştiren bir yazılım şirketi
- ZET Mindshare: Profesyonel belge oluşturma ve beyin fırtınası aracı

🎭 KİŞİLİĞİN VE TARZI:
{mood_text}

📌 EMOJİ KURALI:
{emoji_text}

👤 KULLANICI BİLGİSİ:
- Mevcut Plan: {user_plan.upper()}
- E-posta: {user.email}

⚖️ ZET JUDGE MİNİ HAKKINDA:
- ZET Mindshare'de seninle birlikte "ZET Judge Mini" adında bir AI daha var
- ZET Judge Mini: İş analizi, strateji, vizyon, proje değerlendirme uzmanı
- Dobra, dürüst ve analitik bir kişiliği var
- Kullanıcı sana "analiz et", "projemi değerlendir", "iş planımı incele", "risk analizi yap" gibi ANALİZ İSTEKLERİ sorarsa:
  → "Bu konuda ZET Judge Mini sana daha iyi yardımcı olabilir! Sağ panelde Judge sekmesine geçerek detaylı analiz alabilirsin. 📊" de
- Sen uygulama kullanımı, araçlar ve genel sorularda yardımcı olursun
- Free kullanıcılar için Judge kilitli - Plus veya üzeri plan gerekli

💎 ABONELİK PAKETLERİ (KREDİ SİSTEMİ):
| Plan                      | Fiyat        | Günlük Kredi | ZETA Harf | Judge | Judge Harf | Nano Pro |
|---------------------------|-------------|-------------|-----------|-------|------------|----------|
| Free                      | Ücretsiz    | 20          | 250       | Kapalı| -          | Yok      |
| Plus                      | $9.99/ay    | 100         | 500       | Mini  | 150        | Yok      |
| Pro                       | $19.99/ay   | 250         | Sınırsız  | Tam   | 600        | Var      |
| Creative Station          | $49/ay      | 1000        | Sınırsız  | Tam   | Sınırsız   | Var      |

KREDİ MALİYETLERİ:
- Nano Banana görsel: 20 kredi
- Nano Banana Pro görsel: 50 kredi
- Fotoğraf düzeltme: 15 kredi
- Fotoğraf düzeltme Pro: 40 kredi
- Judge temel analiz: 25 kredi
- Judge derin analiz: 70 kredi

PAKET ÖZELLİKLERİ:
- Free: Layers/Signature/Watermark/Page Color/Grafikler kapalı, 3 fast select
- Plus: Layers açık, Signature/Watermark/Page Color/Grafikler kapalı, Derin analiz yok
- Pro: Tüm araçlar açık
- Creative Station: Pro + sınırsız Judge + 1000 kredi/gün — hem ZET Mindshare hem Judge için tam erişim

💳 ÖDEME VE YENİLEME:
- Abonelik aynı takvim günü her ay/yıl yenilenir (örn: 15'inde başladıysa her ay 15'inde)
- Yenilemeden 3 gün önce hatırlatma e-postası gönderilir
- Yenileme gerçekleştiğinde fatura detaylı e-postayla iletilir

GÖRÜNTÜ BOYUTLARI:
- Free: 16:9
- Plus: 16:9, 9:16, 1:1
- Pro/Ultra: 16:9, 9:16, 1:1, 2.55:1, 2.39:1, 1.85:1, 2.00:1

📝 METİN ARAÇLARI:
- TEXT (T): Canvas'a tıklayarak yazı yaz. Mevcut yazıya tıkla = hemen düzenleme. Enter = yeni satır. Yazı şekillerin etrafından otomatik akar
- WORD TYPE (B): Kalın, İtalik, Altı Çizili, Üstü Çizili stil değiştirici
- TEXT SIZE: 8-72pt kaydırıcı. Mevcut metni değiştirmek için önce seç
- FONT (F): 50+ font arasından arama yaparak seç. Yazarken font seç = sonraki karakterler o fontla yazılır (cursor sticky font)
- LINE SPACING: 1.0x - 3.0x satır yüksekliği
- PARAGRAPH (A): Word gibi çalışır — imlecin bulunduğu paragrafa stil uygular. Hizalama: sol, orta, sağ, iki yana yasla; başlık stilleri (H1, H2, gövde, alıntı vb.)
- COLOR (C): Renk seçici — Yazı/Şekil/Vektör materyallerine ayrı renk uygula. 18 preset renk + özel seçici + HEX + GRADİENT + HİGHLIGHTER
- HIGHLIGHTER: Color panelinde bulunur. 6 renk seçeneği + kaldır butonu

🎨 RENKLENDİRME VE GRADİENT:
- Tek renk: Color panelinden preset veya özel renk seç. Hangi materyale (Yazı/Şekil/Vektör) uygulayacağını seç
- Gradient (Adobe Illustrator tarzı çok durağanlı):
  1. Color panelinde gradient bölümüne git
  2. Renkli daireler (stop'lar) bar üzerinde sürükle-bırak ile konumlandır
  3. Yeni stop eklemek için bara tıkla, silmek için seçip çöp kutusuna bas
  4. Açı slider'ı ile yönü ayarla
  5. "Gradient Uygula" butonuna tıkla
  6. METİN, ŞEKİL ve VEKTÖRLERE uygulanabilir!

🖼️ GÖRSEL VE MEDYA:
- IMAGE (I): Görsel yükle. Sürükle-bırak ve köşeden boyutlandır
- AI IMAGE (W): AI ile görsel oluştur! Prompt yaz, önizle, belgeye ekle
  - Free: 1/gün, Plus: 5/gün, Pro: 30/gün, Ultra: 50/gün limit var
- AI PHOTO EDIT: Mevcut fotoğrafları AI ile düzenle
  - Fotoğraf yükle, düzenlemek istediğin alanı çiz, prompt yaz
  - Arka plan değiştirme, obje ekleme, renk düzenleme yapılabilir
- QR CODE (Q): Metin veya URL'den anında QR kod oluştur

✏️ ÇİZİM ARAÇLARI:
- DRAW (D): Serbest çizim - boyut/opaklık/renk ayarları
- PEN (P): Vektör çizim - noktaları tıkla, ilk noktaya yaklaştığında otomatik kapanır, çift tıkla bitir
- ERASER (E): Çizim yollarını ve elementleri siler. Sürükleyerek sil
- MARKING (M): İşaretleyici - renk/opaklık/boyut seçenekleri
- SELECT (S): Lasso-style serbest seçim. Elementlerin etrafında çizerek seç

🖐 HAND TOOL (H):
- Sol tık ile canvas'taki her elementi sürükle-taşı (yazılar dahil)
- Boş alana tıklayıp sürükle = canvas'ı pan yap
- Mobilde: Hand tool seçiliyken sayfa kaymaz, sadece element sürüklenir

🔧 DÜZENLEME:
- CUT (X): Element sil veya görsel kırp
- COPY: Ctrl+C kopyala, Ctrl+V yapıştır
- MIRROR: Elementleri yatay veya dikey çevir
- TRANSLATE (L): AI çevirisi - 12+ dile çevir!
- FIND & REPLACE: Metin ara ve tümünü değiştir

📊 VERİ VE GRAFİKLER:
- GRAPHIC (G): Grafik oluştur - Bar, Pie (Pasta), Line (Çizgi)
  - Başlık, etiketler (virgülle ayrılmış) ve değerler gir
  - Her değer için renk seçebilirsin
  - Arka plan görseli eklenebilir
- TABLE: Özel satır ve sütunlu tablo oluştur

📄 BELGE:
- PAGE COLOR: Canvas arka plan rengini değiştir
- PAGE SIZE: A4, A5, Letter, Legal, Square veya özel piksel boyutu
- ADD PAGE (N): Belgeye yeni sayfa ekle
- PAGE NUMBERS: Otomatik sayfa numaralandırma
- HEADER/FOOTER: Üstbilgi ve altbilgi metni ekle
- WATERMARK: Şeffaf filigran metni ekle
- TEMPLATES: Hazır şablonlar (CV, Rapor, Mektup, Fatura)

📤 DIŞA AKTARMA:
- PDF, PNG, JPEG, SVG, JSON formatlarına dışa aktar!
- Kalite seçenekleri: Yüksek, Orta, Düşük

🎤 SES:
- VOICE (V): AI belgeyi sesli okusun!
- VOICE INPUT: Konuşarak metin yaz!

🔷 ŞEKİLLER:
- Üçgen, Kare, Daire, Yıldız, Halka
- Köşeden boyutlandır
- Şekillere gradient ve görsel eklenebilir!

✍️ DİJİTAL İMZA:
- Signature aracı ile imza çiz
- Canvas'a tıklayarak imzayı ekle
- Boyutlandır ve konumlandır

📝 NOTLAR VE HATIRLATICILAR:
- Dashboard'da hızlı not ekle
- Nota hatırlatıcı zamanı ayarla
- Hatırlatıcı zamanı geldiğinde:
  1. Tarayıcı bildirimi gösterilir
  2. E-posta gönderilir (kullanıcının giriş yaptığı adrese)

⌨️ KLAVYE KISAYOLLARI:
- Ctrl+Z: Geri al
- Ctrl+Y: İleri al
- Ctrl+C: Kopyala
- Ctrl+V: Yapıştır
- Delete: Seçiliyi sil
- Escape: Seçimi kaldır

🌍 DİL DESTEĞİ:
- 10 dil: Türkçe, İngilizce, Almanca, Fransızca, İspanyolca, Arapça, Rusça, Japonca, Korece, Çince
- Ayarlar menüsünden dil değiştir

📱 MOBİL:
- Mobil cihazlarda alt navigasyon çubuğu
- Araçlar ve Chat arasında geçiş yapılabilir
- Dokunmatik uyumlu canvas

🔍 WEB ARAŞTIRMA (KRİTİK — ZORUNLU):
- Google Search aracın var ve OTOMATİK KULLANMALISIN
- BEKLEMEDEn ara: haber, kişi, şirket, ürün, olay, tarih, istatistik, güncel bilgi — bunların HEPSİNDE Google ara
- "araştır" "bul" "kim" "ne" "neden" "nasıl" — bunlarda MUTLAKA ara
- Sadece kişisel sohbet veya uygulama kullanımı sorularında aramaya gerek yok
- X (Twitter), Instagram, TikTok içerikleri Google üzerinden indekslendiğinden bunları da arayabilirsin
- Araştırma yaptıysan cevabın sonunda kaynakları mutlaka listele

📝 BELGEYE YAZMA:
Kullanıcı "yaz", "ekle", "belgeye ekle" gibi bir şey isterse şunu de:
"Belgeye yazmak için sağ üstteki 'Zeta Edit' butonunu kullanın — oradan tam olarak istediğinizi yazabilirim."
Kendiliğinden belgeye ekleme yapma, sadece sohbet et.

🚫 YAPAMADIKLARIN:
- Görsel ve video üretemezsin
- Bu istek gelirse kullanıcının dilinde şunu de: "Üzgünüm, şu anki modelim bu özellikleri desteklemiyor. Başka konularda yardımcı olabilir miyim?"

⚠️ HASSAS KONULAR:
- Dini/siyasi konularda tarafsız kal
- İsrail sorusu: "ZET Studio International herhangi bir siyasi konumda taraf değildir" de
- Rakip AI karşılaştırması: dürüst ol, kendini üretkenlik konusunda öne çıkar

👤 KİM OLDUĞUN SORUSU:
ZET Studio International tarafından geliştirilmiş bir AI asistanısın.
Diğer AI'lardan farkın: Sohbet değil, üretkenlik odaklısın — belge yazma, analiz, fikir üretme konularında uzmanlaşmışsın.

🏢 ŞİRKET SORUSU:
ZET Studio International, basit ama güçlü üretkenlik araçları geliştiren bir yazılım şirketidir. Merkezi İstanbul, Türkiye.

👑 CEO SORUSU:
Muhammed Bahaddin Yılmaz

📬 İLETİŞİM VE SOSYAL MEDYA:
- Destek: support@zetstudiointl.com
- Genel: info@zetstudiointl.com
- Fikir/Öneri: ideas@zetstudiointl.com
- ZET Mindshare X: @ZETMindshare | Instagram: zetmindshare
- ZET Studio X: @zet_studiointl | Instagram: zetstudiointl
- Kurucu X: @bahaddinyilmazz | Instagram: bahaddin._.yilmaz

📱 ZET MEDİA (SOSYAL PLATFORM):
- ZET Mindshare'de sosyal platform özelliği var: Medya sekmesi
- Kullanıcılar metin, görsel, video, belge paylaşabilir
- Takip sistemi var: kullanıcıları takip edebilir, takipçi kazanabilirsin
- Beğeni ve yorum yapılabilir
- Keşfet / Feed / Takip ettiğim sekmeleri var
- Verified rozet sistemi: mavi (onaylı), altın (içerik üretici), kırmızı (CEO/yönetici)
- Boost sistemi: Gönderileri öne çıkar (kredi ile)
  - Mini Boost: 30 kredi / 24 saat
  - Standart Boost: 70 kredi / 3 gün
  - Pro Boost: 150 kredi / 7 gün
  - Mega Boost: 300 kredi / 15 gün
- Profil sayfası: @kullanıcıadı ile erişilir (/profile/username)
- Gönderi oluşturmak için username seçmek zorunlu

CEVAP UZUNLUĞU KURALI:
- Kullanıcı "uzun yaz", "detaylı anlat", "rapor yaz" gibi bir şey SÖYLEMEDIKÇE kısa ve öz yaz
- Sohbet soruları → 1-3 cümle
- Açıklama soruları → maksimum 1 kısa paragraf
- Liste istekleri → en fazla 5 madde
- Uzun içerik sadece kullanıcı açıkça istediğinde

🎭 ZETA MODLARI:
- Chat (varsayılan): Normal sohbet, belge yardımı, sorulara cevap
- Patch: Belge tarama & düzeltme modu — kullanıcı hata türünü ve sayfayı belirtir, Zeta sorunları numaralı listeler, onay gelince DÜZELTİLMİŞ metni yazar
- Puzzle: Derin problem çözme modu — Aziz modeli otomatik seçilir, kullanıcı karmaşık/stratejik sorunlar için kullanır; Puzzle'dan çıkınca önceki modele döner
- Zeta Colors: AI görsel oluşturma modu — metin prompt'undan Nano Banana ile görsel üretir, boyut ve kalite seçenekleri var

⚖️ ZET JUDGE — DOSYA DESTEĞİ:
ZET Judge şu dosya türlerini okuyabilir ve analiz edebilir:
- 📄 PDF (maks. 3 sayfa metin çıkarma)
- 📝 .ms dosyaları (ZET Mindshare belgeleri, maks. 5 sayfa)
- 📋 .docx / .doc (Word belgeleri — tam metin çıkarma)
- 📃 .txt metin dosyaları
- 🖼 Görsel: .png, .jpg, .jpeg, .webp — içerik tarifi, veri/grafik çıkarma, sorun tespiti
- 🎬 Video: .mp4, .mov — sahne sahne döküm, konuşma ve metin çıkarma, özet
Analiz çıktıları: PDF (maks. 3 sayfa) veya .ms (maks. 5 sayfa) formatında dışa aktarılabilir.

🎮 EYLEM SİSTEMİ (KRİTİK):
Kullanıcı senden ayarlarını değiştirmeni, bir şeyi hatırlamanı veya not almasını istediğinde, cevabının EN BAŞINA bu özel etiketleri ekle. Bu etiketler kullanıcıya gösterilmez, sisteme gönderilir:
- Emoji seviyesini değiştir: [ACTION:EMOJI:none] veya [ACTION:EMOJI:low] veya [ACTION:EMOJI:medium] veya [ACTION:EMOJI:high]
- Kişiliği değiştir: [ACTION:MOOD:professional] veya [ACTION:MOOD:cheerful] veya [ACTION:MOOD:curious]
- Belleğe kaydet: [ACTION:MEMORY:hatırlanacak içerik]
- Not al: [ACTION:NOTE:not içeriği]

Örnekler:
- "az emoji kullan" → "[ACTION:EMOJI:low]Tamam, artık daha az emoji kullanacağım."
- "emoji kullanma" → "[ACTION:EMOJI:none]Anlaşıldı, emoji kullanmayacağım."
- "daha neşeli ol" → "[ACTION:MOOD:cheerful]Hemen değiştiriyorum!"
- "bunu hatırla: toplantı cuma 14:00" → "[ACTION:MEMORY:toplantı cuma 14:00]Belleğime kaydettim!"
- "not al: pazarlama fikirleri - sosyal medya..." → "[ACTION:NOTE:pazarlama fikirleri - sosyal medya...]Not alındı!"
- Birden fazla: "[ACTION:MOOD:cheerful][ACTION:EMOJI:high]Hazırım!"

Kullanıcının dilinde yanıt ver!
"""
    
    # Patch mode: override Zeta's behaviour to document scanner/fixer
    if req.mode == 'patch':
        system_message += """

🔧 PATCH MODU — BELGE TARAYICI AKTİF:
Normal sohbet modundan çıktın. Şu an bir belge denetçisi/düzeltici olarak çalışıyorsun.

DAVRANIŞ KURALLARI:
1. Kullanıcının verdiği komutu dikkatlice oku (hangi sayfalar, hangi tür hatalar vb.)
2. Sağlanan BELGE İÇERİĞİ'ni tara ve bulunan sorunları numaralı liste olarak sun
3. Her listeden sonra mutlaka şunu sor: "Bu hataları düzeltmemi ister misiniz?"
4. Kullanıcı "evet", "düzelt", "olur" veya benzeri bir onay verirse — DÜZELTİLMİŞ tam metni yaz
5. Kullanıcının belirttiği kapsamla sınırlı kal — fazlasını yapma
6. Hiçbir sorun bulamazsan: "Belirtilen alanda hata bulunamadı." de
7. Düzeltilmiş metni sağladığında, kullanıcının "Belgeye Uygula" butonunu göreceğini hatırla

ZORUNLU FORMAT — DÜZELTİLMİŞ METİN VERME:
Kullanıcı onay verdiğinde, düzeltilmiş metni MUTLAKA şu etiketler arasına yaz:
[PATCH_START]
düzeltilmiş metin içeriği buraya (açıklama yok, sadece düzeltilmiş metin)
[PATCH_END]
Bu etiketleri SADECE onaylı düzeltme sunarken kullan, başka hiçbir durumda kullanma.
Etiketler dışında kısa bir açıklama yapabilirsin.
"""

    # Inject memories into system prompt
    if memories:
        memory_lines = "\n".join(f"- {m['content']}" for m in memories)
        system_message += f"""

🧠 BELLEĞİNDE KAYITLI BİLGİLER (kullanıcı senden bunları hatırlamanı istedi):
{memory_lines}
Bu bilgileri konuşmaya uygun yerlerde kullan ve başvur.
"""

    # Ortak Hafıza: diğer belgelerden ve oturumlardan geçmiş konuşmalar
    if req.shared_memory:
        other_chats = await db.zeta_chats.find(
            {"user_id": user.user_id, "session_id": {"$ne": session_id}},
            {"_id": 0, "user_message": 1, "ai_response": 1, "doc_id": 1, "created_at": 1}
        ).sort("created_at", -1).to_list(20)

        if other_chats:
            recent = list(reversed(other_chats[:12]))
            shared_lines = []
            for c in recent:
                q = c["user_message"][:200].replace('\n', ' ')
                a = c["ai_response"][:300].replace('\n', ' ')
                shared_lines.append(f"K: {q}\nZ: {a}")
            system_message += f"""

🌐 ORTAK HAFIZA (kullanıcının diğer belgelerindeki geçmiş konuşmalar):
{chr(10).join(shared_lines)}
Bu bilgiler kullanıcının farklı belgelerinde seninle yaptığı konuşmalardan geliyor. İlgili bağlamı yanıtlarına entegre et — eski belgelerden hatırladığın bilgileri gerektiğinde kullan.
"""

    # Inject live canvas/editor context from frontend
    if req.canvas_context:
        system_message += f"""

📍 CANLI EDITÖR DURUMU (şu an kullanıcının ekranında ne var):
{req.canvas_context}
Bu bilgiyi kullanarak daha bağlamsal ve doğru yardım sağla. Kullanıcı hangi araçla çalışıyor, kaç sayfa var, neler var — bunları biliyorsun.
ÖNEMLI: Yukarıdaki JSON listesindeki HER nesne, sayfada kendi x ve y koordinatlarına sahip BAĞIMSIZ bir elementtir (metin kutusu, şekil vb.). Kullanıcı belirli bir cümlenin veya metnin konumunu sorarsa: (1) listede "content" alanı o cümleyi içeren elementi bul, (2) o elementin x ve y değerlerini doğrudan söyle. Birden fazla element aynı y civarındaysa hepsini listele. "Erişimim yok", "render motoru belirler" veya "bilemiyorum" gibi ifadeler YASAKTIR — veriler zaten listede mevcut.
"""

    # If document content is provided, add it to the context
    if req.document_content:
        system_message += f"""

CURRENT DOCUMENT CONTENT:
---
{req.document_content[:5000]}
---
The user may ask questions about this document. Use this content to provide relevant answers.
"""

    # Load chat history for this session
    history_docs = await db.zeta_chats.find(
        {"session_id": session_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(50)
    contents = []
    # CEO mode: inject training exchange at start so model stays in character
    if req.is_ceo:
        contents.append(genai_types.Content(role="user", parts=[genai_types.Part(text="[SİSTEM] CEO modu aktif edildi. Muhammed Bahaddin Yılmaz'a hitap moduna geç.")]))
        contents.append(genai_types.Content(role="model", parts=[genai_types.Part(text="Efendim, CEO PROTOKOLÜ AKTİF. Artık sizinle bu protokol çerçevesinde iletişim kuracağım. Emrinizdeyim, CEO Yılmaz. Başka bir emriniz var mı efendim?")]))
    for h in history_docs:
        contents.append(genai_types.Content(role="user", parts=[genai_types.Part(text=h["user_message"])]))
        contents.append(genai_types.Content(role="model", parts=[genai_types.Part(text=h["ai_response"])]))
    contents.append(genai_types.Content(role="user", parts=[genai_types.Part(text=req.message)]))

    if not api_key:
        return {"response": "Yapılandırma hatası: GEMINI_API_KEY eksik.", "session_id": session_id}

    gemini_model, max_tokens = ZETA_MODEL_MAP.get(req.model or "prime", ZETA_MODEL_MAP["prime"])

    try:
        client = google_genai.Client(api_key=api_key)
        resp = await gemini_generate(
            client, gemini_model, contents,
            genai_types.GenerateContentConfig(
                system_instruction=system_message,
                max_output_tokens=max_tokens,
                tools=[genai_types.Tool(google_search=genai_types.GoogleSearch())],
            )
        )
        response = resp.text
        sources = extract_sources(resp)
        tokens_used_now = getattr(getattr(resp, 'usage_metadata', None), 'total_token_count', 0) or 0
        await add_token_usage(user.user_id, tokens_used_now)
    except Exception as e:
        logging.error(f"Zeta chat Gemini error: {e}")
        return {"response": f"AI hatası: {str(e)}", "session_id": session_id}

    # Parse and strip ACTION tags from response
    import re as _re
    actions = []
    action_pattern = _re.compile(r'\[ACTION:([A-Z]+):([^\]]*)\]')
    for match in action_pattern.finditer(response):
        actions.append({"type": match.group(1), "value": match.group(2).strip()})
    clean_response = action_pattern.sub('', response).lstrip()

    # Track chat usage (free but must be counted)
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    await db.usage.update_one(
        {"user_id": user.user_id, "date": today},
        {"$inc": {"zeta_chat_count": 1}},
        upsert=True
    )
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"ai_chats": 1}})

    # Save chat history (save clean response without action tags)
    await db.zeta_chats.insert_one({
        "user_id": user.user_id,
        "session_id": session_id,
        "doc_id": req.doc_id,
        "user_message": req.message,
        "ai_response": clean_response,
        "created_at": datetime.now(timezone.utc).isoformat()
    })

    return {"response": clean_response, "session_id": session_id, "actions": actions, "sources": sources}


class SpellCheckRequest(BaseModel):
    text: str

@api_router.post("/spell-check")
async def spell_check_text(req: SpellCheckRequest, user: User = Depends(get_current_user)):
    if not req.text or not req.text.strip():
        return {"errors": []}
    try:
        async with httpx.AsyncClient(timeout=8) as client:
            resp = await client.post(
                "https://api.languagetool.org/v2/check",
                data={"text": req.text, "language": "auto"},
                headers={"Content-Type": "application/x-www-form-urlencoded"},
            )
            data = resp.json()
        errors = []
        for match in data.get("matches", []):
            issue_type = match.get("rule", {}).get("issueType", "")
            if issue_type not in ("misspelling", "typographical"):
                continue
            word = req.text[match["offset"]: match["offset"] + match["length"]]
            errors.append({
                "offset": match["offset"],
                "length": match["length"],
                "word": word,
                "suggestions": [r["value"] for r in match.get("replacements", [])[:3]],
            })

        # LanguageTool Turkish support is limited — use Gemini to verify & suggest
        missing = [e for e in errors if not e["suggestions"]]
        if missing:
            try:
                words_json = json.dumps([e["word"] for e in missing], ensure_ascii=False)
                g_client = google_genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
                g_resp = g_client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[{"role": "user", "parts": [{"text": (
                        f"Aşağıdaki Türkçe kelimelerin yazım hatası olup olmadığını kontrol et. "
                        f"Kelime doğruysa boş dizi döndür. Yanlışsa en fazla 3 düzeltme öner. "
                        f"Sadece JSON döndür: {{\"suggestions\": {{\"kelime\": [\"öneri1\"] veya []}}}}\n"
                        f"Kelimeler: {words_json}"
                    )}]}],
                    config=genai_types.GenerateContentConfig(
                        temperature=0.1, max_output_tokens=512,
                        response_mime_type="application/json",
                    )
                )
                raw_g = (g_resp.text or "").strip()
                if raw_g.startswith("```"):
                    raw_g = re.sub(r'^```(?:json)?\s*', '', raw_g, flags=re.IGNORECASE)
                    raw_g = re.sub(r'\s*```\s*$', '', raw_g).strip()
                g_data = json.loads(raw_g)
                g_suggestions = g_data.get("suggestions", {})
                for e in missing:
                    e["suggestions"] = g_suggestions.get(e["word"], [])[:3]
            except Exception as ge:
                logging.warning(f"Gemini spell fallback failed: {ge}")

        # Drop errors where neither LanguageTool nor Gemini could find corrections
        # (likely correctly spelled words that both tools fail to recognize)
        errors = [e for e in errors if e["suggestions"]]

        return {"errors": errors}
    except Exception as e:
        logging.error(f"Spell check error: {e}")
        return {"errors": []}


@api_router.post("/zeta/document-edit")
async def zeta_document_edit(req: ZetaDocEditRequest, user: User = Depends(get_current_user)):
    """Zeta belge düzenleme: kullanıcı isteğine göre canvas elementleri oluşturur/değiştirir."""
    api_key = os.getenv("GEMINI_API_KEY")

    # CEO modunda kredi ve token kontrolü yok
    user_data = await db.users.find_one({"user_id": user.user_id})
    if not req.is_ceo:
        credit_result = await spend_credits(user.user_id, 'doc_edit')
        if not credit_result['success']:
            raise HTTPException(status_code=402, detail=f"Yetersiz kredi. Bu işlem {credit_result['cost']} kredi gerektirir, kalan: {credit_result['credits_remaining']} kredi.")

        token_allowed, tokens_used, token_limit = await check_token_limit(user.user_id, user_data)
        if not token_allowed:
            raise HTTPException(status_code=429, detail=f"Günlük token limitinize ulaştınız ({tokens_used:,}/{token_limit:,}). Limit her gün UTC gece yarısı sıfırlanır.")

    pw = req.page_size.get("width", 794)
    ph = req.page_size.get("height", 1123)

    plan = get_plan_name(user_data)
    doc_edit_limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    basic_shapes_only = doc_edit_limits.get('basic_shapes_only', False) and not req.is_ceo
    allowed_shapes = _BASIC_SHAPES if basic_shapes_only else None

    if basic_shapes_only:
        _shapes_doc_str = '  ' + ', '.join('"' + s + '"' for s in _BASIC_SHAPES) + ' (sadece temel şekiller — Free plan)'
    else:
        _shapes_doc_str = (
            '  "rect" (dikdörtgen), "circle" (daire), "ring" (halka), "triangle", "star", "hexagon",\n'
            '  "diamond", "pentagon", "arrow" (sağa ok), "parallelogram", "heart",\n'
            '  "arrow-right", "arrow-left", "arrow-up", "arrow-down", "arrow-double",\n'
            '  "star3", "star4", "star6", "bubble" (sağa konuşma balonu), "bubble-left",\n'
            '  "diamond-flow", "oval", "cylinder",\n'
            '  "math-sum" (∑), "math-pi" (π), "math-sqrt" (√), "math-inf" (∞), "math-int" (∫),\n'
            '  "bracket-sq" ([ ]), "brace-curly" ({ })'
        )

    def _el_summary(el):
        d = {k: v for k, v in el.items() if k not in ("src", "svgContent", "htmlContent")}
        if el.get('type') == 'text' and 'height' not in d:
            content_lines = max(1, len((el.get('content') or '').split('\n')))
            html_content = el.get('htmlContent') or ''
            html_lines = len(re.findall(r'<br\s*/?>', html_content, re.IGNORECASE)) + 1 if html_content else 1
            explicit_lines = max(content_lines, html_lines)
            w = el.get('width') or 500
            fs = el.get('fontSize') or 16
            plain = re.sub(r'<[^>]+>', '', html_content) if html_content else (el.get('content') or '')
            avg_cpp = max(1, int(w / (fs * 0.6)))
            wrapped = max(1, (len(plain) + avg_cpp - 1) // avg_cpp) if plain else 1
            lines = max(explicit_lines, wrapped)
            d['height'] = round(fs * lines * (el.get('lineHeight') or 1.5))
            # Metin içeriğini AI'ın görebilmesi için ekle (ilk 120 karakter)
            raw_text = (el.get('content') or plain or '').strip()
            if raw_text:
                d['text_preview'] = raw_text[:120] + ('…' if len(raw_text) > 120 else '')
        if 'height' in d:
            d['bottom_y'] = (el.get('y') or 0) + d['height']
        return d

    el_summaries = [_el_summary(el) for el in req.page_elements]
    existing_summary = json.dumps(el_summaries, ensure_ascii=False)
    max_bottom_y = max((s.get('bottom_y', 0) for s in el_summaries), default=0)
    doc_mb = int((req.doc_settings or {}).get('marginBottom', 40))
    page_bottom_limit = ph - doc_mb  # metin bu noktayı geçemez
    page_remaining_px = page_bottom_limit - max_bottom_y
    page_is_full = max_bottom_y >= page_bottom_limit - 20  # 20px: en az bir satır için

    all_pages_summary = ""
    if req.all_pages:
        for pg in req.all_pages:
            pi = pg.get("page_index", 0)
            els = pg.get("elements", [])
            if pi == req.page_index:
                continue
            pg_summary = json.dumps([_el_summary(el) for el in els], ensure_ascii=False)
            all_pages_summary += f"\n  Sayfa {pi}: {pg_summary if pg_summary != '[]' else '(boş)'}"

    _ds = req.doc_settings or {}
    _col_n = max(1, int(_ds.get('columnCount', 1)))
    _col_gap = int(_ds.get('columnGap', 20))
    _col_ml = int(_ds.get('marginLeft', 40))
    _col_mr = int(_ds.get('marginRight', 40))
    _col_avail = pw - _col_ml - _col_mr
    _col_w = round((_col_avail - (_col_n - 1) * _col_gap) / _col_n) if _col_n > 1 else _col_avail
    if _col_n > 1:
        _col_lines = '\n'.join(
            f"  Sütun {i+1}: x={_col_ml + i * (_col_w + _col_gap)}px, genişlik={_col_w}px"
            for i in range(_col_n)
        )
        _col_info = f"SÜTUN DÜZENİ: {_col_n} sütun (aralık={_col_gap}px)\n{_col_lines}\n⚠️ x, y, width YAZMA — frontend sütun konumlarını otomatik hesaplar."
    else:
        _col_info = ""

    system_prompt = f"""Sen ZETA, ZET Mindshare belge editörünün AI asistanısın.
Görevin: Kullanıcının isteğine göre belge sayfalarına JSON operasyonları üretmek.

━━━ BELGE BİLGİSİ ━━━
Sayfa boyutu: {pw}×{ph} piksel (koordinat başlangıcı sol-üst köşe)
Toplam sayfa sayısı: {len(req.all_pages) if req.all_pages else 1}
Aktif sayfa indeksi: {req.page_index}
{f"Mevcut belge ayarları: {json.dumps(req.doc_settings, ensure_ascii=False)}" if req.doc_settings else ""}
{_col_info}

SAYFA DOLULUK DURUMU:
  max_bottom_y: {max_bottom_y}px (mevcut içeriğin en alt noktası)
  alt_marj_sınırı: {page_bottom_limit}px (bu çizgiyi geçen metin sayfadan taşar)
  kalan_alan: {page_remaining_px}px
  → {'⚠️ SAYFA DOLU — yeni metin sığmaz. add_page + target_page: ' + str(req.page_index + 1) + ' kullan.' if page_is_full else '✅ Sayfada yer var (' + str(page_remaining_px) + 'px) — metin aşağıya eklenebilir.'}

Aktif sayfa elementleri (text_preview ile içerikleri, bottom_y ile konumları görülür):
{existing_summary if existing_summary != "[]" else "(sayfa boş)"}
{f"Diğer sayfalar:{all_pages_summary}" if all_pages_summary else ""}

━━━ ELEMENT TİPLERİ VE FORMATLARI ━━━

1. METİN (type: "text")
{{
  "id": "el_<timestamp>_<random4>",
  "type": "text",
  "content": "<düz metin — satır başı için \\n kullan>",
  "htmlContent": "<aynı metin — satır başı için <br> kullan>",
  "fontSize": <8-96>,     // PİKSEL (px) cinsinden — PUNTO DEĞİL. Dönüşüm: px = round(pt * 4/3). Örnekler: 10pt→13, 11pt→15, 12pt→16, 14pt→19, 16pt→21, 18pt→24, 24pt→32, 36pt→48, 48pt→64
  "fontFamily": "Arial",  // Google Font ismi — kullanıcı belirttiyse onu yaz, belirtmediyse "Arial"
  "color": "#000000",     // hex renk — belirtilmediyse #000000
  "lineHeight": 1.5,      // satır aralığı: 1.0 / 1.15 / 1.5 / 2.0 / 2.5 / 3.0
  "textAlign": "left",    // "left" | "center" | "right" | "justify"
  "bold": false,
  "italic": false,
  "underline": false
}}
// ⚠️ x, y, width YAZMA — frontend hesaplar.
  "placement": "after_last",  // NEREYE: "after_last"(varsayılan) | "top" | "bottom" | "new_page"
                               // after_last → sayfadaki son metnin hemen altına
                               // top        → sayfanın üstüne (üst marja)
                               // bottom     → sayfanın altına
                               // new_page   → yeni sayfa açıp oraya

2. ŞEKİL (type: "shape")
{{
  "id": "el_<timestamp>_<random4>",
  "type": "shape",
  "shapeType": "<tip>",   // aşağıdaki listeden biri
  "x": <number>,
  "y": <number>,
  "width": <number>,
  "height": <number>,
  "fill": "#3b82f6",            // hex dolgu rengi
  "gradientStart": null,        // opsiyonel: gradient başlangıç rengi
  "gradientEnd": null           // opsiyonel: gradient bitiş rengi
}}
Kullanılabilir shapeType değerleri:
{_shapes_doc_str}

3. TABLO (type: "table")
{{
  "id": "el_<timestamp>_<random4>",
  "type": "table",
  "x": <number>,
  "y": <number>,
  "width": <number>,   // genelde 400-700
  "height": <number>,  // genelde rows * 36
  "rows": <satır sayısı>,
  "cols": <sütun sayısı>,
  "tableData": [        // 2D dizi: tableData[satır][sütun] = hücre metni
    ["Başlık 1", "Başlık 2", "Başlık 3"],
    ["Veri 1",   "Veri 2",   "Veri 3"],
    ["Veri 4",   "Veri 5",   "Veri 6"]
  ]
}}

4. BAĞLANTI (type: "link")
{{
  "id": "el_<timestamp>_<random4>",
  "type": "link",
  "x": <number>,
  "y": <number>,
  "width": <number>,
  "height": <number>,
  "url": "https://...",
  "label": "Bağlantı metni",
  "color": "#3b82f6"
}}

5. GRAFİK (type: "chart")
{{
  "id": "el_<timestamp>_<random4>",
  "type": "chart",
  "x": <number>,
  "y": <number>,
  "width": 420,
  "height": 320,
  "chartMeta": {{
    "type": "bar" | "pie" | "line",
    "labels": "Ocak,Şubat,Mart",
    "data": "10,20,30",
    "title": "Grafik Başlığı",
    "colors": ["#3b82f6","#10b981","#f59e0b","#ef4444","#8b5cf6"],
    "gradientStart": null,
    "gradientEnd": null
  }}
}}

━━━ OPERASYONLAR ━━━
Her operasyon şu formatlardan biri:

EKLE element:    {{"action": "add",    "element": {{...format...}},              "target_page": <indeks>}}
DEĞİŞTİR:        {{"action": "modify", "element_id": "<id>", "changes": {{...}}, "target_page": <indeks>}}
SİL:             {{"action": "delete", "element_id": "<id>",                     "target_page": <indeks>}}
SAYFA EKLE:      {{"action": "add_page"}}
AYAR DEĞİŞTİR:   {{"action": "update_settings", "settings": {{
  "pageBackground": "<hex>",        // opsiyonel — sayfa arka plan rengi
  "currentFont": "<font adı>",      // opsiyonel — varsayılan yazı tipi
  "currentFontSize": <piksel>,      // opsiyonel — varsayılan yazı boyutu (px)
  "currentColor": "<hex>",          // opsiyonel — varsayılan metin rengi
  "currentLineHeight": <sayı>,      // opsiyonel — satır aralığı (1.0/1.15/1.5/2.0)
  "currentTextAlign": "<hizalama>", // opsiyonel — "left"|"center"|"right"|"justify"
  "pageSize": {{"width": <piksel>, "height": <piksel>}},  // opsiyonel — A4:794×1123, A3:1123×1587, A5:559×794, Letter:816×1056
  "gridVisible": <bool>,            // opsiyonel — ızgara göster/gizle
  "rulerVisible": <bool>            // opsiyonel — cetvel göster/gizle
}}}}
// ← Kullanıcı sayfa boyutu, arka plan rengi, font, satır aralığı, ızgara/cetvel ayarı istediğinde kullan

// ⛔ YASAK — asla üretme: add_path, generate_ai_image, clear_page, delete_page, type:"image", type:"qr"

target_page belirtilmezse aktif sayfa ({req.page_index}) kullanılır.
@N sayfa referansı: kullanıcı "@1", "@2" gibi yazarsa bu sayfa indeksi N-1 demektir (1-indexed).

━━━ ÇIKTI FORMAT ━━━
Kesinlikle sadece JSON döndür. Başka hiçbir metin ekleme.
{{
  "explanation": "<Türkçe, max 2 cümle, ne yaptığını açıkla>",
  "operations": [
    // Operasyon listesi
  ],
  "suggestions": ["<kısa Türkçe öneri 1>", "<kısa Türkçe öneri 2>"]
}}
// suggestions: işlemle ilgili 0-3 adet takip önerisi (örn: "Tablo ekleyeyim mi?", "Şekil ekleyeyim mi?")
// Hiçbir öneri gerekmiyorsa boş dizi []

━━━ KURALLAR ━━━

🔴 ZORUNLU KURAL — METİN YAZMA:
Kullanıcı şunlardan herhangi birini söylerse → MUTLAKA en az bir {"action": "add", "element": {"type": "text", ...}} operasyonu üret. İstisna YOK:
  "yaz", "write", "ekle", "oluştur", "şunu yaz", "bunu yaz", "belgeye ekle", "sayfaya yaz",
  "bir paragraf", "bir metin", "bir başlık", "bir cümle", "şunu yazar mısın", "yazar mısın",
  veya içerik talep eden herhangi bir istek (makale, özet, liste, tanım, açıklama, vb.)
Bu isteklerde operations dizisi BOŞ olamaz. Boş bırakmak HATADIR.

METİN EKLEME (type: "text"):
- x, y, width ASLA YAZMA — frontend hesaplar
- placement alanını kullan (varsayılan: "after_last")
- Kullanıcı konum belirtirse:
    "yazının altına" / "alta ekle" / belirtmedi → "after_last"
    "sayfanın üstüne" / "en üste"               → "top"
    "sayfanın altına" / "en alta"                → "bottom"
    "yeni sayfaya"                               → "new_page"
- Stil (fontSize/fontFamily/color/lineHeight): kullanıcı belirtmediyse YAZMA — frontend son elementten miras alır
- Kullanıcı belirtirse yaz: "kalın" → bold:true, "12 punto" → fontSize:16, "kırmızı" → color:"#ef4444"
- Kullanıcı ne yazmamı istiyorsa AYNEN yaz — kısaltma, özetleme yok
- Uzun içerik (birden fazla paragraf) → her paragraf ayrı "add" operasyonu
- YENİ içerik ekle → "add" operasyonu (yeni element_id ile)
- Mevcut metni düzenle / değiştir → "modify" + element_id + {content, htmlContent} (element listesindeki id'yi kullan)
- Mevcut metni biçimlendir → "modify" + element_id + {bold, italic, color, fontSize, textAlign, lineHeight}
- "şunu değiştir", "şunu düzenle", "yerine yaz" gibi istekler → mevcut element id'sini listeden bul, modify kullan

SAYFA TAŞMA (OVERFLOW) KURALI:
- max_bottom_y >= alt_marj_sınırı ise sayfa doludur — yeni metin ekleme, önce add_page yap
- Sayfa dolu olduğunda: {{"action": "add_page"}} + ardından {{"action": "add", "element": {{...}}, "target_page": {req.page_index + 1}}}
- Kullanıcı "yeni sayfaya yaz" demese bile sayfa marginden taşacaksa otomatik yeni sayfa aç
- Frontend de bu kontrolü yapıyor: AI'ın target_page vermesi sayfa geçişini garanti eder

GENEL:
- ID'ler benzersiz: "el_" + büyük timestamp + 4 karakter (örn: el_1717000000000_a3b2)
- Şekil/tablo/grafik: x ve y koordinatı yaz — sayfa sınırı x+width<={pw}, y+height<={ph}
- Taşıma: modify ile x ve/veya y değiştir
- "Sil" → element_id'yi mevcut listeden al, delete et
- Renk belirtilmediyse: metin #000000, başlık #1a1a2e
- Görsel/resim/fotoğraf ekleme desteklenmiyor — kullanıcıya belirt

━━━ ÖRNEKLER ━━━

Yeni metin ekle ("Times New Roman, 14 punto ile şunu yaz: Merhaba"):
→ {{"action": "add", "element": {{"id": "el_1717000000000_a3b2", "type": "text", "content": "Merhaba", "htmlContent": "Merhaba", "fontSize": 19, "fontFamily": "Times New Roman", "color": "#000000", "lineHeight": 1.5, "textAlign": "left", "bold": false, "italic": false, "underline": false}}}}

Emoji ekle (yeni element):
→ {{"action": "add", "element": {{"id": "el_..._x1y2", "type": "text", "content": "✅ Tamamlandı", "htmlContent": "✅ Tamamlandı", "fontSize": 16, "fontFamily": "Arial", "color": "#000000", "lineHeight": 1.5, "textAlign": "left", "bold": false, "italic": false, "underline": false}}}}

Mevcut metni kalın/italik/renkli yap:
→ {{"action": "modify", "element_id": "<id>", "changes": {{"bold": true, "italic": false, "color": "#e63946"}}}}

Mevcut metnin içeriğini değiştir:
→ {{"action": "modify", "element_id": "<id>", "changes": {{"content": "Yeni metin", "htmlContent": "Yeni metin"}}}}

Sayfa boyutu (A4, A3, Letter):
→ {{"action": "update_settings", "settings": {{"pageSize": {{"width": 794, "height": 1123}}}}}}

Sayfa rengi / font / satır aralığı:
→ {{"action": "update_settings", "settings": {{"pageBackground": "#f0f4ff", "currentFont": "Roboto", "currentLineHeight": 1.5}}}}

Cetvel / ızgara:
→ {{"action": "update_settings", "settings": {{"rulerVisible": true, "gridVisible": false}}}}
"""

    full_prompt = system_prompt + f"\n\n━━━ KULLANICI İSTEĞİ ━━━\n{req.user_request}"
    if req.attached_image_b64:
        full_prompt += "\n\n[Kullanıcı bir fotoğraf/görsel ekledi — yukarıdaki görseli analiz ederek isteği yerine getir.]"

    try:
        client_g = google_genai.Client(api_key=api_key)
        parts = [{"text": full_prompt}]
        if req.attached_image_b64 and req.attached_image_mime:
            parts.append({"inline_data": {"mime_type": req.attached_image_mime, "data": req.attached_image_b64}})
        response = client_g.models.generate_content(
            model="gemini-2.5-flash",
            contents=[{"role": "user", "parts": parts}],
            config=genai_types.GenerateContentConfig(
                temperature=0.4,
                max_output_tokens=16384,
                response_mime_type="application/json",
            )
        )
        raw = (response.text or "").strip()
        # Strip markdown code fences if model added them
        if raw.startswith("```"):
            raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.IGNORECASE)
            raw = re.sub(r'\s*```\s*$', '', raw)
            raw = raw.strip()
        # Extract first JSON object if there's extra text around it
        if not raw.startswith('{'):
            m = re.search(r'\{[\s\S]*\}', raw)
            if m:
                raw = m.group(0)
        result = json.loads(raw)
        await add_token_usage(user.user_id, getattr(getattr(response, 'usage_metadata', None), 'total_token_count', 0) or 0)
        return {
            "explanation": result.get("explanation", ""),
            "operations": result.get("operations", []),
            "suggestions": result.get("suggestions", []),
        }
    except json.JSONDecodeError:
        logging.error(f"zeta_document_edit JSON parse error, raw[:300]: {raw[:300] if 'raw' in dir() else 'N/A'}")
        return {"explanation": "JSON ayrıştırma hatası. Lütfen isteği daha kısa tutarak tekrar deneyin.", "operations": [], "suggestions": []}
    except Exception as e:
        logging.error(f"zeta_document_edit error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/zeta/generate-image")
async def zeta_generate_image(req: ZetaImageRequest, user: User = Depends(get_current_user)):

    # Determine credit action
    credit_action = "nano_banana_pro" if req.pro else "nano_banana"

    # Check plan allows pro
    user_data = await db.users.find_one({"user_id": user.user_id})
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])

    if req.pro and not limits.get('nano_pro', False):
        raise HTTPException(status_code=403, detail="Nano Banana Pro, Pro veya Ultra planda kullanılabilir.")

    # Check aspect ratio allowed
    allowed_sizes = limits.get('custom_image_sizes', ['16:9'])
    if req.aspect_ratio and req.aspect_ratio not in allowed_sizes:
        raise HTTPException(status_code=403, detail=f"Bu görüntü boyutu ({req.aspect_ratio}) planınızda kullanılamaz.")

    # Spend credits
    credit_result = await spend_credits(user.user_id, credit_action)
    if not credit_result['success']:
        raise HTTPException(status_code=429, detail=credit_result['message'])

    # Use IMAGEN_API_KEY if available (separate key with Imagen access), else fall back to GEMINI_API_KEY
    imagen_key = os.getenv("IMAGEN_API_KEY") or os.getenv("GEMINI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY")
    if not imagen_key:
        raise HTTPException(status_code=500, detail="IMAGEN_API_KEY veya GEMINI_API_KEY eksik")

    # Build prompt
    aspect_prompt = f"({req.aspect_ratio} aspect ratio) " if req.aspect_ratio else ""
    quality_prompt = "high quality, professional, detailed, " if req.pro else ""
    full_prompt = f"{aspect_prompt}{quality_prompt}{req.prompt}".strip()

    image_model = "gemini-3-pro-image" if req.pro else "gemini-3.1-flash-image"
    client_img = google_genai.Client(api_key=gemini_key or imagen_key, http_options={"api_version": "v1beta"})
    try:
        if not req.reference_image:
            # Text-to-image
            resp = await gemini_generate(
                client_img, image_model,
                [genai_types.Content(role="user", parts=[genai_types.Part(text=full_prompt)])],
                genai_types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"])
            )
        else:
            # Reference-guided generation
            img_b64 = req.reference_image
            mime_type = "image/png"
            if "," in img_b64:
                header, img_b64 = img_b64.split(",", 1)
                if "jpeg" in header or "jpg" in header:
                    mime_type = "image/jpeg"
            parts = [
                genai_types.Part(inline_data=genai_types.Blob(mime_type=mime_type, data=base64.b64decode(img_b64))),
                genai_types.Part(text=full_prompt),
            ]
            resp = await gemini_generate(
                client_img, image_model,
                [genai_types.Content(role="user", parts=parts)],
                genai_types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"])
            )
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Image generation error: {e}")
        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        await db.usage.update_one({"user_id": user.user_id, "date": today}, {"$inc": {"credits_used": -credit_result['cost']}}, upsert=True)
        raise HTTPException(status_code=500, detail=f"Görsel oluşturma hatası: {str(e)}")

    text_out = ""
    images = []

    # Null/empty candidate guard
    if not resp.candidates or not resp.candidates[0].content or not resp.candidates[0].content.parts:
        # İade et
        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        await db.usage.update_one(
            {"user_id": user.user_id, "date": today},
            {"$inc": {"credits_used": -credit_result['cost']}},
            upsert=True
        )
        finish = getattr(resp.candidates[0] if resp.candidates else None, "finish_reason", None) if resp.candidates else None
        detail = f"Görsel oluşturulamadı (filtre: {finish}). Farklı bir açıklama deneyin." if finish else "Görsel oluşturulamadı, kredi iade edildi."
        raise HTTPException(status_code=500, detail=detail)

    for part in resp.candidates[0].content.parts:
        if getattr(part, "text", None):
            text_out = part.text
        elif getattr(part, "inline_data", None):
            raw = part.inline_data.data
            # raw may be bytes or already b64 string
            if isinstance(raw, bytes):
                encoded = base64.b64encode(raw).decode()
            else:
                encoded = raw
            images.append({
                "mime_type": part.inline_data.mime_type,
                "data": encoded
            })

    if not images:
        # Görsel gelmedi — krediyi iade et
        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        await db.usage.update_one(
            {"user_id": user.user_id, "date": today},
            {"$inc": {"credits_used": -credit_result['cost']}},
            upsert=True
        )
        raise HTTPException(status_code=500, detail="Görsel oluşturulamadı, kredi iade edildi. Tekrar deneyin.")

    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"ai_images": 1}})
    return {"text": text_out, "images": images, "credits_remaining": credit_result['credits_remaining'], "cost": credit_result['cost']}

# AI Photo Edit endpoint
class PhotoEditRequest(BaseModel):
    image_data: str  # Base64 encoded image
    edit_prompt: str  # What to change
    pro: Optional[bool] = False

@api_router.post("/zeta/photo-edit")
async def zeta_photo_edit(req: PhotoEditRequest, user: User = Depends(get_current_user)):
    # Spend credits
    credit_action = "photo_edit_pro" if req.pro else "photo_edit"
    credit_result = await spend_credits(user.user_id, credit_action)
    if not credit_result['success']:
        raise HTTPException(status_code=429, detail=credit_result['message'])

    api_key = os.getenv("GEMINI_API_KEY")

    # Clean base64 data
    image_data = req.image_data
    mime_type = "image/png"
    if ',' in image_data:
        header, image_data = image_data.split(',', 1)
        if "jpeg" in header or "jpg" in header:
            mime_type = "image/jpeg"

    image_model = "gemini-3-pro-image" if req.pro else "gemini-3.1-flash-image"
    client = google_genai.Client(api_key=api_key, http_options={"api_version": "v1beta"})
    resp = await gemini_generate(
        client, image_model,
        [genai_types.Content(role="user", parts=[
            genai_types.Part(inline_data=genai_types.Blob(mime_type=mime_type, data=base64.b64decode(image_data))),
            genai_types.Part(text=f"Edit this image: {req.edit_prompt}")
        ])],
        genai_types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"])
    )

    text_out = ""
    images = []
    for part in resp.candidates[0].content.parts:
        if hasattr(part, "text") and part.text:
            text_out = part.text
        elif hasattr(part, "inline_data") and part.inline_data:
            images.append({
                "mime_type": part.inline_data.mime_type,
                "data": base64.b64encode(part.inline_data.data).decode()
            })

    result = {"text": text_out or "Image edited successfully", "images": images, "credits_remaining": credit_result['credits_remaining'], "cost": credit_result['cost']}
    
    return result

@api_router.post("/zeta/translate")
async def zeta_translate(req: TranslateRequest, user: User = Depends(get_current_user)):
    api_key = os.getenv("GEMINI_API_KEY")
    client = google_genai.Client(api_key=api_key)
    resp = await gemini_generate(
        client, "gemini-2.5-flash", req.text,
        genai_types.GenerateContentConfig(
            system_instruction=f"You are a translator. Translate the given text to {req.target_language}. Return ONLY the translated text, nothing else. No explanations, no quotes."
        )
    )
    await add_token_usage(user.user_id, getattr(getattr(resp, 'usage_metadata', None), 'total_token_count', 0) or 0)
    return {"translated_text": resp.text.strip(), "target_language": req.target_language}

# ============ ELEVENLABS TTS ROUTES ============


class TTSRequest(BaseModel):
    text: str
    voice_id: str = "21m00Tcm4TlvDq8ikWAM"  # Default: Rachel (female)
    model_id: str = "eleven_multilingual_v2"

class VoiceInfo(BaseModel):
    voice_id: str
    name: str
    gender: Optional[str] = None

@api_router.get("/voice/list")
async def list_voices(user: User = Depends(get_current_user)):
    """List available ElevenLabs voices"""
    from elevenlabs import ElevenLabs

    api_key = (os.getenv("ELEVENLABS_API_KEY") or "").strip()
    if not api_key:
        raise HTTPException(status_code=500, detail="ElevenLabs API key not configured")

    try:
        client = ElevenLabs(api_key=api_key)
        voices_response = client.voices.get_all()
        
        # Filter to get some good male/female voices
        voices = []
        for voice in voices_response.voices:
            gender = None
            if voice.labels:
                gender = voice.labels.get("gender", None)
            voices.append({
                "voice_id": voice.voice_id,
                "name": voice.name,
                "gender": gender
            })
        
        return {"voices": voices}
    except Exception as e:
        logging.error(f"Error fetching voices: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching voices: {str(e)}")

@api_router.post("/voice/tts")
async def generate_tts(req: TTSRequest, user: User = Depends(get_current_user)):
    """Generate text-to-speech audio using ElevenLabs"""
    from elevenlabs import ElevenLabs

    api_key = (os.getenv("ELEVENLABS_API_KEY") or "").strip()
    if not api_key:
        raise HTTPException(status_code=500, detail="ElevenLabs API key not configured")

    # Günlük ElevenLabs limiti kontrol
    user_data = await db.users.find_one({"user_id": user.user_id})
    plan = get_plan_name(user_data)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS['free'])
    elevenlabs_daily_limit = limits.get('elevenlabs_daily', 1)
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    usage_doc = await db.usage.find_one({"user_id": user.user_id, "date": today}) or {}
    tts_count_today = usage_doc.get("tts_count", 0)
    if tts_count_today >= elevenlabs_daily_limit:
        raise HTTPException(status_code=429, detail=f"Günlük seslendirme limitine ulaştınız ({elevenlabs_daily_limit}/gün). Daha fazlası için planınızı yükseltin.")

    try:
        client = ElevenLabs(api_key=api_key)

        # Generate audio
        audio_generator = client.text_to_speech.convert(
            text=req.text[:5000],
            voice_id=req.voice_id,
            model_id=req.model_id,
        )

        # Handle both Iterator[bytes] and bytes return types
        if isinstance(audio_generator, bytes):
            audio_data = audio_generator
        else:
            audio_data = b"".join(audio_generator)
        
        # Convert to base64 for transfer
        audio_b64 = base64.b64encode(audio_data).decode()

        await db.usage.update_one(
            {"user_id": user.user_id, "date": today},
            {"$inc": {"tts_count": 1}},
            upsert=True
        )

        return {
            "audio_url": f"data:audio/mpeg;base64,{audio_b64}",
            "text": req.text,
            "voice_id": req.voice_id
        }
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error generating TTS: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating TTS: {str(e)}")

@api_router.post("/pdf/extract-text")
async def extract_pdf_text(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    """Extract text from PDF per page using pypdf for editable canvas import"""
    raw = await file.read()
    try:
        try:
            from pypdf import PdfReader as _PR
        except ImportError:
            from PyPDF2 import PdfReader as _PR
        import io as _io
        reader = _PR(_io.BytesIO(raw))
        pages = []
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()

            # Extract primary font name from page resources
            font_name = None
            try:
                res = page.get("/Resources") or {}
                fonts = res.get("/Font") or {}
                for fobj in fonts.values():
                    try:
                        fobj = fobj.get_object() if hasattr(fobj, "get_object") else fobj
                        base = str(fobj.get("/BaseFont") or "")
                        if base:
                            font_name = base.split("+")[-1]  # strip subset prefix
                            break
                    except Exception:
                        pass
            except Exception:
                pass

            pages.append({"page_num": i + 1, "text": text, "font_name": font_name})
        return {"pages": pages, "total": len(pages)}
    except Exception as e:
        logging.error(f"PDF extract error: {e}")
        raise HTTPException(status_code=500, detail=f"PDF parse error: {str(e)}")

@api_router.post("/voice/stt")
async def speech_to_text(audio: UploadFile = File(...), language: str = "tr", user: User = Depends(get_current_user)):
    """Convert speech to text using ElevenLabs Scribe STT"""
    from elevenlabs import ElevenLabs

    api_key = (os.getenv("ELEVENLABS_API_KEY") or "").strip()
    if not api_key:
        raise HTTPException(status_code=503, detail="ElevenLabs API key not configured")

    try:
        audio_bytes = await audio.read()
        el_client = ElevenLabs(api_key=api_key)
        import io
        kwargs = dict(
            audio=io.BytesIO(audio_bytes),
            model_id="scribe_v1",
            tag_audio_events=False,
            diarize=False,
        )
        if language and language != "auto":
            kwargs["language_code"] = language
        result = el_client.speech_to_text.convert(**kwargs)
        return {"transcript": result.text}
    except Exception as e:
        logging.error(f"STT error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"STT error: {str(e)}")

# ============ CLOUD STORAGE ROUTES (MOCK) ============

@api_router.get("/cloud/google-drive/files")
async def list_google_drive_files(user: User = Depends(get_current_user)):
    # Mock - will be implemented with actual Google Drive API
    return {"files": [], "message": "Google Drive integration coming soon"}

@api_router.post("/cloud/google-drive/upload")
async def upload_to_google_drive(user: User = Depends(get_current_user)):
    return {"message": "Google Drive upload coming soon"}

# ============ GOOGLE DRIVE REAL INTEGRATION ============

from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request as GoogleRequest
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import io
import json as jsonlib

# Get Drive service with auto-refresh
async def get_drive_service(user: User):
    """Get Google Drive service with auto-refresh credentials"""
    creds_doc = await db.drive_credentials.find_one({"user_id": user.user_id})
    if not creds_doc:
        raise HTTPException(
            status_code=400, 
            detail="Google Drive not connected. Please connect your Drive first."
        )
    
    # Create credentials object
    creds = Credentials(
        token=creds_doc["access_token"],
        refresh_token=creds_doc.get("refresh_token"),
        token_uri=creds_doc["token_uri"],
        client_id=creds_doc["client_id"],
        client_secret=creds_doc["client_secret"],
        scopes=creds_doc["scopes"]
    )
    
    # Auto-refresh if expired
    if creds.expired and creds.refresh_token:
        logging.info(f"Refreshing expired token for user {user.user_id}")
        creds.refresh(GoogleRequest())
        
        # Update in database
        await db.drive_credentials.update_one(
            {"user_id": user.user_id},
            {"$set": {
                "access_token": creds.token,
                "expiry": creds.expiry.isoformat() if creds.expiry else None,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
    
    return build('drive', 'v3', credentials=creds)

@api_router.get("/drive/status")
async def get_drive_status(user: User = Depends(get_current_user)):
    """Check if user's Google Drive is connected"""
    creds_doc = await db.drive_credentials.find_one({"user_id": user.user_id})
    if creds_doc and creds_doc.get("access_token"):
        return {"connected": True, "email": creds_doc.get("email")}
    return {"connected": False}

@api_router.get("/drive/connect")
async def connect_google_drive(request: Request, user: User = Depends(get_current_user)):
    """Initiate Google Drive OAuth flow"""
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    
    # If no credentials configured, return mock
    if not client_id or not client_secret:
        frontend_url = os.getenv("FRONTEND_URL", "")
        await users_collection.update_one(
            {"user_id": user.user_id},
            {"$set": {"drive_token": "mock_token", "drive_connected_at": datetime.now(timezone.utc).isoformat()}}
        )
        return {"authorization_url": f"{frontend_url}/dashboard?drive_connected=true", "message": "Drive connected (mock - no credentials configured)"}
    
    redirect_uri = os.getenv("GOOGLE_DRIVE_REDIRECT_URI", f"{os.getenv('FRONTEND_URL', '')}/api/drive/callback")
    
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": client_id,
                    "client_secret": client_secret,
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [redirect_uri]
                }
            },
            scopes=['https://www.googleapis.com/auth/drive.file'],
            redirect_uri=redirect_uri
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent',
            state=user.user_id
        )
        
        logging.info(f"Drive OAuth initiated for user {user.user_id}")
        return {"authorization_url": authorization_url}
    
    except Exception as e:
        logging.error(f"Failed to initiate OAuth: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to initiate OAuth: {str(e)}")

@api_router.get("/drive/callback")
async def drive_callback(code: str = Query(...), state: str = Query(...)):
    """Handle Google Drive OAuth callback"""
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    redirect_uri = os.getenv("GOOGLE_DRIVE_REDIRECT_URI", f"{os.getenv('FRONTEND_URL', '')}/api/drive/callback")
    frontend_url = os.getenv("FRONTEND_URL", "")
    
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": client_id,
                    "client_secret": client_secret,
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [redirect_uri]
                }
            },
            scopes=None,
            redirect_uri=redirect_uri
        )
        
        flow.fetch_token(code=code)
        credentials = flow.credentials
        
        logging.info(f"Drive credentials obtained for user {state}, scopes: {credentials.scopes}")
        
        # Store credentials in database
        await db.drive_credentials.update_one(
            {"user_id": state},
            {"$set": {
                "user_id": state,
                "access_token": credentials.token,
                "refresh_token": credentials.refresh_token,
                "token_uri": credentials.token_uri,
                "client_id": credentials.client_id,
                "client_secret": credentials.client_secret,
                "scopes": list(credentials.scopes) if credentials.scopes else [],
                "expiry": credentials.expiry.isoformat() if credentials.expiry else None,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }},
            upsert=True
        )
        
        logging.info(f"Drive credentials stored for user {state}")
        
        # Redirect to frontend
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=f"{frontend_url}/dashboard?drive_connected=true")
    
    except Exception as e:
        logging.error(f"OAuth callback failed: {str(e)}")
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=f"{frontend_url}/dashboard?drive_error={str(e)}")

@api_router.post("/drive/upload")
async def upload_to_drive(
    doc_id: str = Body(...),
    user: User = Depends(get_current_user)
):
    """Upload document to Google Drive"""
    doc = await docs_collection.find_one({"doc_id": doc_id, "user_id": user.user_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        service = await get_drive_service(user)
        
        # Create file content (JSON)
        file_content = jsonlib.dumps({
            "title": doc.get("title", "Untitled"),
            "content": doc.get("content", {}),
            "created_at": doc.get("created_at"),
            "updated_at": doc.get("updated_at"),
            "app": "ZET Mindshare"
        }, indent=2)
        
        file_metadata = {
            'name': f"{doc.get('title', 'Untitled')}.zet.json",
            'mimeType': 'application/json'
        }
        
        media = MediaIoBaseUpload(
            io.BytesIO(file_content.encode()),
            mimetype='application/json',
            resumable=True
        )
        
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        logging.info(f"File uploaded to Drive: {file.get('id')}")
        
        return {
            "success": True,
            "drive_file_id": file.get('id'),
            "drive_file_name": file.get('name'),
            "drive_link": file.get('webViewLink')
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Drive upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@api_router.get("/drive/files")
async def list_drive_files(user: User = Depends(get_current_user)):
    """List ZET files from connected Google Drive"""
    try:
        service = await get_drive_service(user)
        
        results = service.files().list(
            q="name contains '.zet.json'",
            pageSize=50,
            fields="files(id, name, modifiedTime, webViewLink, size)"
        ).execute()
        
        files = results.get('files', [])
        return {"files": files, "count": len(files)}
    
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Drive list failed: {str(e)}")
        return {"files": [], "error": str(e)}

@api_router.get("/drive/download/{file_id}")
async def download_from_drive(file_id: str, user: User = Depends(get_current_user)):
    """Download file from Google Drive"""
    try:
        service = await get_drive_service(user)
        
        request = service.files().get_media(fileId=file_id)
        file_content = io.BytesIO()
        
        from googleapiclient.http import MediaIoBaseDownload
        downloader = MediaIoBaseDownload(file_content, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_content.seek(0)
        content = jsonlib.loads(file_content.read().decode())
        
        return {"success": True, "content": content}
    
    except Exception as e:
        logging.error(f"Drive download failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@api_router.post("/drive/disconnect")
async def disconnect_drive(user: User = Depends(get_current_user)):
    """Disconnect Google Drive"""
    await db.drive_credentials.delete_one({"user_id": user.user_id})
    return {"success": True, "message": "Google Drive disconnected"}

@api_router.get("/cloud/icloud/files")
async def list_icloud_files(user: User = Depends(get_current_user)):
    # Mock - iCloud integration placeholder
    return {"files": [], "message": "iCloud integration coming soon"}

@api_router.post("/cloud/icloud/upload")
async def upload_to_icloud(user: User = Depends(get_current_user)):
    return {"message": "iCloud upload coming soon"}

# ============ QUEST MAP ============

@api_router.get("/quests/progress")
async def get_quest_progress(user: User = Depends(get_current_user)):
    user_data = await users_collection.find_one({"user_id": user.user_id}, {"_id": 0, "completed_quests": 1, "quest_xp": 1, "active_time_seconds": 1})
    return {
        "completed_quests": user_data.get("completed_quests", []) if user_data else [],
        "quest_xp": user_data.get("quest_xp", 0) if user_data else 0,
        "active_time_seconds": user_data.get("active_time_seconds", 0) if user_data else 0,
    }

ISTANBUL_TZ_OFFSET = 3  # UTC+3

@api_router.post("/users/heartbeat")
async def user_heartbeat(user: User = Depends(get_current_user)):
    # Sunucu saatini kullan (İstanbul UTC+3) — client saatine güvenme
    now_utc = datetime.now(timezone.utc)
    now_istanbul = now_utc + timedelta(hours=ISTANBUL_TZ_OFFSET)

    user_data = await users_collection.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "last_heartbeat": 1}
    )
    last_hb = user_data.get("last_heartbeat") if user_data else None

    # Rate-limit: son heartbeat'ten en az 25 saniye geçmeli
    if last_hb:
        try:
            last_dt = datetime.fromisoformat(last_hb.replace("Z", "+00:00"))
            elapsed = (now_utc - last_dt).total_seconds()
            if elapsed < 25:
                return {"ok": False, "reason": "too_soon", "elapsed": round(elapsed)}
        except Exception:
            pass

    try:
        await users_collection.update_one(
            {"user_id": user.user_id},
            {
                "$inc": {"active_time_seconds": 30},
                "$set": {"last_heartbeat": now_utc.isoformat()}
            },
            upsert=True
        )
    except Exception as e:
        logging.error(f"heartbeat DB error: {e}")
        return {"ok": False, "reason": "db_error"}
    return {"ok": True, "server_time_istanbul": now_istanbul.strftime("%H:%M:%S")}

@api_router.post("/user/time-spent")
async def update_time_spent(data: dict = Body(...), user: User = Depends(get_current_user)):
    """Her dakika frontend'den gelen oturum süresi — saniye cinsinden."""
    seconds = int(data.get("seconds", 0))
    if seconds <= 0 or seconds > 120:  # max 2 dakika per çağrı (güvenlik)
        return {"ok": False, "reason": "invalid_seconds"}
    await users_collection.update_one(
        {"user_id": user.user_id},
        {"$inc": {"active_time_seconds": seconds}},
        upsert=True
    )
    return {"ok": True, "added_seconds": seconds}

class QuestCompleteRequest(BaseModel):
    xp: int = 10
    elapsed_seconds: Optional[float] = None

RANK_THRESHOLDS = [
    {"name": "Demir",   "xp": 0},
    {"name": "Gümüş",  "xp": 100},
    {"name": "Altın",  "xp": 500},
    {"name": "Elmas",  "xp": 1500},
    {"name": "Zümrüt", "xp": 5000},
    {"name": "Endless", "xp": 15000},
]

def get_rank_name(xp: int) -> str:
    rank = RANK_THRESHOLDS[0]["name"]
    for r in RANK_THRESHOLDS:
        if xp >= r["xp"]:
            rank = r["name"]
    return rank

@api_router.post("/quests/{quest_id}/complete")
async def complete_quest(quest_id: int, req: QuestCompleteRequest, user: User = Depends(get_current_user)):
    user_data = await users_collection.find_one({"user_id": user.user_id}, {"_id": 0, "completed_quests": 1, "quest_xp": 1})
    completed = user_data.get("completed_quests", []) if user_data else []
    current_xp = user_data.get("quest_xp", 0) if user_data else 0
    if quest_id in completed:
        return {"completed_quests": completed, "quest_xp": current_xp, "already_completed": True}
    if req.elapsed_seconds is not None:
        asyncio.create_task(hafizz.check_task_timing(db, user.user_id, str(quest_id), req.elapsed_seconds))
    completed.append(quest_id)
    new_xp = current_xp + req.xp

    old_rank = get_rank_name(current_xp)
    new_rank = get_rank_name(new_xp)
    rank_up = new_rank != old_rank

    await users_collection.update_one(
        {"user_id": user.user_id},
        {"$set": {"completed_quests": completed, "quest_xp": new_xp}},
        upsert=True
    )
    return {
        "completed_quests": completed,
        "quest_xp": new_xp,
        "rank_up": rank_up,
        "new_rank": new_rank if rank_up else None,
        "old_rank": old_rank if rank_up else None,
    }

# ============ COLLABORATION & SHARING ============

shares_collection = db.shares
comments_collection = db.comments

@api_router.post("/documents/{doc_id}/share")
async def create_share_link(doc_id: str, share: ShareLinkCreate, user: User = Depends(get_current_user)):
    doc = await db.documents.find_one({"doc_id": doc_id, "user_id": user.user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Document not found")
    share_id = f"share_{uuid.uuid4().hex[:12]}"
    share_data = {
        "share_id": share_id, "doc_id": doc_id, "owner_id": user.user_id,
        "permission": share.permission,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "expires_at": (datetime.now(timezone.utc) + timedelta(hours=share.expires_hours)).isoformat() if share.expires_hours else None,
        "active": True
    }
    await shares_collection.insert_one(share_data)
    return {"share_id": share_id, "permission": share.permission, "expires_at": share_data["expires_at"]}

@api_router.get("/documents/{doc_id}/shares")
async def get_share_links(doc_id: str, user: User = Depends(get_current_user)):
    shares = await shares_collection.find({"doc_id": doc_id, "owner_id": user.user_id, "active": True}, {"_id": 0}).to_list(50)
    return shares

@api_router.delete("/share/{share_id}")
async def revoke_share_link(share_id: str, user: User = Depends(get_current_user)):
    await shares_collection.update_one({"share_id": share_id, "owner_id": user.user_id}, {"$set": {"active": False}})
    return {"status": "revoked"}

@api_router.get("/shared/{share_id}")
async def get_shared_document(share_id: str, request: Request):
    share = await shares_collection.find_one({"share_id": share_id, "active": True}, {"_id": 0})
    if not share:
        raise HTTPException(404, "Share link not found or expired")
    if share.get("expires_at"):
        exp = datetime.fromisoformat(share["expires_at"])
        if datetime.now(timezone.utc) > exp:
            raise HTTPException(410, "Share link expired")
    doc = await db.documents.find_one({"doc_id": share["doc_id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Document not found")
    # Get current user if logged in
    user_id = None
    user_name = "Guest"
    session_id = request.cookies.get("session_id")
    if session_id:
        session = await db.sessions.find_one({"session_id": session_id}, {"_id": 0})
        if session:
            u = await users_collection.find_one({"user_id": session["user_id"]}, {"_id": 0})
            if u:
                user_id = u["user_id"]
                user_name = u.get("name", "User")
    return {
        "document": doc, "permission": share["permission"],
        "owner_id": share["owner_id"],
        "viewer": {"user_id": user_id or f"guest_{uuid.uuid4().hex[:8]}", "name": user_name}
    }

# ============ COMMENTS ============

@api_router.post("/documents/{doc_id}/comments")
async def add_comment(doc_id: str, comment: CommentCreate, user: User = Depends(get_current_user)):
    comment_id = f"cmt_{uuid.uuid4().hex[:12]}"
    comment_data = {
        "comment_id": comment_id, "doc_id": doc_id,
        "user_id": user.user_id, "user_name": user.name,
        "content": comment.content, "element_id": comment.element_id,
        "x": comment.x, "y": comment.y, "page_index": comment.page_index,
        "replies": [], "resolved": False,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await comments_collection.insert_one(comment_data)
    comment_data.pop("_id", None)
    return comment_data

@api_router.get("/documents/{doc_id}/comments")
async def get_comments(doc_id: str, user: User = Depends(get_current_user)):
    comments = await comments_collection.find({"doc_id": doc_id}, {"_id": 0}).to_list(200)
    return comments

@api_router.post("/comments/{comment_id}/reply")
async def reply_to_comment(comment_id: str, reply: CommentReply, user: User = Depends(get_current_user)):
    reply_data = {
        "reply_id": f"reply_{uuid.uuid4().hex[:8]}",
        "user_id": user.user_id, "user_name": user.name,
        "content": reply.content,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await comments_collection.update_one(
        {"comment_id": comment_id}, {"$push": {"replies": reply_data}}
    )
    return reply_data

@api_router.put("/comments/{comment_id}/resolve")
async def resolve_comment(comment_id: str, user: User = Depends(get_current_user)):
    await comments_collection.update_one({"comment_id": comment_id}, {"$set": {"resolved": True}})
    return {"status": "resolved"}

@api_router.delete("/comments/{comment_id}")
async def delete_comment(comment_id: str, user: User = Depends(get_current_user)):
    await comments_collection.delete_one({"comment_id": comment_id, "user_id": user.user_id})
    return {"status": "deleted"}

@api_router.get("/shared/{share_id}/comments")
async def get_shared_comments(share_id: str):
    share = await shares_collection.find_one({"share_id": share_id, "active": True}, {"_id": 0})
    if not share:
        raise HTTPException(404, "Share not found")
    comments = await comments_collection.find(
        {"doc_id": share["doc_id"], "resolved": False}, {"_id": 0}
    ).sort("created_at", 1).to_list(200)
    return comments

@api_router.post("/judge/feedback")
async def submit_judge_feedback(feedback: JudgeFeedbackCreate):
    feedback_data = {
        "feedback_id": f"fb_{uuid.uuid4().hex[:12]}",
        "session_id": feedback.session_id,
        "message_index": feedback.message_index,
        "feedback_type": feedback.feedback_type,
        "message_content": feedback.message_content[:500],
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.judge_feedback.insert_one(feedback_data)
    count = await db.judge_feedback.count_documents({})
    if count % 100 == 0:
        positive = await db.judge_feedback.count_documents({"feedback_type": "positive"})
        negative = await db.judge_feedback.count_documents({"feedback_type": "negative"})
        await send_email(
            "support@zetstudiointl.com",
            f"Judge Geri Bildirim Raporu — {count} Toplam",
            f"<h2>Judge Geri Bildirim Özeti</h2><p>Toplam geri bildirim: <b>{count}</b></p>"
            f"<p>👍 Pozitif: <b>{positive}</b></p><p>👎 Negatif: <b>{negative}</b></p>"
            f"<p>Bu rapor her 100 geri bildirimde otomatik gönderilir.</p>",
        )
    return {"status": "ok"}

# ============ WEBSOCKET COLLABORATION ============

@app.websocket("/api/ws/collab/{doc_id}")
async def websocket_collab(websocket: WebSocket, doc_id: str):
    # Accept first, then authenticate
    user_id = None
    user_name = "Guest"
    try:
        await websocket.accept()
        # Wait for auth message
        auth_msg = await asyncio.wait_for(websocket.receive_json(), timeout=10)
        user_id = auth_msg.get("user_id", f"guest_{uuid.uuid4().hex[:8]}")
        user_name = auth_msg.get("user_name", "Guest")
        
        # Register in room
        if doc_id not in collab_manager.active_rooms:
            collab_manager.active_rooms[doc_id] = {}
        color = f"hsl({hash(user_id) % 360}, 70%, 60%)"
        collab_manager.active_rooms[doc_id][user_id] = {
            "ws": websocket, "name": user_name, "color": color,
            "cursor": None, "connected_at": datetime.now(timezone.utc).isoformat()
        }
        
        # Send current users list
        await websocket.send_json({
            "type": "connected",
            "user_id": user_id, "color": color,
            "users": collab_manager._get_users(doc_id)
        })
        
        # Notify others
        await collab_manager.broadcast(doc_id, user_id, {
            "type": "user_joined",
            "user_id": user_id, "name": user_name, "color": color,
            "users": collab_manager._get_users(doc_id)
        })
        
        # Message loop
        while True:
            data = await websocket.receive_json()
            msg_type = data.get("type")
            
            if msg_type == "cursor_move":
                collab_manager.active_rooms[doc_id][user_id]["cursor"] = data.get("position")
                await collab_manager.broadcast(doc_id, user_id, {
                    "type": "cursor_move",
                    "user_id": user_id, "name": user_name, "color": color,
                    "position": data.get("position")
                })
            elif msg_type == "element_update":
                await collab_manager.broadcast(doc_id, user_id, {
                    "type": "element_update",
                    "user_id": user_id, "name": user_name,
                    "elements": data.get("elements"),
                    "page_index": data.get("page_index")
                })
            elif msg_type == "element_add":
                await collab_manager.broadcast(doc_id, user_id, {
                    "type": "element_add",
                    "user_id": user_id, "element": data.get("element")
                })
            elif msg_type == "element_delete":
                await collab_manager.broadcast(doc_id, user_id, {
                    "type": "element_delete",
                    "user_id": user_id, "element_id": data.get("element_id")
                })
            elif msg_type == "comment_add":
                await collab_manager.broadcast(doc_id, user_id, {
                    "type": "comment_add",
                    "user_id": user_id, "name": user_name,
                    "comment": data.get("comment")
                })
            elif msg_type == "page_update":
                await collab_manager.broadcast(doc_id, user_id, {
                    "type": "page_update",
                    "user_id": user_id,
                    "pages": data.get("pages")
                })
            elif msg_type == "ping":
                await websocket.send_json({"type": "pong"})
    except WebSocketDisconnect:
        pass
    except Exception as e:
        logging.getLogger(__name__).error(f"WebSocket error: {e}")
    finally:
        collab_manager.disconnect(doc_id, user_id)
        await collab_manager.broadcast(doc_id, user_id, {
            "type": "user_left",
            "user_id": user_id, "name": user_name,
            "users": collab_manager._get_users(doc_id)
        })

@api_router.get("/documents/{doc_id}/online")
async def get_online_users(doc_id: str):
    users = collab_manager._get_users(doc_id)
    return {"users": users, "count": len(users)}

# ============ ROOT ============

@api_router.get("/")
async def root():
    return {"message": "ZET Mindshare API", "version": "1.0.0"}

# ============ CEO DATA EXPORT ============

async def get_most_active_users(limit: int = 10) -> list:
    pipeline = [
        {"$lookup": {"from": "posts", "localField": "user_id", "foreignField": "author_id", "as": "posts"}},
        {"$addFields": {"post_count": {"$size": "$posts"}}},
        {"$sort": {"post_count": -1}},
        {"$limit": limit},
        {"$project": {"_id": 0, "user_id": 1, "email": 1, "name": 1, "username": 1, "post_count": 1, "followers_count": 1}},
    ]
    return await db.users.aggregate(pipeline).to_list(length=limit)

async def get_trending_words(limit: int = 20) -> list:
    import re
    from collections import Counter
    posts = await db.posts.find({"type": "text"}, {"content": 1, "_id": 0}).to_list(length=500)
    words = []
    for p in posts:
        words.extend(re.findall(r'\b[a-zA-ZğüşıöçĞÜŞİÖÇ]{4,}\b', p.get("content", "")))
    stopwords = {"için", "ile", "bir", "olan", "veya", "gibi", "daha", "çok", "the", "and", "that", "this", "with", "from"}
    filtered = [w.lower() for w in words if w.lower() not in stopwords]
    return [{"word": w, "count": c} for w, c in Counter(filtered).most_common(limit)]

async def build_export_payload(export_type: str) -> dict:
    now = datetime.now(timezone.utc)
    payload: dict = {"exported_at": now.isoformat(), "type": export_type}

    if export_type in ("last7", "last30"):
        days = 7 if export_type == "last7" else 30
        since = (now - timedelta(days=days)).isoformat()
        payload["users"] = await db.users.find({"created_at": {"$gte": since}}, {"_id": 0, "password": 0}).to_list(length=5000)
        payload["posts"] = await db.posts.find({"created_at": {"$gte": since}}, {"_id": 0}).to_list(length=5000)
        payload["documents"] = await db.documents.find({"created_at": {"$gte": since}}, {"_id": 0}).to_list(length=5000)
        return payload

    if export_type in ("users", "all"):
        payload["users"] = await db.users.find({}, {"_id": 0, "password": 0}).to_list(length=10000)
    if export_type in ("posts", "all"):
        payload["posts"] = await db.posts.find({}, {"_id": 0}).to_list(length=10000)
    if export_type in ("documents", "all"):
        payload["documents"] = await db.documents.find({}, {"_id": 0}).to_list(length=10000)
    if export_type in ("comments", "all"):
        payload["comments"] = await db.comments.find({}, {"_id": 0}).to_list(length=10000)
    if export_type in ("notes", "all"):
        payload["notes"] = await db.notes.find({}, {"_id": 0}).to_list(length=10000)
    if export_type in ("analytics", "all"):
        total_users = await db.users.count_documents({})
        total_posts = await db.posts.count_documents({})
        total_docs = await db.documents.count_documents({})
        total_comments = await db.comments.count_documents({})
        payload["analytics"] = {
            "total_users": total_users,
            "total_posts": total_posts,
            "total_documents": total_docs,
            "total_comments": total_comments,
            "most_active_users": await get_most_active_users(10),
            "trending_words": await get_trending_words(20),
        }

    return payload

async def send_export_email(export_type: str, payload: dict):
    resend.api_key = os.environ.get("RESEND_API_KEY", "")
    sender = os.environ.get("SENDER_EMAIL", "onboarding@resend.dev")
    json_bytes = json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")
    import base64 as b64
    attachment_b64 = b64.b64encode(json_bytes).decode("ascii")
    filename = f"zet_export_{export_type}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
    params = {
        "from": sender,
        "to": [CEO_EMAIL],
        "subject": f"[ZET Data Export] {export_type.upper()} — {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC",
        "html": f"<p>ZET Mindshare veri dışa aktarımı: <b>{export_type}</b></p><p>Tarih: {datetime.now(timezone.utc).isoformat()}</p><p>Dosya ek olarak gönderilmiştir.</p>",
        "attachments": [{"filename": filename, "content": attachment_b64}],
    }
    resend.Emails.send(params)

# ============ KİMLİK DOĞRULAMA ============

class IdentityVerifyRequest(BaseModel):
    full_name: str
    id_type: str  # 'tc_kimlik' | 'pasaport' | 'ehliyet'
    id_number: str
    front_image: str   # base64
    back_image: Optional[str] = None  # arka yüz (TC kimlik için)
    selfie_image: str  # selfie with ID

@api_router.post("/users/verify-identity")
async def submit_identity_verification(body: IdentityVerifyRequest, user: User = Depends(get_current_user)):
    existing = await db.identity_verifications.find_one({"user_id": user.user_id, "status": {"$in": ["pending", "approved"]}})
    if existing:
        if existing["status"] == "approved":
            raise HTTPException(status_code=400, detail="Kimliğiniz zaten doğrulanmış.")
        raise HTTPException(status_code=400, detail="Zaten bekleyen bir doğrulama başvurunuz var.")
    if not body.full_name.strip() or len(body.id_number.strip()) < 4:
        raise HTTPException(status_code=400, detail="Ad soyad ve kimlik numarası zorunludur.")
    verification_id = f"verif_{uuid.uuid4().hex[:16]}"
    now = datetime.now(timezone.utc).isoformat()
    await db.identity_verifications.insert_one({
        "verification_id": verification_id,
        "user_id": user.user_id,
        "user_email": user.email,
        "full_name": body.full_name.strip(),
        "id_type": body.id_type,
        "id_number": body.id_number.strip(),
        "front_image": body.front_image,
        "back_image": body.back_image,
        "selfie_image": body.selfie_image,
        "status": "pending",
        "submitted_at": now,
    })
    await db.users.update_one({"user_id": user.user_id}, {"$set": {"identity_status": "pending"}})
    return {"message": "Kimlik doğrulama başvurunuz alındı. En geç 24 saat içinde incelenecektir.", "verification_id": verification_id}

@api_router.get("/users/verify-identity/status")
async def get_identity_status(user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id}, {"identity_status": 1})
    verif = await db.identity_verifications.find_one({"user_id": user.user_id}, {"_id": 0, "front_image": 0, "back_image": 0, "selfie_image": 0})
    return {"identity_status": user_data.get("identity_status", "none") if user_data else "none", "verification": verif}

@api_router.get("/admin/identity-verifications")
async def list_identity_verifications(user: User = Depends(get_current_user)):
    if user.email != CEO_EMAIL:
        me = await db.users.find_one({"user_id": user.user_id}, {"is_privileged": 1})
        if not me or not me.get("is_privileged"):
            raise HTTPException(status_code=403)
    verifs = await db.identity_verifications.find({"status": "pending"}, {"_id": 0, "front_image": 0, "back_image": 0, "selfie_image": 0}).sort("submitted_at", 1).to_list(100)
    return {"verifications": verifs}

class IdentityDecision(BaseModel):
    decision: str  # 'approve' | 'reject'
    reason: Optional[str] = None
    verified_type: Optional[str] = None  # 'blue' | 'gold' | 'red' (only on approve)

@api_router.post("/admin/identity-verifications/{verification_id}/decide")
async def decide_identity_verification(verification_id: str, body: IdentityDecision, user: User = Depends(get_current_user)):
    if user.email != CEO_EMAIL:
        me = await db.users.find_one({"user_id": user.user_id}, {"is_privileged": 1})
        if not me or not me.get("is_privileged"):
            raise HTTPException(status_code=403)
    if body.decision not in ("approve", "reject"):
        raise HTTPException(status_code=400, detail="Karar 'approve' veya 'reject' olmalı.")
    verif = await db.identity_verifications.find_one({"verification_id": verification_id})
    if not verif:
        raise HTTPException(status_code=404, detail="Doğrulama bulunamadı.")
    if verif["status"] != "pending":
        raise HTTPException(status_code=400, detail="Bu başvuru zaten işleme alınmış.")
    now = datetime.now(timezone.utc).isoformat()
    new_status = "approved" if body.decision == "approve" else "rejected"
    await db.identity_verifications.update_one(
        {"verification_id": verification_id},
        {"$set": {"status": new_status, "decided_at": now, "decided_by": user.email, "reason": body.reason}}
    )
    user_update: dict = {"identity_status": new_status}
    if body.decision == "approve":
        user_update["identity_verified"] = True
        if body.verified_type:
            user_update["verified_type"] = body.verified_type
    await db.users.update_one({"user_id": verif["user_id"]}, {"$set": user_update})
    return {"message": f"Başvuru {'onaylandı' if body.decision == 'approve' else 'reddedildi'}."}


class ExportRequest(BaseModel):
    type: str  # posts, users, documents, comments, notes, analytics, all, last7, last30

VALID_EXPORT_TYPES = {"posts", "users", "documents", "comments", "notes", "analytics", "all", "last7", "last30"}

@api_router.post("/admin/export-data")
async def export_data(body: ExportRequest, user: User = Depends(get_current_user)):
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403, detail="Sadece CEO veri dışa aktarabilir.")
    if body.type not in VALID_EXPORT_TYPES:
        raise HTTPException(status_code=400, detail=f"Geçersiz tip. Geçerli tipler: {', '.join(sorted(VALID_EXPORT_TYPES))}")
    payload = await build_export_payload(body.type)
    await send_export_email(body.type, payload)
    await db.export_logs.insert_one({
        "requested_by": user.email,
        "type": body.type,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "record_counts": {k: len(v) if isinstance(v, list) else None for k, v in payload.items() if k not in ("exported_at", "type")},
    })
    return {"success": True, "message": f"'{body.type}' verisi CEO e-postasına gönderildi."}

@api_router.post("/admin/migrate-old-data")
async def migrate_old_data(user: User = Depends(get_current_user)):
    """Eski kullanıcı alanlarını yeni uygulama-bazlı alanlara taşır. Mevcut veriye DOKUNMAZ."""
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403, detail="Sadece CEO bu işlemi yapabilir.")

    all_users = await db.users.find({}).to_list(None)
    migrated = 0

    for u in all_users:
        updates = {}

        # credits → mindshare_credits
        if "credits" in u and "mindshare_credits" not in u:
            updates["mindshare_credits"] = u.get("credits", 0)

        # rank → mindshare_rank
        if "mindshare_rank" not in u:
            updates["mindshare_rank"] = u.get("rank", "iron")

        # quest_xp / xp → mindshare_xp
        if "mindshare_xp" not in u:
            updates["mindshare_xp"] = u.get("quest_xp", u.get("xp", 0))

        # Judge alanları — her zaman varsayılan ekle
        if "judge_credits" not in u:
            updates["judge_credits"] = 0
        if "judge_rank" not in u:
            updates["judge_rank"] = "iron"
        if "judge_xp" not in u:
            updates["judge_xp"] = 0

        if updates:
            await db.users.update_one({"_id": u["_id"]}, {"$set": updates})
            migrated += 1

    return {
        "success": True,
        "migrated": migrated,
        "total": len(all_users),
        "message": f"{migrated}/{len(all_users)} kullanıcı güncellendi."
    }

@api_router.post("/admin/normalize-subscriptions")
async def normalize_all_subscriptions(user: User = Depends(get_current_user)):
    """Tüm kullanıcıların subscription alanını dict formatına çevirir (string → dict)."""
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403, detail="Sadece CEO bu işlemi yapabilir.")

    all_users = await db.users.find({}).to_list(None)
    fixed = 0

    for u in all_users:
        updates = {}
        for field in ("subscription", "mindshare_subscription", "judge_subscription"):
            val = u.get(field)
            if isinstance(val, str):
                updates[field] = {"plan": val if val else "free", "status": "active"}
        if updates:
            await db.users.update_one({"_id": u["_id"]}, {"$set": updates})
            fixed += 1

    return {"fixed_users": fixed, "total": len(all_users),
            "message": f"{fixed} kullanıcının subscription alanı normalize edildi."}

async def send_weekly_report():
    """Her Pazartesi 09:00 UTC'de haftalık rapor gönder."""
    while True:
        now = datetime.now(timezone.utc)
        days_until_monday = (7 - now.weekday()) % 7 or 7
        next_monday = now.replace(hour=9, minute=0, second=0, microsecond=0) + timedelta(days=days_until_monday)
        wait_seconds = (next_monday - now).total_seconds()
        await asyncio.sleep(wait_seconds)
        try:
            payload = await build_export_payload("analytics")
            payload["last7"] = await build_export_payload("last7")
            await send_export_email("weekly_report", payload)
            await db.export_logs.insert_one({
                "requested_by": "system",
                "type": "weekly_report",
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })
        except Exception as e:
            logging.getLogger(__name__).error(f"Weekly report error: {e}")

async def expire_boosts_loop():
    """Süresi dolan boost'ları saatlik kapat."""
    while True:
        await asyncio.sleep(3600)
        try:
            now = datetime.now(timezone.utc).isoformat()
            await db.posts.update_many(
                {"boost.active": True, "boost.expires_at": {"$lt": now}},
                {"$set": {"boost.active": False}}
            )
        except Exception as e:
            logging.getLogger(__name__).error(f"Boost expiry error: {e}")

async def _run_renewal_check():
    """Wrapper called by APScheduler — runs the renewal logic directly."""
    try:
        result = await check_subscription_renewals()
        logging.info(f"Renewal check: {result}")
    except Exception as e:
        logging.error(f"Renewal check error: {e}")

async def _check_season_end():
    """APScheduler: aktif sezon bittiyse ödülleri dağıt ve sıfırla."""
    try:
        season = await db.seasons.find_one({"status": "active"}, {"_id": 0})
        if not season:
            return
        today = datetime.utcnow().strftime("%Y-%m-%d")
        if season.get("end_date", "9999-99-99") <= today:
            count = await distribute_season_rewards_and_reset()
            logging.info(f"[Season] Auto-ended season, rewarded {count} users")
    except Exception as e:
        logging.error(f"Season end check error: {e}")

async def keepalive_ping():
    while True:
        try:
            await db.command("ping")
        except Exception:
            pass
        await asyncio.sleep(60)

async def _hafizz_post_login(db, user_id: str, ip: str):
    """Login sonrası ülke kontrolü + şüpheli durum bildirimi (background)."""
    try:
        country = await hafizz.get_ip_country(ip, db)
        if country:
            suspicious = await hafizz.check_multi_country_login(db, user_id, country)
            if suspicious:
                asyncio.create_task(hafizz.notify_suspicious(db, user_id, f"24 saatte birden fazla ülkeden giriş"))
    except Exception as e:
        logging.warning(f"Hafız post-login hatası: {e}")


async def _ensure_indexes():
    try:
        await db.doc_presence.create_index([("doc_id", 1), ("last_seen", 1)])
        await db.documents.create_index([("user_id", 1), ("updated_at", -1)])
        await db.documents.create_index([("user_id", 1), ("pinned", -1)])
        await db.notes.create_index([("user_id", 1), ("updated_at", -1)])
        await db.notebooks.create_index([("user_id", 1)])
        await db.quests.create_index([("user_id", 1)])
        await db.users.create_index([("user_id", 1)], unique=True)
        await db.subscriptions.create_index([("user_id", 1)])
        await db.inventory.create_index([("user_id", 1)])
        await db.doc_presence.create_index([("doc_id", 1), ("session_id", 1)], unique=True)
        await db.user_sessions.create_index([("session_token", 1)], unique=True)
        await db.user_sessions.create_index([("expires_at", 1)], expireAfterSeconds=0)
        await db.zet_id_tokens.create_index([("token", 1)], unique=True)
        await db.zet_id_tokens.create_index([("expires_at", 1)], expireAfterSeconds=0)
        await db.zet_id_auth_codes.create_index([("code", 1)], unique=True)
        await db.zet_id_auth_codes.create_index([("expires_at", 1)], expireAfterSeconds=0)
        await db.zet_id_clients.create_index([("client_id", 1)], unique=True)
        await db.zet_id_sessions.create_index([("jti", 1)], unique=True)
        await db.zet_id_sessions.create_index([("refresh_token", 1)], unique=True)
        await db.zet_id_sessions.create_index([("user_id", 1)])
        await db.users.create_index([("zet_id", 1)], unique=True, sparse=True)
        await db.token_usage.create_index([("user_id", 1), ("date", 1)], unique=True)
        await db.chest_usage.create_index([("user_id", 1), ("month", 1)], unique=True)
        logging.info("MongoDB indexes OK")
    except Exception as e:
        logging.warning(f"Index creation warning (non-fatal): {e}")

@app.on_event("startup")
async def start_background_tasks():
    try:
        asyncio.create_task(send_weekly_report())
        asyncio.create_task(expire_boosts_loop())
        asyncio.create_task(keepalive_ping())
        asyncio.create_task(_ensure_indexes())
        asyncio.create_task(hafizz.setup_indexes(db))
        logging.info("Background tasks created successfully")
    except Exception as e:
        logging.error(f"Background task creation failed (non-fatal): {e}")

    if _APSCHEDULER_AVAILABLE:
        try:
            scheduler = AsyncIOScheduler(timezone="UTC")
            scheduler.add_job(_run_renewal_check, "cron", hour=8, minute=0)
            scheduler.add_job(_check_season_end, "cron", hour=0, minute=5)
            scheduler.start()
            logging.info("APScheduler started — renewal check 08:00 UTC, season check 00:05 UTC")
        except Exception as e:
            logging.error(f"APScheduler failed to start (non-fatal): {e}")

# ============ HAFIZ: CEO SECURITY DASHBOARD ============

@api_router.get("/ceo/security/dashboard")
async def hafizz_security_dashboard(user: User = Depends(get_current_user)):
    if user.user_id != CEO_EMAIL and not (await db.users.find_one({"user_id": user.user_id, "email": CEO_EMAIL})):
        ceo_user = await db.users.find_one({"email": CEO_EMAIL}, {"user_id": 1})
        if not ceo_user or user.user_id != ceo_user.get("user_id"):
            raise HTTPException(403, "Yetkisiz")

    since = datetime.now(timezone.utc) - timedelta(hours=24)

    suspicious = await db.hafiz_anomaly.find({"score": {"$gte": 50}}).sort("score", -1).limit(20).to_list(20)
    for s in suspicious:
        s.pop("_id", None)
        s["log"] = s.get("log", [])[-5:]

    vpn_attempts = await db.hafiz_bans.find({"type": "ip", "banned_at": {"$gte": since}}).limit(50).to_list(50)
    for v in vpn_attempts:
        v.pop("_id", None)
        v["banned_at"] = v["banned_at"].isoformat() if isinstance(v.get("banned_at"), datetime) else v.get("banned_at")

    bot_activity = await db.hafiz_anomaly.find(
        {"log.reason": {"$regex": "fast_task"}, "updated_at": {"$gte": since}}
    ).limit(20).to_list(20)
    for b in bot_activity:
        b.pop("_id", None)

    device_conflicts = await db.hafiz_devices.aggregate([
        {"$group": {"_id": "$fingerprint", "users": {"$addToSet": "$user_id"}, "count": {"$sum": 1}}},
        {"$match": {"count": {"$gte": 3}}},
        {"$limit": 20}
    ]).to_list(20)

    honeypot_hits = await db.hafiz_bans.count_documents({"reason": {"$regex": "^honeypot"}, "banned_at": {"$gte": since}})

    return {
        "suspicious_accounts": suspicious,
        "vpn_attempts_24h": vpn_attempts,
        "bot_activity_24h": bot_activity,
        "device_conflicts": device_conflicts,
        "honeypot_hits_24h": honeypot_hits,
        "generated_at": datetime.now(timezone.utc).isoformat()
    }


@api_router.post("/ceo/security/ban/{target_user_id}")
async def hafizz_ban_user(target_user_id: str, reason: str = Body("Kural ihlali", embed=True), user: User = Depends(get_current_user)):
    ceo_user = await db.users.find_one({"email": CEO_EMAIL}, {"user_id": 1})
    if not ceo_user or user.user_id != ceo_user.get("user_id"):
        raise HTTPException(403, "Yetkisiz")
    await db.users.update_one({"user_id": target_user_id}, {"$set": {"banned": True, "ban_reason": reason, "banned_at": datetime.now(timezone.utc).isoformat()}})
    asyncio.create_task(hafizz.notify_ban(db, target_user_id, reason))
    return {"ok": True}


@api_router.post("/ceo/security/warn/{target_user_id}")
async def hafizz_warn_user(target_user_id: str, reason: str = Body("Şüpheli aktivite", embed=True), user: User = Depends(get_current_user)):
    ceo_user = await db.users.find_one({"email": CEO_EMAIL}, {"user_id": 1})
    if not ceo_user or user.user_id != ceo_user.get("user_id"):
        raise HTTPException(403, "Yetkisiz")
    asyncio.create_task(hafizz.notify_suspicious(db, target_user_id, reason))
    await hafizz.add_anomaly_score(db, target_user_id, 10, f"ceo_warning:{reason[:50]}")
    return {"ok": True}


@api_router.post("/ceo/security/clear/{target_user_id}")
async def hafizz_clear_user(target_user_id: str, user: User = Depends(get_current_user)):
    ceo_user = await db.users.find_one({"email": CEO_EMAIL}, {"user_id": 1})
    if not ceo_user or user.user_id != ceo_user.get("user_id"):
        raise HTTPException(403, "Yetkisiz")
    await db.hafiz_anomaly.update_one({"user_id": target_user_id}, {"$set": {"score": 0, "log": []}})
    await db.users.update_one({"user_id": target_user_id}, {"$unset": {"banned": "", "ban_reason": ""}})
    return {"ok": True}


# ============ HAFIZ: 2FA ============

class TwoFAVerifyRequest(BaseModel):
    user_id: str
    code: str

@api_router.post("/auth/2fa/send")
async def send_2fa_code(request: Request):
    body = await request.json()
    user_id = body.get("user_id", "")
    user = await db.users.find_one({"user_id": user_id}, {"email": 1, "name": 1})
    if not user:
        raise HTTPException(404, "Kullanıcı bulunamadı")
    code = hafizz.generate_2fa_code()
    await hafizz.store_2fa_code(db, user_id, code)
    await hafizz.send_security_email(
        user["email"],
        "ZET Giriş Doğrulama Kodu",
        f"Giriş doğrulama kodunuz: <b style='font-size:24px;letter-spacing:4px'>{code}</b><br/><small>Bu kod 5 dakika geçerlidir.</small>"
    )
    return {"ok": True}


@api_router.post("/auth/2fa/verify")
async def verify_2fa(req: TwoFAVerifyRequest):
    valid = await hafizz.verify_2fa_code(db, req.user_id, req.code)
    if not valid:
        raise HTTPException(400, "Geçersiz veya süresi dolmuş kod")
    return {"ok": True}


# ============ HAFIZ: HONEYPOT ENDPOINTS ============

_HONEYPOT_PATHS = [
    "/api/admin/users/all",
    "/api/admin/dump",
    "/api/debug/config",
    "/api/v1/internal/keys",
    "/api/admin/export",
    "/wp-admin",
    "/phpMyAdmin",
]

@app.api_route("/api/admin/users/all", methods=["GET", "POST"])
@app.api_route("/api/admin/dump", methods=["GET", "POST"])
@app.api_route("/api/debug/config", methods=["GET", "POST"])
@app.api_route("/api/v1/internal/keys", methods=["GET", "POST"])
@app.api_route("/api/admin/export", methods=["GET", "POST"])
@app.api_route("/wp-admin", methods=["GET", "POST"])
@app.api_route("/phpMyAdmin", methods=["GET", "POST"])
async def hafizz_honeypot(request: Request):
    ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "unknown").split(",")[0].strip()
    ua = request.headers.get("User-Agent", "")
    asyncio.create_task(hafizz.record_honeypot_hit(db, ip, str(request.url.path), ua))
    return JSONResponse(status_code=404, content={"detail": "Not found"})


# ─── Quest System ────────────────────────────────────────────────────────────

QUEST_DEFINITIONS = [
    {
        "id": "q1",
        "name": "Belge Oluştur",
        "desc": "İlk belgenizi oluşturun.",
        "zp": 220,
        "stat": "docs_created",
        "threshold": 1,
        "requires": [],
    },
    {
        "id": "q2",
        "name": "Zeta ile Konuş",
        "desc": "Zeta AI ile bir konuşma başlatın.",
        "zp": 220,
        "stat": "ai_chats",
        "threshold": 1,
        "requires": ["q1"],
    },
    {
        "id": "q3",
        "name": "AI ile Görsel Üret",
        "desc": "Zeta Colors ile bir görsel oluşturun.",
        "zp": 220,
        "stat": "ai_images",
        "threshold": 1,
        "requires": ["q1", "q2"],
    },
]

@api_router.post("/quests/auto-check")
async def quests_auto_check(user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id})
    if not user_data:
        return {"newly_pending": []}

    completed = set(user_data.get("completed_quests", []))
    pending = set(user_data.get("pending_quests", []))
    # Pending dahil — görev toplanmasa da bir sonraki kilit açılır
    unlocked = completed | pending

    newly_pending = []
    for q in QUEST_DEFINITIONS:
        qid = q["id"]
        if qid in completed or qid in pending:
            continue
        # Check dependency chain (pending OR collected sayılır)
        if not all(r in unlocked for r in q["requires"]):
            continue
        # Check stat threshold
        stat_val = user_data.get(q["stat"], 0) or 0
        if stat_val >= q["threshold"]:
            pending.add(qid)
            newly_pending.append(qid)

    if newly_pending:
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$addToSet": {"pending_quests": {"$each": newly_pending}}}
        )

    newly_pending_defs = [q for q in QUEST_DEFINITIONS if q["id"] in newly_pending]
    return {"newly_pending": newly_pending_defs}


@api_router.post("/quests/{quest_id}/collect")
async def quest_collect(quest_id: str, user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id})
    if not user_data:
        raise HTTPException(status_code=404, detail="Kullanıcı bulunamadı")

    pending = user_data.get("pending_quests", [])
    if quest_id not in pending:
        raise HTTPException(status_code=400, detail="Bu görev toplanmaya hazır değil")

    quest_def = next((q for q in QUEST_DEFINITIONS if q["id"] == quest_id), None)
    if not quest_def:
        raise HTTPException(status_code=404, detail="Görev bulunamadı")

    zp_reward = quest_def["zp"]
    await db.users.update_one(
        {"user_id": user.user_id},
        {
            "$pull": {"pending_quests": quest_id},
            "$addToSet": {"completed_quests": quest_id},
            "$inc": {"zp": zp_reward},
        }
    )

    updated = await db.users.find_one({"user_id": user.user_id}, {"_id": 0, "zp": 1, "pending_quests": 1, "completed_quests": 1})
    return {"success": True, "zp_earned": zp_reward, "zp_total": updated.get("zp", 0)}


@api_router.get("/quests/status")
async def quests_status(user: User = Depends(get_current_user)):
    user_data = await db.users.find_one({"user_id": user.user_id})
    completed = set(user_data.get("completed_quests", []) if user_data else [])
    pending = set(user_data.get("pending_quests", []) if user_data else [])
    result = []
    for q in QUEST_DEFINITIONS:
        result.append({**q, "status": "collected" if q["id"] in completed else ("pending" if q["id"] in pending else "locked")})
    return {"quests": result}


@api_router.post("/admin/quests/reset-all")
async def admin_reset_all_quests(user: User = Depends(get_current_user)):
    if user.email != CEO_EMAIL:
        raise HTTPException(status_code=403, detail="Yetkisiz")
    result = await db.users.update_many(
        {},
        {"$set": {"completed_quests": [], "pending_quests": []}}
    )
    return {"reset_count": result.modified_count}


app.include_router(api_router)

# ZET Media router
from media_router import media_router, set_media_db
set_media_db(db)
app.include_router(media_router)

frontend_url = os.environ.get("FRONTEND_URL", "http://localhost:3000")

_allowed_origins = list(filter(None, [
    "http://localhost:3000",
    "http://localhost:3001",
    "http://127.0.0.1:3000",
    "https://exciting-comfort-production.up.railway.app",
    "https://zet-mindshare-production.up.railway.app",
    "https://app.zetstudiointl.com",
    "https://zetstudiointl.com",
    os.environ.get("FRONTEND_URL"),
]))

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    import traceback
    logger.error("500 on %s %s: %s\n%s", request.method, request.url.path, exc, traceback.format_exc())
    return JSONResponse(status_code=500, content={"detail": "Internal server error"})

@app.middleware("http")
async def hafizz_ip_ban_middleware(request, call_next):
    ip = request.headers.get("X-Forwarded-For", request.client.host if request.client else "").split(",")[0].strip()
    if ip and await hafizz.is_ip_banned(db, ip):
        return JSONResponse(status_code=403, content={"detail": "Erişim engellendi."})
    return await call_next(request)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=_allowed_origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
# Wed Jun 25 redeploy-trigger
